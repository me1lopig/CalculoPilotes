import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Configuración de la página
st.set_page_config(page_title="Colapsabilidad - Gibbs", layout="wide")

# ==========================================
# UTILIDADES PARA WORD (Sombreado y Fórmulas)
# ==========================================
def set_cell_bg(cell, color_hex):
    """
    Función auxiliar para establecer el color de fondo de una celda en Word.
    """
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def crear_formula_gibbs(doc):
    """
    Inserta la fórmula de Gibbs en formato matemático nativo de Word (XML/OMML).
    Crea una fracción vertical visualmente perfecta.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # XML que define la ecuación matemática: Ig = 2.6 / (1 + 0.026 * LL)
    omml = """
    <m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:oMath>
        <m:r>
            <m:rPr><m:sty m:val="p"/></m:rPr>
            <m:t>Ig = </m:t>
        </m:r>
        <m:f>
          <m:num><m:r><m:t>2.6</m:t></m:r></m:num>
          <m:den>
            <m:r><m:t>1 + 0.026 </m:t></m:r>
            <m:r><m:t>•</m:t></m:r>
            <m:r><m:t> LL</m:t></m:r>
          </m:den>
        </m:f>
      </m:oMath>
    </m:oMathPara>
    """
    p._element.append(parse_xml(omml))

# ==========================================
# FUNCIÓN PARA GENERAR EL informe WORD
# ==========================================
def generar_informe_word(df_resultados, fig_grafico):
    doc = Document()
    
    # --- 1. TÍTULO Y ENCABEZADO ---
    titulo = doc.add_heading('Informe de Colapsabilidad de Suelos', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Análisis basado en el Criterio de Gibbs (1961)')
    run.italic = True
    run.font.size = Pt(12)

    # --- 2. FUNDAMENTO TÉCNICO (EXPLICACIÓN DE FÓRMULA) ---
    doc.add_heading('1. Fundamento Teórico', level=1)
    doc.add_paragraph('El análisis se fundamenta en la relación crítica entre la densidad seca del suelo y su límite líquido. La formulación matemática empleada es:')
    
    # >>> CAMBIO 1: INSERTAR FÓRMULA CIENTÍFICA (XML) <<<
    crear_formula_gibbs(doc)

    doc.add_paragraph('Donde los términos se definen técnicamente como:')
    
    # Lista de definiciones técnicas
    items = [
        ("Ig (Densidad Crítica)", "Representa la densidad seca teórica correspondiente a un estado donde el volumen de vacíos es justo el suficiente para contener el agua correspondiente al Límite Líquido. Si el suelo tiene una densidad menor a esta, su estructura es metaestable."),
        ("LL (Límite Líquido)", "Contenido de humedad expresado en porcentaje bajo el cual el suelo pasa de un comportamiento plástico a un comportamiento viscoso/líquido."),
        ("2.6 (Valor del peso específico relativo)", "Valor constante asumido por el método representando el peso específico relativo media de las partículas sólidas (Gs) para suelos típicos susceptibles al colapso."),
        ("γd (Densidad Seca In-situ)", "Es la densidad seca del suelo en su estado natural. Es el parámetro comparativo determinante.")
    ]
    
    for termino, def_texto in items:
        p_item = doc.add_paragraph(style='List Bullet')
        r_term = p_item.add_run(termino + ": ")
        r_term.bold = True
        p_item.add_run(def_texto)

    # --- 3. TABLA DE RESULTADOS (ESTÉTICA) ---
    doc.add_heading('2. Resumen de Resultados', level=1)
    
    # Crear tabla con autofit
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = True 
    
    # Encabezados
    headers = ['ID Muestra', 'Peso Esp.\n(KN/m³)', 'Densidad\n(T/m³)', 'Lim. Liq\n(LL)', 'DIAGNÓSTICO']
    hdr_cells = table.rows[0].cells
    
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = text
        
        # >>> CAMBIO 2: COLOR AZUL SUAVE EN ENCABEZADOS <<<
        # Hex: #D9E2F3 (Azul claro similar a Word "Emphasis 1")
        set_cell_bg(cell, "D9E2F3") 
        
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(10)

    # Rellenar filas
    for index, row in df_resultados.iterrows():
        row_cells = table.add_row().cells
        
        # Datos normales
        datos = [
            str(row['Identificador']),
            f"{row['Peso Específico Seco (KN/m³)']:.2f}",
            f"{row['γd (T/m³)']:.2f}",
            f"{row['Límite Líquido (LL)']:.1f}"
        ]
        
        for i, dato in enumerate(datos):
            cell = row_cells[i]
            cell.text = dato
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = 1 
        
        # Celda de Estado (Coloreada según resultado)
        estado_cell = row_cells[4]
        estado = row['Estado']
        estado_cell.text = estado
        p_estado = estado_cell.paragraphs[0]
        p_estado.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_estado = p_estado.runs[0]
        run_estado.bold = True
        run_estado.font.size = Pt(9)
        
        # Mantenemos los colores rojo/verde para el diagnóstico porque es información de alerta
        if estado == 'COLAPSABLE':
            set_cell_bg(estado_cell, "FFCCCC") # Rojo claro fondo
            run_estado.font.color.rgb = RGBColor(150, 0, 0)
        else:
            set_cell_bg(estado_cell, "CCFFCC") # Verde claro fondo
            run_estado.font.color.rgb = RGBColor(0, 100, 0)

    # Nota al pie de tabla
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_note = note.add_run('*Diagnóstico basado en la comparación γd vs Ig.')
    run_note.font.size = Pt(8)
    run_note.italic = True

    # --- 4. GRÁFICO ---
    doc.add_heading('3. Representación Gráfica', level=1)
    
    img_stream = BytesIO()
    fig_grafico.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    
    doc.add_picture(img_stream, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar en buffer
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# ==========================================
# INTERFAZ STREAMLIT
# ==========================================
st.sidebar.title("Panel de Control")
st.sidebar.header("📖 Flujo de Trabajo")
st.sidebar.markdown("""
**Instrucciones:**
1. Introduce los datos de tus muestras directamente en la tabla de **"Datos de Entrada"**.
2. Puedes **añadir filas** haciendo clic en la fila vacía al final o en el icono `+`.
3. Puedes **borrar filas** seleccionándolas y pulsando `Supr` o usando el icono de papelera.
4. Los resultados y el gráfico se actualizarán automáticamente.
5. Se puede crear un informe de resultados de los cálculos realizados
""")

st.sidebar.divider()
download_placeholder = st.sidebar.empty()
st.sidebar.info("Cálculo según Gibbs, H.J. y Bara, J.P. (1962).")

# Título Principal
st.title("Análisis de Colapsabilidad (Criterio de Gibbs)")

# Inicializar datos
if 'df_input' not in st.session_state:
    data_inicial = {
        "Identificador": ["C-1", "C-2", "C-3"],
        "Peso Específico Seco (KN/m³)": [12.00, 15.00, 11.50],
        "Límite Líquido (LL)": [30.00, 25.00, 45.00]
    }
    st.session_state.df_input = pd.DataFrame(data_inicial)

# --- 1. TABLA DE ENTRADA ---
st.subheader("1. Introducción de Datos")
edited_df = st.data_editor(
    st.session_state.df_input,
    num_rows="dynamic",
    use_container_width=True,
    column_config={
        "Identificador": st.column_config.TextColumn("ID Muestra"),
        "Peso Específico Seco (KN/m³)": st.column_config.NumberColumn(
            "γd (KN/m³)", min_value=0.0, max_value=30.0, step=0.1, format="%.2f"
        ),
        "Límite Líquido (LL)": st.column_config.NumberColumn(
            "Límite Líquido (LL)", min_value=0.0, max_value=150.0, step=0.5, format="%.1f"
        ),
    },
    key="editor_data"
)

# --- PROCESAMIENTO ---
if not edited_df.empty:
    res_df = edited_df.copy()
    
    # Cálculos
    res_df["γd (T/m³)"] = res_df["Peso Específico Seco (KN/m³)"] / 9.81
    res_df["Densidad Crítica (Ig)"] = 2.6 / (1 + (0.026 * res_df["Límite Líquido (LL)"]))
    
    def determinar_estado(row):
        return "COLAPSABLE" if row["Densidad Crítica (Ig)"] > row["γd (T/m³)"] else "NO COLAPSABLE"

    res_df["Estado"] = res_df.apply(determinar_estado, axis=1)

    # --- 2. VISTA PREVIA RESULTADOS ---
    st.subheader("2. Resultados Calculados")
    
    def estilo_estado(val):
        color = '#d9534f' if val == 'COLAPSABLE' else '#5cb85c' 
        return f'color: {color}; font-weight: bold'

    cols_mostrar = ["Identificador", "γd (T/m³)", "Densidad Crítica (Ig)", "Estado"]
    st.dataframe(
        res_df[cols_mostrar].style.map(estilo_estado, subset=["Estado"])
                                  .format({"γd (T/m³)": "{:.2f}", "Densidad Crítica (Ig)": "{:.2f}"}),
        use_container_width=True
    )

    # --- 3. GRÁFICO ---
    st.markdown("---")
    st.subheader("3. Gráfico de Gibbs")

    # Rango corregido desde 0
    ll_range = np.linspace(0, 100, 200)
    ig_curve = 2.6 / (1 + (0.026 * ll_range))

    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Curva y zonas
    ax.plot(ig_curve, ll_range, label="Límite Teórico", color="#1f77b4", linewidth=2.5)
    ax.fill_betweenx(ll_range, 0, ig_curve, color='red', alpha=0.1, label="Colapsable")
    ax.fill_betweenx(ll_range, ig_curve, 3.5, color='green', alpha=0.1, label="No colapsable")

    # Puntos
    for index, row in res_df.iterrows():
        color_punto = '#d62728' if row["Estado"] == "COLAPSABLE" else '#2ca02c'
        ax.scatter(row["γd (T/m³)"], row["Límite Líquido (LL)"], c=color_punto, s=120, edgecolors="black", zorder=5)
        # Etiqueta
        ax.text(row["γd (T/m³)"] + 0.03, row["Límite Líquido (LL)"], f" {row['Identificador']}", 
                fontsize=10, fontweight='bold', verticalalignment='center',
                bbox=dict(facecolor='white', alpha=0.6, edgecolor='none', pad=1))

    ax.set_xlabel("Peso Específico Seco (T/m³)", fontsize=11)
    ax.set_ylabel("Límite Líquido (LL%)", fontsize=11)
    ax.set_title("Criterio de Colapsabilidad de Gibbs", fontsize=13)
    ax.set_xlim(0.5, 2.5)
    ax.set_ylim(0, 100)
    ax.grid(True, linestyle='--', alpha=0.5)
    ax.legend(loc="upper right")
    
    st.pyplot(fig)

    # --- GENERACIÓN DEL informe ---
    try:
        archivo_word = generar_informe_word(res_df, fig)
        
        download_placeholder.download_button(
            label="📄 Descargar Informe  (.docx)",
            data=archivo_word,
            file_name="Informe_Gibbs.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    except Exception as e:
        st.error(f"Error generando el informe: {e}")

else:
    st.info("Introduce datos en la tabla para comenzar el análisis.")