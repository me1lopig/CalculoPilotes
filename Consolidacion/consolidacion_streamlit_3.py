import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import io
import json
from datetime import datetime
import warnings

import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches

try:
    warnings.simplefilter('ignore', np.exceptions.RankWarning)
except AttributeError:
    warnings.simplefilter('ignore', np.RankWarning)

st.set_page_config(page_title="Consolidación 1D", layout="wide")
st.title("Modelo de Consolidación 1D Avanzado")

# --- Función de Seguridad (Callback) ---
def reset_estado():
    st.session_state.calculado = False
    st.session_state.docs_generados = False

# --- Inicializar memoria y parámetros por defecto ---
if 'calculado' not in st.session_state:
    st.session_state.calculado = False
if 'docs_generados' not in st.session_state:
    st.session_state.docs_generados = False

default_params = {
    'longitud': 10.0, 'Ti': 100.0, 'c': 0.05, 'mv': 0.0002, 
    'h': 1.0, 'k': 1.0, 'max_U_pct': 95.0, 'intervalo_dias_curvas': 10.0, 
    'tipo_calculo': 1, 'metodo_numerico': "Explícito (Original)"
}

for key, val in default_params.items():
    if key not in st.session_state:
        st.session_state[key] = val

# --- PANEL LATERAL ---
st.sidebar.header("📁 Gestión de Perfiles")
uploaded_file = st.sidebar.file_uploader("Cargar Configuración (.json)", type=["json"], on_change=reset_estado)
if uploaded_file is not None:
    try:
        loaded_params = json.load(uploaded_file)
        for key in default_params.keys():
            if key in loaded_params:
                st.session_state[key] = loaded_params[key]
        reset_estado()
        st.sidebar.success("Perfil cargado correctamente.")
    except Exception:
        st.sidebar.error("Error al leer el archivo JSON.")

st.sidebar.header("Parámetros de Entrada")
longitud = st.sidebar.number_input("Espesor del estrato [m]", min_value=0.1, step=0.5, key='longitud', on_change=reset_estado)
Ti = st.sidebar.number_input("Carga exterior (Ti) [kPa]", min_value=1.0, step=10.0, key='Ti', on_change=reset_estado)
T0 = 0.0 
TL = 0.0 

c = st.sidebar.number_input("Coef. consolidación (Cv) [m2/día]", min_value=1e-8, format="%.5e", key='c', on_change=reset_estado)
mv = st.sidebar.number_input("Coef. compresibilidad vol. (mv) [m2/kN]", min_value=1e-8, format="%.5e", key='mv', on_change=reset_estado)

st.sidebar.header("Datos del Mallado y Visualización")
h = st.sidebar.number_input("Incremento de x (h) [m]", min_value=0.01, step=0.1, key='h', on_change=reset_estado)
k = st.sidebar.number_input("Incremento de t (k) [días]", min_value=0.01, step=0.5, key='k', on_change=reset_estado)
max_U_pct = st.sidebar.slider("Máximo grado de consolidación a calcular [%]", min_value=10.0, max_value=99.9, key='max_U_pct', on_change=reset_estado)

intervalo_dias_curvas = st.sidebar.number_input("Mostrar isócronas cada (días):", min_value=0.1, step=1.0, key='intervalo_dias_curvas', on_change=reset_estado)

tipo_calculo = st.sidebar.selectbox(
    "Tipos de condiciones de contorno", options=[1, 2, 3],
    format_func=lambda x: {1: "[1] Permeable - Permeable (Doble)", 2: "[2] Permeable - Impermeable (Sup.)", 3: "[3] Impermeable - Permeable (Inf.)"}[x],
    key='tipo_calculo', on_change=reset_estado
)
str_contorno = {1: "Doble Drenaje", 2: "Drenaje Superior", 3: "Drenaje Inferior"}[tipo_calculo]

metodo_numerico = st.sidebar.radio(
    "Motor Numérico", 
    options=["Explícito (Original)", "Implícito (Incondicionalmente Estable)"],
    key='metodo_numerico', on_change=reset_estado
)

perfil_actual = {key: st.session_state[key] for key in default_params.keys()}
st.sidebar.download_button(
    label="💾 Guardar Configuración (.json)",
    data=json.dumps(perfil_actual, indent=4),
    file_name="perfil_consolidacion.json",
    mime="application/json"
)

# --- PESTAÑAS ---
tab_simulacion, tab_datos, tab_teoria = st.tabs(["⚙️ Simulación y Resultados", "📊 Datos Tabulados", "📚 Fundamento Teórico"])

# ==========================================
# PESTAÑA 1: SIMULACIÓN
# ==========================================
with tab_simulacion:
    
    if st.button("🚀 Iniciar Cálculo del Modelo", type="primary"):
        s_max = longitud * mv * Ti
        permeabilidad = c * mv * 10  
        alfa = c * k / (h**2)
        
        if metodo_numerico == "Explícito (Original)" and alfa > 0.5:
            st.error(f"Error: El método explícito no es convergente para alfa={alfa:.3f} (debe ser ≤ 0.5). Reduce el incremento temporal 'k', aumenta 'h', o cambia al Método Implícito.")
            st.session_state.calculado = False
        else:
            nx = int(np.floor(longitud / h))
            x = np.arange(0, longitud + h/2, h)
            if len(x) != nx + 1: x = np.linspace(0, longitud, nx + 1)
                
            u0 = np.ones(nx + 1) * Ti
            u = np.zeros(nx + 1)
            
            hist_t, hist_Q, hist_S, hist_U = [], [], [], []
            hist_presiones_completas = [(0.0, np.copy(u0))] 
            grado_consolidacion = 0.0
            t = 0.0 
            progreso = st.progress(0)

            # --- MOTOR IMPLÍCITO ---
            if metodo_numerico == "Implícito (Incondicionalmente Estable)":
                A = np.zeros((nx + 1, nx + 1))
                for i in range(1, nx):
                    A[i, i-1] = -alfa
                    A[i, i]   = 1 + 2 * alfa
                    A[i, i+1] = -alfa
                
                if tipo_calculo == 1:
                    A[0, 0] = 1; A[nx, nx] = 1
                    u0[0] = T0; u0[nx] = TL
                elif tipo_calculo == 2:
                    A[0, 0] = 1
                    A[nx, nx-1] = -2 * alfa; A[nx, nx] = 1 + 2 * alfa
                    u0[0] = T0
                elif tipo_calculo == 3:
                    A[0, 0] = 1 + 2 * alfa; A[0, 1] = -2 * alfa
                    A[nx, nx] = 1
                    u0[nx] = TL

                A_inv = np.linalg.inv(A)

                while grado_consolidacion <= max_U_pct / 100.0:
                    t += k
                    u = np.dot(A_inv, u0)
                    
                    if tipo_calculo == 1: u[0] = T0; u[nx] = TL
                    elif tipo_calculo == 2: u[0] = T0
                    elif tipo_calculo == 3: u[nx] = TL
                        
                    u0 = np.copy(u)
                    hist_presiones_completas.append((t, np.copy(u0)))
                    
                    if tipo_calculo in [1, 2]: derivada = u0[1] / h 
                    else: derivada = abs(u0[nx] - u0[nx-1]) / h 
                    Q = permeabilidad * derivada
                    
                    uajuste = np.polyfit(x, u0, 2)
                    sumasimpson = sum([((x[ix+1] - x[ix]) / 6.0) * (u0[ix+1] + u0[ix] + 4 * np.polyval(uajuste, (x[ix+1] + x[ix]) * 0.5)) for ix in range(len(x) - 1)])
                    
                    area_total = longitud * Ti
                    grado_consolidacion = (area_total - sumasimpson) / area_total
                    asiento = s_max * grado_consolidacion
                    
                    hist_t.append(t); hist_Q.append(Q); hist_S.append(asiento * 100); hist_U.append(grado_consolidacion * 100) 
                    progreso.progress(min(1.0, grado_consolidacion / (max_U_pct / 100.0)))
                    if t > 500000: break

            # --- MOTOR EXPLÍCITO ---
            else:
                if tipo_calculo == 1: 
                    u0[0] = (T0 + Ti) / 2; u0[nx] = (TL + Ti) / 2
                    for i in range(1, nx): u[i] = alfa * (u0[i+1] + u0[i-1] - 2*u0[i]) + u0[i]
                    u[0] = T0; u[nx] = TL
                elif tipo_calculo == 2: 
                    u0[0] = 0
                    for i in range(1, nx): u[i] = alfa * (u0[i+1] + u0[i-1] - 2*u0[i]) + u0[i]
                    u[nx] = alfa * (u0[nx-1] + u0[nx-1] - 2*u0[nx]) + u0[nx]
                elif tipo_calculo == 3: 
                    u0[0] = alfa * (2 * u0[1] - 2 * u0[0]) + u0[0]
                    for i in range(1, nx): u[i] = alfa * (u0[i+1] + u0[i-1] - 2*u0[i]) + u0[i]
                    u0[nx] = 0; u[nx] = 0

                while grado_consolidacion <= max_U_pct / 100.0:
                    t += k
                    for i in range(1, nx): u[i] = alfa * (u0[i+1] + u0[i-1]) + (1 - 2*alfa) * u0[i]
                    
                    if tipo_calculo == 2: u[nx] = alfa * (u0[nx-1] + u0[nx-1] - 2*u0[nx]) + u0[nx]
                    elif tipo_calculo == 3: u[0] = alfa * (2 * u0[1] - 2 * u0[0]) + u0[0]; u[nx] = 0
                        
                    u0 = np.copy(u)
                    hist_presiones_completas.append((t, np.copy(u0)))
                    
                    if tipo_calculo in [1, 2]: derivada = u0[1] / h 
                    else: derivada = abs(u0[nx] - u0[nx-1]) / h 
                    Q = permeabilidad * derivada
                    
                    uajuste = np.polyfit(x, u0, 2)
                    sumasimpson = sum([((x[ix+1] - x[ix]) / 6.0) * (u0[ix+1] + u0[ix] + 4 * np.polyval(uajuste, (x[ix+1] + x[ix]) * 0.5)) for ix in range(len(x) - 1)])
                        
                    area_total = longitud * Ti
                    grado_consolidacion = (area_total - sumasimpson) / area_total
                    asiento = s_max * grado_consolidacion
                    
                    hist_t.append(t); hist_Q.append(Q); hist_S.append(asiento * 100); hist_U.append(grado_consolidacion * 100) 
                    progreso.progress(min(1.0, grado_consolidacion / (max_U_pct / 100.0)))
                    if t > 500000: break
                    
            progreso.empty()
            st.success(f"Cálculos completados exitosamente en {t:.1f} días usando el motor {metodo_numerico.split(' ')[0]}.")
            st.info(f"Para un grado de consolidación de **{hist_U[-1]:.2f} %**, el asiento es **{hist_S[-1]:.2f} cm**.")

            historial_isocronas = [(0.0, hist_presiones_completas[0][1])] 
            tiempo_objetivo = intervalo_dias_curvas
            for tiempo_pres, u0_pres in hist_presiones_completas[1:]:
                if tiempo_pres >= tiempo_objetivo:
                    historial_isocronas.append((tiempo_pres, u0_pres))
                    tiempo_objetivo += intervalo_dias_curvas

            st.session_state.resultados = {
                'hist_t': np.array(hist_t), 'hist_Q': np.array(hist_Q), 'hist_S': np.array(hist_S), 'hist_U': np.array(hist_U),
                'x': x, 'historial_isocronas': historial_isocronas, 'hist_presiones_completas': hist_presiones_completas,
                'c': c, 'longitud': longitud, 'Ti': Ti, 'mv': mv, 's_max': s_max, 'permeabilidad': permeabilidad, 
                'str_contorno': str_contorno, 'intervalo_dias': intervalo_dias_curvas,
                'fecha': datetime.now().strftime("%d_%m_%y_%H_%M_%S")
            }
            st.session_state.calculado = True
            st.session_state.docs_generados = False 

    # --- MOSTRAR LOS RESULTADOS ---
    if st.session_state.calculado:
        res = st.session_state.resultados
        factor_tiempo = res['hist_t'] * res['c'] / (res['longitud']**2)

        st.markdown("---")
        st.subheader("Gráficas Interactivas de Resultados")

        col1, col2 = st.columns(2)
        
        with col1:
            # --- MEJORA: Gráfica de Presiones con Gradiente de Color y Sin Leyenda Múltiple ---
            fig_iso = go.Figure()
            num_iso = len(res['historial_isocronas'])
            cmap = plt.get_cmap('viridis') # Mapa de color científico (de oscuro a claro)

            for i, (tiempo_iso, u_iso) in enumerate(res['historial_isocronas']):
                if tiempo_iso == 0.0:
                    color = 'red'
                    dash = 'dash'
                    show_l = True # Mostramos solo el t=0 en la leyenda por referencia
                else:
                    rgba = cmap(i / max(1, num_iso - 1))
                    color = f'rgba({int(rgba[0]*255)}, {int(rgba[1]*255)}, {int(rgba[2]*255)}, 0.8)'
                    dash = 'solid'
                    show_l = False # Ocultamos la leyenda para evitar la saturación
                    
                fig_iso.add_trace(go.Scatter(
                    x=u_iso, y=res['x'], mode='lines', 
                    line=dict(color=color, dash=dash), 
                    name=f't={tiempo_iso:.1f}d',
                    showlegend=show_l,
                    hoverinfo='name+x+y' # Al pasar el ratón se ve todo claro
                ))
            fig_iso.update_layout(
                title=f"Presión de poro (cada {res['intervalo_dias']} días)", 
                xaxis_title="Presión de poro [kPa]", 
                yaxis_title="Profundidad [m]", 
                yaxis=dict(autorange="reversed")
            )
            st.plotly_chart(fig_iso, use_container_width=True)

            fig_S_t = go.Figure()
            fig_S_t.add_trace(go.Scatter(x=res['hist_t'], y=res['hist_S'], mode='lines', line=dict(color='red')))
            fig_S_t.update_layout(title="Asientos vs Tiempo", xaxis_title="Tiempo [días]", yaxis_title="Asientos [cm]", yaxis=dict(autorange="reversed"))
            st.plotly_chart(fig_S_t, use_container_width=True)
            
            fig_U_t = go.Figure()
            fig_U_t.add_trace(go.Scatter(x=res['hist_t'], y=res['hist_U'], mode='lines', line=dict(color='brown')))
            fig_U_t.update_layout(title="Grado de Consolidación vs Tiempo", xaxis_title="Tiempo [días]", yaxis_title="Grado de Consolidación [%]", yaxis=dict(autorange="reversed"))
            st.plotly_chart(fig_U_t, use_container_width=True)

        with col2:
            fig_Q_t = go.Figure()
            fig_Q_t.add_trace(go.Scatter(x=res['hist_t'], y=res['hist_Q'], mode='lines', line=dict(color='green')))
            fig_Q_t.update_layout(title="Caudal vs Tiempo", xaxis_title="Tiempo [días]", yaxis_title="Caudal [m3/día/m2]")
            st.plotly_chart(fig_Q_t, use_container_width=True)

            fig_S_U = go.Figure()
            fig_S_U.add_trace(go.Scatter(x=res['hist_U'], y=res['hist_S'], mode='lines', line=dict(color='orange')))
            fig_S_U.update_layout(title="Asientos vs Grado de Consolidación", xaxis_title="Grado de Consolidación [%]", yaxis_title="Asientos [cm]", yaxis=dict(autorange="reversed"))
            st.plotly_chart(fig_S_U, use_container_width=True)
            
            fig_U_Tv = go.Figure()
            fig_U_Tv.add_trace(go.Scatter(x=factor_tiempo, y=res['hist_U'], mode='lines', line=dict(color='purple')))
            fig_U_Tv.update_layout(title="Grado Consolidación vs Factor Tiempo", xaxis_title="Factor Tiempo (Tv)", yaxis_title="Grado de Consolidación [%]", xaxis_type="log", yaxis=dict(autorange="reversed"))
            st.plotly_chart(fig_U_Tv, use_container_width=True)

        # --- BOTÓN 2: GENERAR DOCUMENTOS ---
        st.markdown("---")
        st.subheader("Generación de Documentos")
        
        if st.button("📝 Generar Informes (Word y Excel)"):
            with st.spinner("Creando documentos (generando imágenes vectoriales en segundo plano)..."):
                
                imagenes_buffers = []
                def save_plt(fig_plt):
                    buf = io.BytesIO()
                    fig_plt.savefig(buf, format="png", bbox_inches="tight", dpi=150)
                    buf.seek(0)
                    imagenes_buffers.append(buf)
                    plt.close(fig_plt)
                
                # --- MEJORA MATPLOTLIB (WORD): Gráfica de Presiones con Gradiente ---
                f_iso, ax = plt.subplots(figsize=(6,4))
                num_iso_w = len(res['historial_isocronas'])
                cmap_w = plt.get_cmap('viridis')
                
                for i, (t_iso, u_iso) in enumerate(res['historial_isocronas']):
                    if t_iso == 0.0:
                        ax.plot(u_iso, res['x'], color='red', linestyle='--', label='t=0')
                    else:
                        ax.plot(u_iso, res['x'], color=cmap_w(i / max(1, num_iso_w - 1)), alpha=0.7)
                
                ax.invert_yaxis(); ax.set_title("Presión de poro"); save_plt(f_iso)

                f_St, ax = plt.subplots(figsize=(6,4)); ax.plot(res['hist_t'], res['hist_S'], color='red'); ax.invert_yaxis(); ax.set_title("Asientos vs Tiempo"); save_plt(f_St)
                f_Ut, ax = plt.subplots(figsize=(6,4)); ax.plot(res['hist_t'], res['hist_U'], color='brown'); ax.invert_yaxis(); ax.set_title("U vs Tiempo"); save_plt(f_Ut)
                f_Qt, ax = plt.subplots(figsize=(6,4)); ax.plot(res['hist_t'], res['hist_Q'], color='green'); ax.set_title("Caudal vs Tiempo"); save_plt(f_Qt)
                f_SU, ax = plt.subplots(figsize=(6,4)); ax.plot(res['hist_U'], res['hist_S'], color='orange'); ax.invert_yaxis(); ax.set_title("Asientos vs U"); save_plt(f_SU)
                f_UTv, ax = plt.subplots(figsize=(6,4)); ax.semilogx(factor_tiempo, res['hist_U'], color='purple'); ax.invert_yaxis(); ax.set_title("U vs Factor Tiempo"); save_plt(f_UTv)

                excel_buffer = io.BytesIO()
                df_params = pd.DataFrame({"Parámetro": ["Espesor", "Carga", "Cv", "mv", "Permeabilidad", "Asiento Max", "Contorno"], "Valor": [res['longitud'], res['Ti'], res['c'], res['mv'], res['permeabilidad'], res['s_max']*100, res['str_contorno']]})
                df_evolucion = pd.DataFrame({"Tiempo [dias]": res['hist_t'], "U [%]": res['hist_U'], "Asientos [cm]": res['hist_S'], "Caudal": res['hist_Q']})
                columnas_profundidad = [f"z={xi:.2f}m" for xi in res['x']]
                df_presiones = pd.DataFrame([p[1] for p in res['hist_presiones_completas']], columns=columnas_profundidad)
                df_presiones.insert(0, "Tiempo [dias]", [p[0] for p in res['hist_presiones_completas']])

                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    df_params.to_excel(writer, sheet_name='1. Parámetros', index=False)
                    df_evolucion.to_excel(writer, sheet_name='2. Evolución temporal', index=False)
                    df_presiones.to_excel(writer, sheet_name='3. Presiones Isocronas', index=False)
                
                doc = Document()
                doc.add_heading('Informe de Consolidación 1D', 0)
                doc.add_heading('Datos del Modelo', level=1)
                for index, row in df_params.iterrows(): doc.add_paragraph(f"• {row['Parámetro']}: {row['Valor']}")
                doc.add_heading('Resultados Gráficos', level=1)
                for img_buf in imagenes_buffers: doc.add_picture(img_buf, width=Inches(5.5))
                word_buffer = io.BytesIO()
                doc.save(word_buffer)

                st.session_state.excel_data = excel_buffer.getvalue()
                st.session_state.word_data = word_buffer.getvalue()
                st.session_state.docs_generados = True

        if st.session_state.docs_generados:
            st.success("¡Documentos generados! Puedes descargarlos aquí.")
            col_dw1, col_dw2 = st.columns(2)
            with col_dw1: st.download_button("📊 Descargar Excel", data=st.session_state.excel_data, file_name=f"Datos_{res['fecha']}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")
            with col_dw2: st.download_button("📄 Descargar Word", data=st.session_state.word_data, file_name=f"Informe_{res['fecha']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
    else:
        st.info("👈 Configura los parámetros en el panel lateral y pulsa **'Iniciar Cálculo del Modelo'** para ver las gráficas y generar informes.")

# ==========================================
# PESTAÑA 2: DATOS TABULADOS
# ==========================================
with tab_datos:
    st.header("Datos Numéricos de la Simulación")
    if st.session_state.calculado:
        res = st.session_state.resultados
        st.subheader("1. Evolución Temporal")
        st.dataframe(pd.DataFrame({"Tiempo [días]": res['hist_t'], "Grado Consolidación [%]": res['hist_U'], "Asientos [cm]": res['hist_S'], "Caudal [m3/día/m2]": res['hist_Q']}), use_container_width=True, height=300)
        
        st.subheader("2. Matriz de Presiones Intersticiales (Isócronas)")
        df_presiones_tab = pd.DataFrame([p[1] for p in res['hist_presiones_completas']], columns=[f"z={xi:.2f}m" for xi in res['x']])
        df_presiones_tab.insert(0, "Tiempo [días]", [p[0] for p in res['hist_presiones_completas']])
        st.dataframe(df_presiones_tab, use_container_width=True, height=400)
    else:
        st.info("ℹ️ Inicia el cálculo en la pestaña de Simulación para ver los datos tabulados.")

# ==========================================
# PESTAÑA 3: FUNDAMENTO TEÓRICO AMPLIADO
# ==========================================
with tab_teoria:
    st.header("Teoría de la Consolidación Unidimensional de Terzaghi")
    st.write("""
    La teoría de la consolidación unidimensional, formulada por Karl von Terzaghi (1925), describe cómo el exceso de presión 
    intersticial del agua generado por una carga externa se disipa a lo largo del tiempo. Este proceso provoca una reducción 
    de volumen en los suelos de grano fino saturados, traduciéndose en asientos en la superficie.
    """)

    st.subheader("1. La Ecuación Diferencial Gobernadora")
    st.write("""
    Partiendo de las leyes de flujo continuo de Darcy y la conservación de masa, Terzaghi dedujo la siguiente ecuación 
    diferencial en derivadas parciales (EDP) parabólica:
    """)
    st.latex(r" \frac{\partial u}{\partial t} = c_v \frac{\partial^2 u}{\partial x^2} ")
    st.write("""
    **Donde:**
    * **$u$**: Exceso de presión intersticial de poro [kPa].
    * **$t$**: Tiempo [días].
    * **$x$**: Coordenada de profundidad geométrica [m].
    * **$c_v$**: Coeficiente de consolidación [m²/día], el cual agrupa las propiedades hidro-mecánicas del suelo:
    """)
    st.latex(r" c_v = \frac{k_{perm}}{\gamma_w \cdot m_v} ")
    st.write("""
    *(Nota: En este software, el incremento temporal numérico se denomina $k$ y el incremento espacial $h$. Para evitar confusiones, nos referimos a la permeabilidad de Darcy como $k_{perm}$. En el código, se asume un peso específico del agua $\gamma_w \approx 10 \text{ kN/m}^3$).*
    """)

    st.markdown("---")
    st.subheader("2. Aproximación Numérica por Diferencias Finitas")
    st.write("""
    Dado que la solución analítica de la ecuación de Terzaghi mediante series de Fourier es muy compleja para condiciones 
    de contorno variables, este software implementa el **Método de Diferencias Finitas (FDM)**. 
    
    A través de las expansiones de las Series de Taylor, podemos aproximar las derivadas continuas convirtiendo el estrato 
    continuo en una malla discreta de nodos $i$ separados por una distancia **$h$**, y el tiempo continuo en incrementos **$k$**.

    Para la derivada espacial de segundo grado (usando **Diferencias Centrales**):
    """)
    st.latex(r" \frac{\partial^2 u}{\partial x^2} \approx \frac{u_{i-1} - 2u_i + u_{i+1}}{h^2} ")

    st.markdown("---")
    st.subheader("3. Modelos de Resolución Temporal")

    st.markdown("#### A. Método Explícito (Euler hacia adelante)")
    st.write("""
    Si evaluamos la derivada espacial en el tiempo presente ($t$) y la derivada temporal hacia adelante, la ecuación se despeja 
    de manera directa (explícita) para predecir el futuro:
    """)
    st.latex(r" \frac{u_i^{t+k} - u_i^t}{k} = c_v \frac{u_{i+1}^t - 2u_i^t + u_{i-1}^t}{h^2} ")
    st.write("Agrupando términos y definiendo el Módulo de Fourier o factor de estabilidad como **$\\alpha = \frac{c_v \cdot k}{h^2}$**:")
    st.latex(r" u_i^{t+k} = \alpha u_{i+1}^t + (1 - 2\alpha) u_i^t + \alpha u_{i-1}^t ")
    st.info("⚠️ **Problema de Estabilidad:** Este método es computacionalmente rápido pero condicionalmente estable. Si el término $(1 - 2\\alpha)$ se vuelve negativo, la matemática oscila hacia el infinito. Por tanto, exige estrictamente que **$\\alpha \le 0.5$**.")

    st.markdown("#### B. Método Implícito (Euler hacia atrás)")
    st.write("""
    Para eliminar la restricción de estabilidad, evaluamos la derivada espacial en el instante futuro ($t+k$). 
    """)
    st.latex(r" \frac{u_i^{t+k} - u_i^t}{k} = c_v \frac{u_{i+1}^{t+k} - 2u_i^{t+k} + u_{i-1}^{t+k}}{h^2} ")
    st.write("Al reorganizar, nos encontramos con tres incógnitas del futuro en una sola ecuación:")
    st.latex(r" -\alpha u_{i-1}^{t+k} + (1+2\alpha)u_i^{t+k} - \alpha u_{i+1}^{t+k} = u_i^t ")
    st.write("""
    Esto genera un sistema de ecuaciones simultáneas que se representa matricialmente de forma tridiagonal ($[A] \cdot \{u^{t+k}\} = \{u^t\}$):
    """)
    st.latex(r"""
    \begin{bmatrix}
    (1+2\alpha) & -\alpha & 0 & \dots & 0 \\
    -\alpha & (1+2\alpha) & -\alpha & \dots & 0 \\
    0 & -\alpha & (1+2\alpha) & \dots & 0 \\
    \vdots & \vdots & \vdots & \ddots & \vdots \\
    0 & 0 & 0 & -\alpha & (1+2\alpha)
    \end{bmatrix}
    \begin{bmatrix}
    u_1^{t+k} \\ u_2^{t+k} \\ u_3^{t+k} \\ \vdots \\ u_n^{t+k}
    \end{bmatrix}
    =
    \begin{bmatrix}
    u_1^t \\ u_2^t \\ u_3^t \\ \vdots \\ u_n^t
    \end{bmatrix}
    """)
    st.success("✅ **Incondicionalmente Estable:** Este método invierte la matriz $[A]$ en cada paso. Aunque consume ligeramente más memoria, permite dar pasos de tiempo $k$ gigantescos sin que el modelo colapse jamás, siendo ideal para consolidaciones a largo plazo.")

    st.markdown("---")
    st.subheader("4. Cálculo Físico Final (Grado de Consolidación y Asientos)")
    st.write("""
    A medida que la presión intersticial $u$ disminuye transfiriendo la carga al esqueleto sólido, el suelo se comprime. 
    El grado de consolidación medio del estrato en un instante de tiempo, **$U(t)$**, se calcula integrando el área bajo 
    la curva isócrona real versus el área de carga inicial rectangular:
    """)
    st.latex(r" U(t) = 1 - \frac{\int_0^H u(x,t) \, dx}{\int_0^H u(x,0) \, dx} ")
    st.write("""
    En el código, esta integral no se resuelve con trapecios simples, sino ajustando un polinomio de segundo grado 
    (`polyfit`) entre nodos e integrando con la Regla de Simpson para obtener una precisión altísima.

    Finalmente, el asiento físico en centímetros para ese instante se obtiene escalando el asiento máximo esperado ($S_{max}$) 
    por el grado de consolidación:
    """)
    st.latex(r" S_{max} = H \cdot m_v \cdot T_i ")
    st.latex(r" S(t) = S_{max} \cdot U(t) ")