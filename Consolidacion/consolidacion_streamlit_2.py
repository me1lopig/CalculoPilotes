import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import io
from datetime import datetime
import warnings

from docx import Document
from docx.shared import Inches

# Ignoramos el RankWarning de polyfit (compatible con NumPy antiguo y 2.0+)
try:
    warnings.simplefilter('ignore', np.exceptions.RankWarning)
except AttributeError:
    warnings.simplefilter('ignore', np.RankWarning)

st.set_page_config(page_title="Consolidación 1D", layout="wide")
st.title("Modelo de Consolidación 1D (Diferencias Finitas)")

# --- Inicializar variables de estado (Memoria de Streamlit) ---
if 'calculado' not in st.session_state:
    st.session_state.calculado = False
if 'docs_generados' not in st.session_state:
    st.session_state.docs_generados = False

# --- 1. ENTRADA DE DATOS (Panel Lateral) ---
st.sidebar.header("Parámetros de Entrada")
longitud = st.sidebar.number_input("Espesor del estrato [m]", min_value=0.1, value=10.0, step=0.5)
Ti = st.sidebar.number_input("Carga exterior (Ti) [kPa]", min_value=1.0, value=100.0, step=10.0)
T0 = 0.0 
TL = 0.0 

c = st.sidebar.number_input("Coef. consolidación (Cv) [m2/día]", min_value=0.001, value=0.05, step=0.01, format="%.3f")
mv = st.sidebar.number_input("Coef. compresibilidad vol. (mv) [m2/kN]", min_value=0.00001, value=0.0002, step=0.00005, format="%.5f")

st.sidebar.header("Datos del Mallado y Visualización")
h = st.sidebar.number_input("Incremento de x (h) [m]", min_value=0.1, value=1.0, step=0.1)
k = st.sidebar.number_input("Incremento de t (k) [días]", min_value=0.1, value=1.0, step=0.5)
max_U_pct = st.sidebar.slider("Máximo grado de consolidación a calcular [%]", min_value=10.0, max_value=99.9, value=95.0)

intervalo_dias_curvas = st.sidebar.number_input("Mostrar curvas isócronas cada (días):", min_value=0.1, value=10.0, step=1.0)

tipo_calculo = st.sidebar.selectbox(
    "Tipos de condiciones de contorno",
    options=[1, 2, 3],
    format_func=lambda x: {
        1: "[1] Permeable - Permeable (Doble)",
        2: "[2] Permeable - Impermeable (Sup.)",
        3: "[3] Impermeable - Permeable (Inf.)"
    }[x]
)

str_contorno = {1: "Doble Drenaje", 2: "Drenaje Superior", 3: "Drenaje Inferior"}[tipo_calculo]

# --- CREACIÓN DE PESTAÑAS (Ahora son 3) ---
tab_simulacion, tab_datos, tab_teoria = st.tabs(["⚙️ Simulación y Resultados", "📊 Datos Tabulados", "📚 Fundamento Teórico"])

# ==========================================
# PESTAÑA 1: SIMULACIÓN
# ==========================================
with tab_simulacion:
    
    # BOTÓN 1: SOLO CALCULAR
    if st.button("🚀 Iniciar Cálculo del Modelo", type="primary"):
        s_max = longitud * mv * Ti
        permeabilidad = c * mv * 10  
        alfa = c * k / (h**2)
        
        if alfa > 0.5:
            st.error(f"El modelo no es convergente (alfa={alfa:.3f} > 0.5). Disminuye k o aumenta h.")
            st.session_state.calculado = False
        else:
            # Inicialización de la malla
            nx = int(np.floor(longitud / h))
            x = np.arange(0, longitud + h/2, h)
            if len(x) != nx + 1:
                x = np.linspace(0, longitud, nx + 1)
                
            u0 = np.ones(nx + 1) * Ti
            u = np.zeros(nx + 1)
            
            # Arranque de contornos
            if tipo_calculo == 1: 
                u0[0] = (T0 + Ti) / 2
                u0[nx] = (TL + Ti) / 2
                for i in range(1, nx):
                    u[i] = alfa * (u0[i+1] + u0[i-1] - 2*u0[i]) + u0[i]
                u[0] = T0
                u[nx] = TL
            elif tipo_calculo == 2: 
                u0[0] = 0
                for i in range(1, nx):
                    u[i] = alfa * (u0[i+1] + u0[i-1] - 2*u0[i]) + u0[i]
                u[nx] = alfa * (u0[nx-1] + u0[nx-1] - 2*u0[nx]) + u0[nx]
            elif tipo_calculo == 3: 
                u0[0] = alfa * (2 * u0[1] - 2 * u0[0]) + u0[0]
                for i in range(1, nx):
                    u[i] = alfa * (u0[i+1] + u0[i-1] - 2*u0[i]) + u0[i]
                u0[nx] = 0
                u[nx] = 0

            grado_consolidacion = 0.0
            t = 1.0 
            
            hist_t, hist_Q, hist_S, hist_U = [], [], [], []
            hist_presiones_completas = [(0.0, np.ones(nx + 1) * Ti)] 
            
            progreso = st.progress(0)
            
            # --- BUCLE PRINCIPAL ---
            while grado_consolidacion <= max_U_pct / 100.0:
                t += k
                
                for i in range(1, nx):
                    u[i] = alfa * (u0[i+1] + u0[i-1]) + (1 - 2*alfa) * u0[i]
                
                if tipo_calculo == 2:
                    u[nx] = alfa * (u0[nx-1] + u0[nx-1] - 2*u0[nx]) + u0[nx]
                elif tipo_calculo == 3:
                    u[0] = alfa * (2 * u0[1] - 2 * u0[0]) + u0[0]
                    u[nx] = 0
                    
                u0 = np.copy(u)
                hist_presiones_completas.append((t, np.copy(u0)))
                
                if tipo_calculo in [1, 2]:
                    derivada = u0[1] / h 
                else:
                    derivada = abs(u0[nx] - u0[nx-1]) / h 
                Q = permeabilidad * derivada
                
                uajuste = np.polyfit(x, u0, 2)
                sumasimpson = 0
                for ix in range(len(x) - 1):
                    punto_medio = (x[ix+1] + x[ix]) * 0.5
                    mid_val = np.polyval(uajuste, punto_medio)
                    termino = ((x[ix+1] - x[ix]) / 6.0) * (u0[ix+1] + u0[ix] + 4 * mid_val)
                    sumasimpson += termino
                    
                area_total = longitud * Ti
                grado_consolidacion = (area_total - sumasimpson) / area_total
                asiento = s_max * grado_consolidacion
                
                hist_t.append(t)
                hist_Q.append(Q)
                hist_S.append(asiento * 100)
                hist_U.append(grado_consolidacion * 100) 
                
                progreso.progress(min(1.0, grado_consolidacion / (max_U_pct / 100.0)))
                
                if t > 50000:
                    break
                    
            progreso.empty()
            st.success(f"Cálculos completados exitosamente en {t:.1f} días.")
            st.info(f"Para un grado de consolidación de **{hist_U[-1]:.2f} %**, el asiento es **{hist_S[-1]:.2f} cm**.")

            historial_isocronas = [(0.0, hist_presiones_completas[0][1])] 
            tiempo_objetivo = intervalo_dias_curvas
            
            for tiempo_pres, u0_pres in hist_presiones_completas[1:]:
                if tiempo_pres >= tiempo_objetivo:
                    historial_isocronas.append((tiempo_pres, u0_pres))
                    tiempo_objetivo += intervalo_dias_curvas

            st.session_state.resultados = {
                'hist_t': np.array(hist_t), 'hist_Q': np.array(hist_Q),
                'hist_S': np.array(hist_S), 'hist_U': np.array(hist_U),
                'x': x, 'historial_isocronas': historial_isocronas,
                'hist_presiones_completas': hist_presiones_completas,
                'c': c, 'longitud': longitud, 'Ti': Ti, 'mv': mv, 's_max': s_max,
                'permeabilidad': permeabilidad, 'str_contorno': str_contorno,
                'intervalo_dias': intervalo_dias_curvas,
                'fecha': datetime.now().strftime("%d_%m_%y_%H_%M_%S")
            }
            st.session_state.calculado = True
            st.session_state.docs_generados = False 

    # --- MOSTRAR LOS RESULTADOS (GRÁFICOS) ---
    if st.session_state.calculado:
        res = st.session_state.resultados
        factor_tiempo = res['hist_t'] * res['c'] / (res['longitud']**2)

        st.markdown("---")
        st.subheader("Gráficas de Resultados")

        col1, col2 = st.columns(2)
        
        with col1:
            fig6, ax6 = plt.subplots(figsize=(6, 4))
            for tiempo_iso, u_iso in res['historial_isocronas']:
                if tiempo_iso == 0.0:
                    ax6.plot(u_iso, res['x'], color='red', linestyle='--', label='t=0')
                else:
                    ax6.plot(u_iso, res['x'], color='blue', alpha=0.4)
            ax6.set_title(f"Presión de poro (cada {res['intervalo_dias']} días)")
            ax6.set_xlabel("Presión de poro [kPa]")
            ax6.set_ylabel("Profundidad [m]")
            ax6.invert_yaxis()
            ax6.legend(loc="upper right")
            st.pyplot(fig6)

            fig2, ax2 = plt.subplots(figsize=(6, 4))
            ax2.plot(res['hist_t'], res['hist_S'], color='red')
            ax2.set_title("Asientos vs Tiempo")
            ax2.set_ylabel("Asientos [cm]")
            ax2.set_xlabel("Tiempo [días]")
            ax2.invert_yaxis()
            st.pyplot(fig2)
            
            fig4, ax4 = plt.subplots(figsize=(6, 4))
            ax4.plot(res['hist_t'], res['hist_U'], color='brown')
            ax4.set_title("Grado de Consolidación vs Tiempo")
            ax4.set_xlabel("Tiempo [días]")
            ax4.set_ylabel("Grado de Consolidación [%]")
            ax4.invert_yaxis()
            st.pyplot(fig4)

        with col2:
            fig1, ax1 = plt.subplots(figsize=(6, 4))
            ax1.plot(res['hist_t'], res['hist_Q'], color='green')
            ax1.set_title("Caudal vs Tiempo")
            ax1.set_ylabel("Caudal [m3/día/m2]")
            ax1.set_xlabel("Tiempo [días]")
            st.pyplot(fig1)

            fig3, ax3 = plt.subplots(figsize=(6, 4))
            ax3.plot(res['hist_U'], res['hist_S'], color='orange')
            ax3.set_title("Asientos vs Grado de Consolidación")
            ax3.set_ylabel("Asientos [cm]")
            ax3.set_xlabel("Grado de Consolidación [%]")
            ax3.invert_yaxis()
            st.pyplot(fig3)
            
            fig5, ax5 = plt.subplots(figsize=(6, 4))
            ax5.semilogx(factor_tiempo, res['hist_U'], color='purple')
            ax5.set_title("Grado de Consolidación vs Factor Tiempo")
            ax5.set_xlabel("Factor Tiempo")
            ax5.set_ylabel("Grado de Consolidación [%]")
            ax5.invert_yaxis()
            st.pyplot(fig5)

        # --- BOTÓN 2: GENERAR DOCUMENTOS ---
        st.markdown("---")
        st.subheader("Generación de Documentos")
        
        if st.button("📝 Generar Informes (Word y Excel)"):
            with st.spinner("Creando documentos, por favor espera..."):
                
                imagenes_buffers = []
                def save_fig(fig):
                    buf = io.BytesIO()
                    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
                    buf.seek(0)
                    imagenes_buffers.append(buf)
                    
                save_fig(fig6); save_fig(fig2); save_fig(fig4)
                save_fig(fig1); save_fig(fig3); save_fig(fig5)

                # Generar Excel
                excel_buffer = io.BytesIO()
                df_params = pd.DataFrame({
                    "Parámetro": ["Espesor estrato [m]", "Carga exterior [kPa]", "Coef. Consolidación [m2/dia]", 
                                  "Coef. Compresibilidad [m2/kN]", "Coef. Permeabilidad [m/dia]", 
                                  "Asiento Máximo [cm]", "Condición de Contorno"],
                    "Valor": [res['longitud'], res['Ti'], res['c'], res['mv'], res['permeabilidad'], res['s_max']*100, res['str_contorno']]
                })
                
                df_evolucion = pd.DataFrame({
                    "Tiempo [dias]": res['hist_t'], "Grado Consolidación [%]": res['hist_U'],
                    "Asientos [cm]": res['hist_S'], "Caudal [m3/dia/m2]": res['hist_Q']
                })
                
                columnas_profundidad = [f"z={xi:.2f}m" for xi in res['x']]
                datos_presiones = [p[1] for p in res['hist_presiones_completas']]
                df_presiones = pd.DataFrame(datos_presiones, columns=columnas_profundidad)
                df_presiones.insert(0, "Tiempo [dias]", [p[0] for p in res['hist_presiones_completas']])

                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    df_params.to_excel(writer, sheet_name='1. Parámetros', index=False)
                    df_evolucion.to_excel(writer, sheet_name='2. Evolución temporal', index=False)
                    df_presiones.to_excel(writer, sheet_name='3. Presiones Isocronas', index=False)
                
                # Generar Word
                doc = Document()
                doc.add_heading('Informe de Consolidación 1D', 0)
                doc.add_heading('Datos del Modelo', level=1)
                for index, row in df_params.iterrows():
                    doc.add_paragraph(f"• {row['Parámetro']}: {row['Valor']}")
                doc.add_heading('Resultados Gráficos', level=1)
                for img_buf in imagenes_buffers:
                    doc.add_picture(img_buf, width=Inches(5.5))
                    
                word_buffer = io.BytesIO()
                doc.save(word_buffer)

                st.session_state.excel_data = excel_buffer.getvalue()
                st.session_state.word_data = word_buffer.getvalue()
                st.session_state.docs_generados = True

        if st.session_state.docs_generados:
            st.success("¡Documentos generados! Puedes descargarlos aquí.")
            col_dw1, col_dw2 = st.columns(2)
            with col_dw1:
                st.download_button(
                    label="📊 Descargar Excel Multipestaña", data=st.session_state.excel_data,
                    file_name=f"Datos_Consolidacion_{res['fecha']}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary"
                )
            with col_dw2:
                st.download_button(
                    label="📄 Descargar Informe en Word", data=st.session_state.word_data,
                    file_name=f"Informe_Consolidacion_{res['fecha']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary"
                )


# ==========================================
# PESTAÑA 2: DATOS TABULADOS
# ==========================================
with tab_datos:
    st.header("Datos Numéricos de la Simulación")
    
    if st.session_state.calculado:
        res = st.session_state.resultados
        st.write("Explora de manera interactiva las tablas generadas por el modelo.")
        
        st.subheader("1. Evolución Temporal")
        # Creamos el dataframe de evolución
        df_evolucion_tab = pd.DataFrame({
            "Tiempo [días]": res['hist_t'],
            "Grado Consolidación [%]": res['hist_U'],
            "Asientos [cm]": res['hist_S'],
            "Caudal [m3/día/m2]": res['hist_Q']
        })
        st.dataframe(df_evolucion_tab, use_container_width=True, height=300)
        
        st.subheader("2. Matriz de Presiones Intersticiales (Isócronas)")
        # Creamos el dataframe de la matriz de presiones
        columnas_profundidad_tab = [f"z={xi:.2f}m" for xi in res['x']]
        datos_presiones_tab = [p[1] for p in res['hist_presiones_completas']]
        df_presiones_tab = pd.DataFrame(datos_presiones_tab, columns=columnas_profundidad_tab)
        df_presiones_tab.insert(0, "Tiempo [días]", [p[0] for p in res['hist_presiones_completas']])
        
        st.dataframe(df_presiones_tab, use_container_width=True, height=400)
        
    else:
        st.info("ℹ️ Aún no hay datos disponibles. Ve a la pestaña de **Simulación y Resultados** e inicia el cálculo para ver las tablas aquí.")


# ==========================================
# PESTAÑA 3: FUNDAMENTO TEÓRICO
# ==========================================
with tab_teoria:
    st.header("Teoría de la Consolidación Unidimensional de Terzaghi")
    st.write("""
    Este programa resuelve numéricamente la ecuación diferencial que gobierna el proceso de consolidación unidimensional 
    formulado por Karl von Terzaghi. El modelo simula cómo se disipa la presión intersticial (o presión de poro) del agua 
    en un estrato de suelo fino (como las arcillas) a lo largo del tiempo, provocando el asiento del terreno.
    """)

    st.subheader("1. La Ecuación Diferencial")
    st.write("La ley fundamental de consolidación establece que el cambio de la presión intersticial en el tiempo es proporcional a su concavidad espacial:")
    st.latex(r" \frac{\partial u}{\partial t} = c_v \frac{\partial^2 u}{\partial x^2} ")
    st.write("""
    Donde:
    * **$u$**: Presión intersticial de poro [kPa].
    * **$t$**: Tiempo [días].
    * **$x$**: Profundidad dentro del estrato [m].
    * **$c_v$**: Coeficiente de consolidación del suelo [m²/día].
    """)

    st.subheader("2. Resolución Numérica: Diferencias Finitas Explícitas")
    st.write("""
    Para que el ordenador pueda resolver esta ecuación, transformamos las derivadas continuas en diferencias discretas 
    (pasos finitos de tiempo $\Delta t$ o `k`, y de espacio $\Delta x$ o `h`).
    
    Aproximando la derivada espacial con diferencias centrales y la temporal hacia adelante, la presión de un nodo $i$ 
    en el siguiente instante de tiempo se calcula como:
    """)
    st.latex(r" u_i^{t+k} = u_i^t + \alpha \left( u_{i+1}^t - 2u_i^t + u_{i-1}^t \right) ")
    
    st.write("Donde **$\alpha$** es el factor de estabilidad de la malla o módulo de Fourier:")
    st.latex(r" \alpha = \frac{c_v \cdot k}{h^2} ")
    
    st.info("💡 **Condición de Convergencia:** Para que el método explícito sea estable y no diverja matemáticamente, el valor de $\\alpha$ debe ser obligatoriamente **$\\alpha \leq 0.5$**.")

    st.subheader("3. Cálculo del Grado de Consolidación y Asientos")
    st.write("A medida que el agua se drena y las presiones bajan, el terreno se asienta. El **grado de consolidación ($U$)** representa qué porcentaje de ese proceso se ha completado. Matemáticamente, es la relación entre el área de presiones disipadas y el área inicial total.")
    st.latex(r" U = 1 - \frac{\int_0^H u(x,t) \, dx}{\int_0^H u(x,0) \, dx} ")
    st.write("En el código, esta integral se resuelve iterativamente usando un ajuste parabólico y la regla de Simpson combinada.")
    
    st.write("Finalmente, el **asiento en el tiempo ($S$)** es proporcional al asiento máximo esperado:")
    st.latex(r" S_{max} = H \cdot m_v \cdot \Delta \sigma ")
    st.latex(r" S(t) = S_{max} \cdot U(t) ")
    st.write("""
    Donde:
    * **$H$**: Espesor del estrato.
    * **$m_v$**: Coeficiente de compresibilidad volumétrico.
    * **$\Delta \sigma$**: Incremento de la carga exterior ($T_i$).
    """)