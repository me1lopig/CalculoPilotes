import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import io
import xlsxwriter
import zipfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 0. CONFIGURACIÓN Y ESTADO
# ==========================================
st.set_page_config(page_title="Geotecnia Pro Informe", layout="wide", page_icon="📝")

if 'archivo_listo' not in st.session_state:
    st.session_state.archivo_listo = False
if 'buffer_zip' not in st.session_state:
    st.session_state.buffer_zip = None

def resetear_descarga():
    st.session_state.archivo_listo = False
    st.session_state.buffer_zip = None

# ==========================================
# 1. FUNCIONES AUXILIARES (WORD)
# ==========================================

def dataframe_a_tabla_word(doc, df):
    """Convierte un DataFrame de Pandas en una tabla de Word con estilo."""
    t = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    t.style = 'Table Grid'

    # --- ENCABEZADOS ---
    for j, col_name in enumerate(df.columns):
        cell = t.cell(0, j)
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9) # Fuente un poco más pequeña para que quepa bien
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- DATOS ---
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            val = df.iloc[i, j]
            cell = t.cell(i + 1, j)
            
            if isinstance(val, (int, float)):
                cell.text = f"{val:.2f}"
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                cell.text = str(val) 
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

def generar_informe_word(df_in, df_out, img_stream, nf_val, usar_nf):
    """Crea el objeto Documento de Word completo."""
    doc = Document()
    
    # Título
    titulo = doc.add_heading('Memoria de Cálculo: Tensiones Verticales', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sección 1: Datos de Entrada
    doc.add_heading('1. Modelo Geotécnico Definido', level=1)
    doc.add_paragraph('Estratigrafía y parámetros de identificación:')
    dataframe_a_tabla_word(doc, df_in)
    
    doc.add_paragraph() 
    p_nf = doc.add_paragraph()
    runner = p_nf.add_run('Condición NF: ')
    runner.font.bold = True
    if usar_nf:
        p_nf.add_run(f"Nivel Freático considerado a {nf_val:.2f} m de profundidad.")
    else:
        p_nf.add_run("Sin Nivel Freático (Suelo no saturado).")

    # Sección 2: Gráfica
    doc.add_heading('2. Perfil de Tensiones', level=1)
    img_stream.seek(0) 
    doc.add_picture(img_stream, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sección 3: Resultados por ahora no se pone tarda mucho
    #doc.add_heading('3. Tabla de Resultados (Listado Completo)', level=1)
    #doc.add_paragraph('A continuación se detallan los valores de tensión calculados metro a metro (paso 0.10m):')
    
    # --- CAMBIO IMPORTANTE: Sin límite de filas ---
    #dataframe_a_tabla_word(doc, df_out)

    return doc

# ==========================================
# 2. MOTOR DE CÁLCULO
# ==========================================

def obtener_maximo_menor(lista, valor):
    filtro = [x for x in lista if x <= valor]
    if not filtro: return 0
    return max(filtro)

def parametro_terreno(cotas, zt):
    if zt <= 0: return 1
    indice = 0
    for z in range(len(cotas)-1):
        cota_sup = cotas[z]
        cota_inf = cotas[z+1]
        if zt > cota_sup and zt <= cota_inf:
            indice = z + 1
            break
    if indice == 0 and zt > cotas[-1]:
        indice = len(cotas) - 1
    return indice

def presion_total(cotas, valor_nf, pe_saturado, pe_seco, valor_cota):
    lista_cotas = sorted(list(set(cotas + [valor_nf]))) 
    resultado = obtener_maximo_menor(lista_cotas, valor_cota)
    idx_terreno = parametro_terreno(cotas, valor_cota)
    
    if idx_terreno >= len(pe_saturado): idx_terreno = len(pe_saturado) - 1

    p_sat_val = pe_saturado[idx_terreno]
    p_seco_val = pe_seco[idx_terreno]
    
    peso = p_seco_val if valor_cota <= valor_nf else p_sat_val
    presion_acumulada = (valor_cota - resultado) * peso
    
    try: idx_resultado = lista_cotas.index(resultado)
    except ValueError: idx_resultado = 0
    
    for j in range(idx_resultado, 0, -1):
        cota_actual = lista_cotas[j]
        cota_anterior = lista_cotas[j-1]
        espesor = cota_actual - cota_anterior
        cota_media = (cota_actual + cota_anterior) / 2
        idx_t = parametro_terreno(cotas, cota_media)
        
        if idx_t >= len(pe_saturado): idx_t = len(pe_saturado) - 1
        
        p_sat_capa = pe_saturado[idx_t]
        p_seco_capa = pe_seco[idx_t]
        
        if cota_actual <= valor_nf: peso_tramo = p_seco_capa
        else: peso_tramo = p_sat_capa
        presion_acumulada += espesor * peso_tramo
        
    return presion_acumulada

def calcular_perfil_tensiones(df_input, nivel_freatico, usar_nf_logic):
    espesores = [0] + df_input['Espesor (m)'].tolist()
    pe_seco = [0] + df_input['Peso ap. (kN/m³)'].tolist()
    pe_saturado = [0] + df_input['Peso Sat. (kN/m³)'].tolist()
    
    cotas = []
    acumulado = 0
    for esp in espesores:
        acumulado += esp
        cotas.append(acumulado)
    
    max_z = max(cotas)
    if max_z == 0: return pd.DataFrame()
    
    paso = 0.1
    profundidades = [round(x, 2) for x in np.arange(0, max_z + paso, paso)]
    resultados = []
    nf_calculo = nivel_freatico if usar_nf_logic else 99999.9

    for z in profundidades:
        if z > max_z: continue 
        sigma_v = presion_total(cotas, nf_calculo, pe_saturado, pe_seco, z)
        if z >= nf_calculo: u = (z - nf_calculo) * 9.81
        else: u = 0
        sigma_eff = sigma_v - u
        resultados.append({
            "Profundidad (m)": z,
            "Presión de Poro (kPa)": round(u, 2),
            "Presión Efectiva (kPa)": round(sigma_eff, 2),
            "Presión Total (kPa)": round(sigma_v, 2)
        })
    return pd.DataFrame(resultados)

# ==========================================
# 3. INTERFAZ DE USUARIO
# ==========================================

st.title("🌍 Calculadora de Tensiones Verticales")

# --- ENTRADA DE DATOS ---
col1, col2 = st.columns([3, 1])

datos_iniciales = pd.DataFrame([
    {"Unidad": "UG-01", "Espesor (m)": 5.0, "Peso ap. (kN/m³)": 18.0, "Peso Sat. (kN/m³)": 20.0},
    {"Unidad": "UG-02", "Espesor (m)": 10.0, "Peso ap. (kN/m³)": 19.0, "Peso Sat. (kN/m³)": 21.0},
])

with col1:
    st.subheader("1. Estratigrafía")
    df_estratos = st.data_editor(
        datos_iniciales,
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "Unidad": st.column_config.TextColumn("Unidad Geotécnica", width="medium", required=True),
            "Espesor (m)": st.column_config.NumberColumn("Espesor (m)", min_value=0.1, format="%.2f"),
            "Peso ap. (kN/m³)": st.column_config.NumberColumn("γ ap. (kN/m³)", min_value=1.0, format="%.1f"),
            "Peso Sat. (kN/m³)": st.column_config.NumberColumn("γ sat. (kN/m³)", min_value=1.0, format="%.1f"),
        },
        key="editor_datos" 
    )

if 'data_anterior' not in st.session_state:
    st.session_state.data_anterior = df_estratos.to_json()

if df_estratos.to_json() != st.session_state.data_anterior:
    st.session_state.data_anterior = df_estratos.to_json()
    resetear_descarga()

prof_total = 0.0
if not df_estratos.empty:
    prof_total = df_estratos["Espesor (m)"].sum()

with col2:
    st.markdown("##### Resumen")
    st.metric("Profundidad Total", f"{prof_total:.2f} m")

# --- SIDEBAR ---
with st.sidebar:
    st.header("💧 Condiciones NF")
    usar_nf = st.checkbox("Considerar Nivel Freático", value=True, on_change=resetear_descarga)
    
    if usar_nf:
        max_limit = float(prof_total) if prof_total > 0 else 100.0
        if 'nf_input' in st.session_state:
            if st.session_state.nf_input > max_limit:
                st.session_state.nf_input = max_limit
                st.toast(f"⚠️ Nivel freático ajustado a {max_limit}m", icon="💧")
        
        nf_input = st.number_input(
            "Profundidad Nivel Freático (m)", 
            min_value=0.0, 
            max_value=max_limit,  
            value=3.5, 
            step=0.1,
            on_change=resetear_descarga,
            key='nf_input',
            format="%.2f"
        )
    else:
        nf_input = 0.0
        st.info("Cálculo sin presión de poros.")

# --- CÁLCULO ---
if not df_estratos.empty:
    
    df_resultados = calcular_perfil_tensiones(df_estratos, nf_input, usar_nf)
    
    st.divider()
    
    tab_graf, tab_data = st.tabs(["📈 Gráfica", "📋 Tabla"])
    
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.plot(df_resultados["Presión Efectiva (kPa)"], df_resultados["Profundidad (m)"], 'darkred', label='Tensión Efectiva')
    ax.plot(df_resultados["Presión Total (kPa)"], df_resultados["Profundidad (m)"], 'blue', linestyle='--', label='Tensión Total')
    ax.plot(df_resultados["Presión de Poro (kPa)"], df_resultados["Profundidad (m)"], 'green', linestyle=':', label='Presión de Poro')
    
    if usar_nf and nf_input <= prof_total:
        ax.axhline(y=nf_input, color='cyan', linestyle='-.', label=f'N.F. ({nf_input:.2f}m)')
    
    ax.invert_yaxis()
    ax.set_xlabel('Tensión [kPa]'); ax.set_ylabel('Profundidad [m]')
    ax.legend(); ax.grid(True, linestyle='--', alpha=0.6)

    with tab_graf:
        col_g1, col_g2 = st.columns([1, 4])
        with col_g2: st.pyplot(fig)
        
    with tab_data:
        st.dataframe(df_resultados, use_container_width=True, height=300)

    # --- GENERACIÓN DE ARCHIVOS ---
    st.divider()
    st.subheader("📥 Exportar Informe Completo")
    
    col_gen, col_down = st.columns([1, 1])
    
    with col_gen:
        if st.button("🔄 Generar Archivos (Excel + Word + PNG)", type="primary"):
            
            # 1. IMAGEN 
            img_buffer = io.BytesIO()
            fig.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            
            # 2. EXCEL 
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
                df_resultados.to_excel(writer, index=False, sheet_name='Resultados')
                workbook = writer.book
                worksheet_res = writer.sheets['Resultados']
                formato_num = workbook.add_format({'num_format': '0.00', 'align': 'center'})
                worksheet_res.set_column('A:D', 18, formato_num)
                
                df_estratos.to_excel(writer, index=False, sheet_name='Datos de Entrada')
                worksheet_in = writer.sheets['Datos de Entrada']
                worksheet_in.set_column('A:A', 25) 
                worksheet_in.set_column('B:D', 18, formato_num)

            # 3. WORD (Ahora incluye tabla completa)
            doc_object = generar_informe_word(df_estratos, df_resultados, img_buffer, nf_input, usar_nf)
            word_buffer = io.BytesIO()
            doc_object.save(word_buffer)
            
            # 4. ZIP 
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                zf.writestr("Resultados_Calculo.xlsx", excel_buffer.getvalue())
                zf.writestr("Informe.docx", word_buffer.getvalue())
                zf.writestr("Grafica_Tensiones.png", img_buffer.getvalue())
            
            st.session_state.buffer_zip = zip_buffer.getvalue()
            st.session_state.archivo_listo = True
            st.rerun() 

    with col_down:
        if st.session_state.archivo_listo:
            st.success("✅ Informe generado.")
            st.download_button(
                label="📦 Descargar Informe (ZIP)",
                data=st.session_state.buffer_zip,
                file_name="Informe.zip",
                mime="application/zip",
                type="secondary"
            )
        else:
            st.info("👆 Pulsa para generar el paquete de documentos.")

else:
    st.warning("Introduce datos en la tabla para comenzar.")