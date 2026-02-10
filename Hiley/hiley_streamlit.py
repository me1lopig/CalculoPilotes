import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
import zipfile
from datetime import datetime

# LIBRERÍAS PARA WORD Y ESTILOS AVANZADOS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn

# LIBRERÍAS EXCEL
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

# --- 1. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="Fórmula de Hiley para DPSH",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 2. ESTADO DE LA SESIÓN ---
if 'resultados' not in st.session_state:
    st.session_state.resultados = None
if 'zip_buffer' not in st.session_state:
    st.session_state.zip_buffer = None
if 'word_buffer' not in st.session_state:
    st.session_state.word_buffer = None
    
if 'config_tramos' not in st.session_state:
    st.session_state.config_tramos = {}
if 'errores_detectados' not in st.session_state:
    st.session_state.errores_detectados = False

# --- 3. CONSTANTES ---
MATERIALES_F = {
    "Arenas o Gravas (F=25)": 25.0,
    "SM - Arenas limosas (F=30)": 30.0,
    "ML - Limos baja plast. (F=35)": 35.0,
    "CL - Arcillas baja plast. (F=40)": 40.0,
    "CH - Arcillas alta plast. (F=50)": 50.0
}

# --- 4. FUNCIONES AUXILIARES DE ESTILO WORD ---
def set_table_header_bg_color(cell, color_hex):
    """
    Aplica un color de fondo a una celda de tabla en Word modificando el XML.
    color_hex: string sin '#' (ej: "D9E1F2")
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# --- 5. FUNCIONES DE GENERACIÓN ---

def _aplicar_estilo_hoja(ws, headers, titulo='Cálculos Hiley'):
    thin_side = Side(style='thin', color='B0B0B0')
    border_thin = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    fill_title = PatternFill('solid', fgColor='1F4E79')
    fill_header = PatternFill('solid', fgColor='D9E1F2')
    font_title = Font(bold=True, color='FFFFFF', size=14)
    font_header = Font(bold=True, color='1F1F1F')
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center', wrap_text=True)

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    cell_title = ws.cell(row=1, column=1, value=titulo)
    cell_title.fill = fill_title
    cell_title.font = font_title
    cell_title.alignment = align_center

    for j, h in enumerate(headers, start=1):
        c = ws.cell(row=2, column=j, value=h)
        c.fill = fill_header
        c.font = font_header
        c.alignment = align_center
        c.border = border_thin
    
    col_widths = {1: 22, 2: 10, 3: 10, 4: 15, 5: 10, 6: 10, 7: 10, 8: 10, 9: 24, 10: 24}
    for col_idx, w in col_widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = w
    
    ws.freeze_panes = 'B3'
    return border_thin, align_left, align_center

def _escribir_df_en_hoja(ws, df_calc, headers, titulo):
    border_thin, align_left, align_center = _aplicar_estilo_hoja(ws, headers, titulo=titulo)
    start_row = 3
    for i in range(len(df_calc)):
        row_vals = df_calc.iloc[i].tolist()
        for j, val in enumerate(row_vals, start=1):
            c = ws.cell(row=start_row + i, column=j, value=val)
            c.border = border_thin
            if j == 1:
                c.alignment = align_left
            else:
                c.alignment = align_center
                if isinstance(val, (int, float, np.integer, np.floating)):
                    if j in [5, 6, 7, 8]: # a, e, n, c
                        c.number_format = '0.000'
                    else:
                        c.number_format = '0.00'

def generar_zip_en_memoria(df_resultados, ensayos_unicos):
    output_zip = BytesIO()
    excel_buffer = BytesIO()
    wb = Workbook()
    ws_def = wb.active
    wb.remove(ws_def)
    headers = list(df_resultados.columns)
    
    with zipfile.ZipFile(output_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
        for ensayo in ensayos_unicos:
            df_e = df_resultados[df_resultados["Ensayo"] == str(ensayo)].sort_values("Z(m)")
            if df_e.empty: continue
            
            # Excel
            ws_e = wb.create_sheet(title=str(ensayo)[:31])
            _escribir_df_en_hoja(ws_e, df_e, headers, titulo=f'Cálculos {ensayo}')
            
            # Gráficos
            plt.style.use('default')
            
            fig1, ax1 = plt.subplots(figsize=(6, 5))
            ax1.plot(df_e['Presión Admisible (kg/cm2)'], df_e['Z(m)'], marker='o', linewidth=2)
            ax1.invert_yaxis()
            ax1.grid(True, alpha=0.3)
            ax1.set_xlabel('Presión Admisible ($kg/cm^2$)')
            ax1.set_ylabel('Profundidad Z (m)')
            ax1.set_title(f'{ensayo} - Presión Admisible')
            img_buf1 = BytesIO()
            plt.tight_layout()
            plt.savefig(img_buf1, format='png', dpi=150)
            plt.close()
            zf.writestr(f"graficos/Pa_vs_Z_{ensayo}.png", img_buf1.getvalue())
            
            fig2, ax2 = plt.subplots(figsize=(6, 5))
            ax2.plot(df_e['N(20)'], df_e['Z(m)'], marker='o', linewidth=2, color='#8B0000')
            ax2.invert_yaxis()
            ax2.grid(True, alpha=0.3)
            ax2.set_xlabel('N(20) (Golpes)')
            ax2.set_ylabel('Profundidad Z (m)')
            ax2.set_title(f'{ensayo} - Golpeos')
            img_buf2 = BytesIO()
            plt.tight_layout()
            plt.savefig(img_buf2, format='png', dpi=150)
            plt.close()
            zf.writestr(f"graficos/Golpeos_vs_Z_{ensayo}.png", img_buf2.getvalue())

        wb.save(excel_buffer)
        zf.writestr("Resultados_Hiley.xlsx", excel_buffer.getvalue())
    
    return output_zip

def generar_word_en_memoria(df_resultados, config_tramos, ensayos_unicos):
    doc = Document()
    
    # 1. ESTILO GLOBAL
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # 2. ENCABEZADO TIPO INFORME TÉCNICO
    # Título Principal
    titulo = doc.add_heading('Informe de resultados', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = titulo.runs[0]
    run_tit.font.color.rgb = RGBColor(0, 0, 0) # Negro
    run_tit.font.bold = True
    
    # Metadatos alineados a la derecha
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_fecha = p_meta.add_run(f'Fecha de emisión: {datetime.now().strftime("%d/%m/%Y")}\n')
    run_metod = p_meta.add_run("Metodología: Fórmula de Hiley (DPSH)")
    run_fecha.italic = True
    run_metod.bold = True
    
    doc.add_paragraph() # Espacio vacío simple para separar

    # 3. CONTENIDO POR ENSAYO
    for i, ensayo in enumerate(ensayos_unicos):
        # Título del Ensayo (Azul profesional)
        h1 = doc.add_heading(f'ENSAYO: {ensayo}', level=1)
        h1.runs[0].font.color.rgb = RGBColor(31, 78, 121) 
        
        # --- A. DATOS DE ENTRADA ---
        doc.add_heading('1.Parámetros del Suelo', level=2)
        
        if ensayo in config_tramos:
            df_input = config_tramos[ensayo]
            
            # Tabla Estilizada
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.autofit = False 
            
            # Encabezados
            hdr_cells = table.rows[0].cells
            headers = ['Prof. Desde (m)', 'Prof. Hasta (m)', 'Tipo de Material (F)']
            
            for j, text in enumerate(headers):
                cell = hdr_cells[j]
                cell.text = text
                # Estilo: Negrita + Fondo Gris Claro
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                set_table_header_bg_color(cell, "D9D9D9") 
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Datos
            for _, row in df_input.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = f"{row['Desde (m)']:.2f}"
                row_cells[1].text = f"{row['Hasta (m)']:.2f}"
                row_cells[2].text = str(row['Material'])
                
                # Centrado
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph("Se han utilizado valores por defecto.", style='Intense Quote')

        doc.add_paragraph()

        # --- B. RESULTADOS ---
        doc.add_heading('2. Resultados de Capacidad de Carga', level=2)
        
        df_e = df_resultados[df_resultados["Ensayo"] == str(ensayo)].sort_values("Z(m)")
        
        if not df_e.empty:
            table_res = doc.add_table(rows=1, cols=3)
            table_res.style = 'Table Grid'
            
            hdr_res = table_res.rows[0].cells
            headers_res = ['Profundidad Z (m)', 'Golpeos N(20)', 'P. Admisible (kg/cm²)']
            
            for j, text in enumerate(headers_res):
                cell = hdr_res[j]
                cell.text = text
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                set_table_header_bg_color(cell, "D9E1F2") # Azul claro
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for _, row in df_e.iterrows():
                row_cells = table_res.add_row().cells
                row_cells[0].text = f"{row['Z(m)']:.2f}"
                row_cells[1].text = f"{row['N(20)']:.0f}"
                row_cells[2].text = f"{row['Presión Admisible (kg/cm2)']:.2f}"
                
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_note = doc.add_paragraph(f"Total de registros calculados: {len(df_e)}")
            p_note.style = 'Caption'

        # --- C. GRÁFICAS ---
        doc.add_heading('3. Gráficas de Análisis', level=2)
        plt.style.use('default')
        
        # Gráfica 1: Presión
        fig1, ax1 = plt.subplots(figsize=(6.5, 4))
        ax1.plot(df_e['Presión Admisible (kg/cm2)'], df_e['Z(m)'], marker='o', markersize=4, linewidth=2, color='#1F4E79')
        ax1.invert_yaxis()
        ax1.grid(True, linestyle='--', alpha=0.5)
        ax1.set_xlabel('Presión Admisible ($kg/cm^2$)', fontsize=10)
        ax1.set_ylabel('Profundidad Z (m)', fontsize=10)
        ax1.set_title(f'Perfil de Presiones - {ensayo}', fontsize=12, fontweight='bold')
        plt.tight_layout()
        
        mem_fig1 = BytesIO()
        plt.savefig(mem_fig1, format='png', dpi=150)
        plt.close()
        
        doc.add_picture(mem_fig1, width=Inches(6.0))
        doc.add_paragraph("Fig 1. Evolución de la Presión Admisible.", style='Caption')
        doc.add_paragraph()

        # Gráfica 2: Golpeos
        fig2, ax2 = plt.subplots(figsize=(6.5, 4))
        ax2.plot(df_e['N(20)'], df_e['Z(m)'], marker='s', markersize=4, linewidth=2, color='#C00000')
        ax2.invert_yaxis()
        ax2.grid(True, linestyle='--', alpha=0.5)
        ax2.set_xlabel('N(20) (Golpes)', fontsize=10)
        ax2.set_ylabel('Profundidad Z (m)', fontsize=10)
        ax2.set_title(f'Perfil de Golpeos - {ensayo}', fontsize=12, fontweight='bold')
        plt.tight_layout()
        
        mem_fig2 = BytesIO()
        plt.savefig(mem_fig2, format='png', dpi=150)
        plt.close()
        
        doc.add_picture(mem_fig2, width=Inches(6.0))
        doc.add_paragraph("Fig 2. Evolución del número de golpes N(20).", style='Caption')

        # --- CONTROL DE SALTO DE PÁGINA ---
        # Solo añadir salto si NO es el último ensayo de la lista
        if i < len(ensayos_unicos) - 1:
            doc.add_page_break()

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# --- 6. SIDEBAR ---
with st.sidebar:
    st.title("🏗️ Fórmula de Hiley para DPSH")
    st.markdown("---")
    
    st.header("1. Carga de Datos")
    uploaded_file = st.file_uploader("Subir Excel (DPRG)", type=["xlsx", "xls"])
    
    st.markdown("---")
    st.header("2. Cálculos")
    
    if st.session_state.errores_detectados:
        st.error("⛔ Corrige los errores indicados en las pestañas.")
        btn_calcular = False
    else:
        btn_calcular = st.button("🚀 Calcular", type="primary", use_container_width=True)

    st.markdown("Después de modificar los datos de entrada, presione el botón 'Calcular' para actualizar los resultados.")
    st.markdown("Y pulsar Crear archivos de Informe")  
    st.write("")
    st.markdown("---")
    st.header("3. Generar Informe")
    
    if st.session_state.resultados is None:
        st.info("ℹ️ Primero realiza el cálculo.")
    else:
        if st.button("📄 Crear Archivos de Informe", use_container_width=True):
            with st.spinner("Generando Word e Imágenes..."):
                ensayos_list = st.session_state.resultados["Ensayo"].unique()
                st.session_state.zip_buffer = generar_zip_en_memoria(
                    st.session_state.resultados, ensayos_list
                )
                st.session_state.word_buffer = generar_word_en_memoria(
                    st.session_state.resultados, 
                    st.session_state.config_tramos, 
                    ensayos_list
                )
            st.success("¡Archivos Generados!")
        
        if st.session_state.zip_buffer and st.session_state.word_buffer:
            c_d1, c_d2 = st.columns(2)
            with c_d1:
                st.download_button(
                    label="📥 Word",
                    data=st.session_state.word_buffer.getvalue(),
                    file_name="Informe_Hiley.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with c_d2:
                st.download_button(
                    label="📥 ZIP Datos",
                    data=st.session_state.zip_buffer.getvalue(),
                    file_name="Datos_Hiley.zip",
                    mime="application/zip",
                    use_container_width=True
                )

    st.markdown("---")
    st.caption("v3.7 - ITQ")

# --- 7. ÁREA PRINCIPAL ---
if not uploaded_file:
    st.info("👋  Carga un archivo Excel (DPRG) en el menú lateral para comenzar.")
else:
    try:
        df_dprg = pd.read_excel(uploaded_file)
        df_dprg.columns = df_dprg.columns.str.strip()
        
        col_desc = 'Descripción Muestra'
        col_depth = 'Profundidad'
        col_blows = 'Número de Golpes'
        
        if not all(col in df_dprg.columns for col in [col_desc, col_depth, col_blows]):
            st.error(f"Error: Faltan columnas ({col_desc}, ...)")
            st.stop()
            
        df_base = df_dprg[[col_desc, col_depth, col_blows]].copy()
        df_base[col_depth] = pd.to_numeric(df_base[col_depth], errors='coerce')
        df_base[col_blows] = pd.to_numeric(df_base[col_blows], errors='coerce')
        df_base = df_base.dropna().sort_values([col_desc, col_depth]).reset_index(drop=True)
        ensayos_unicos = df_base[col_desc].unique()

        st.subheader(f"📂 Proyecto Cargado: {len(ensayos_unicos)} Ensayos DPSH detectados")
        
        hay_error_grave = False
        tabs = st.tabs([f"📍 {e}" for e in ensayos_unicos])

        for i, ensayo in enumerate(ensayos_unicos):
            with tabs[i]:
                st.markdown("##### 🛠️ Configuración de Tipos de suelo")
                max_depth_real = df_base[df_base[col_desc] == ensayo][col_depth].max()
                
                # Sin redondeo (exactitud)
                df_template = pd.DataFrame({
                    "Desde (m)": [0.0],
                    "Hasta (m)": [float(max_depth_real)], 
                    "Material": ["Arenas o Gravas (F=25)"]
                })
                
                if ensayo in st.session_state.config_tramos:
                    data_to_show = st.session_state.config_tramos[ensayo]
                else:
                    data_to_show = df_template

                edited_df = st.data_editor(
                    data_to_show,
                    num_rows="dynamic",
                    column_config={
                        "Desde (m)": st.column_config.NumberColumn(format="%.2f"),
                        "Hasta (m)": st.column_config.NumberColumn(format="%.2f"),
                        "Material": st.column_config.SelectboxColumn(
                            "Tipo de Suelo",
                            options=list(MATERIALES_F.keys()),
                            required=True
                        )
                    },
                    key=f"editor_{ensayo}",
                    use_container_width=True
                )
                
                st.session_state.config_tramos[ensayo] = edited_df

                # Validaciones
                if not edited_df.empty:
                    errores_fila = edited_df[edited_df["Desde (m)"] >= edited_df["Hasta (m)"]]
                    if not errores_fila.empty:
                        st.error("⛔ **Error Lógico:** 'Desde' >= 'Hasta'.")
                        hay_error_grave = True
                    
                    max_definido = edited_df["Hasta (m)"].max()
                    # Integridad (margen error 5cm)
                    if max_definido < (max_depth_real - 0.05):
                        metros_faltantes = max_depth_real - max_definido
                        st.warning(f"⚠️ **Incompleto:** El ensayo llega a **{max_depth_real:.2f}m**.")
                        st.error(f"❌ **Faltan definir: {metros_faltantes:.2f} metros**")

                # Resultados Visuales
                if st.session_state.resultados is not None:
                    df_res = st.session_state.resultados
                    df_ensayo = df_res[df_res["Ensayo"] == str(ensayo)]
                    
                    if not df_ensayo.empty:
                        st.divider()
                        st.markdown("##### 📊 Resultados")
                        # 3 Decimales variables intermedias
                        st.dataframe(
                            df_ensayo.style.format({
                                "Z(m)": "{:.2f}", 
                                "N(20)": "{:.0f}", 
                                "F": "{:.0f}",
                                "a": "{:.3f}", 
                                "e": "{:.3f}", 
                                "n": "{:.3f}", 
                                "c": "{:.3f}",
                                "Presión Admisible (kg/cm2)": "{:.2f}",
                                "Presión Característica (kg/cm2)": "{:.2f}"
                            }),
                            use_container_width=True,
                            height=300,
                            hide_index=True 
                        )
                        c1, c2 = st.columns(2)
                        with c1:
                            fig1, ax1 = plt.subplots(figsize=(4, 3))
                            ax1.plot(df_ensayo['Presión Admisible (kg/cm2)'], df_ensayo['Z(m)'], marker='.')
                            ax1.invert_yaxis()
                            ax1.grid(True, alpha=0.3)
                            ax1.set_xlabel('Presión ($kg/cm^2$)')
                            ax1.set_ylabel('Profundidad (m)')
                            ax1.set_title("Presión Admisible")
                            st.pyplot(fig1)
                        with c2:
                            fig2, ax2 = plt.subplots(figsize=(4, 3))
                            ax2.plot(df_ensayo['N(20)'], df_ensayo['Z(m)'], marker='.', color='brown')
                            ax2.invert_yaxis()
                            ax2.grid(True, alpha=0.3)
                            ax2.set_xlabel('N(20) (Golpes)')
                            ax2.set_ylabel('Profundidad (m)')
                            ax2.set_title("N(20)")
                            st.pyplot(fig2)

        st.session_state.errores_detectados = hay_error_grave

        # --- LÓGICA DE CÁLCULO ---
        if btn_calcular and not hay_error_grave:
            with st.spinner("Calculando..."):
                st.session_state.zip_buffer = None
                st.session_state.word_buffer = None
                
                final_rows = []
                for idx, row in df_base.iterrows():
                    ensayo = row[col_desc]
                    depth = row[col_depth]
                    n20 = row[col_blows]
                    
                    factor_f = 50.0 
                    if ensayo in st.session_state.config_tramos:
                        tramos_df = st.session_state.config_tramos[ensayo]
                        for _, tramo in tramos_df.iterrows():
                            try:
                                if float(tramo["Desde (m)"]) <= depth <= float(tramo["Hasta (m)"]):
                                    factor_f = MATERIALES_F.get(tramo["Material"], 50.0)
                                    break
                            except: pass
                    
                    # --- NUEVA LÓGICA N=0 (SUELOS BLANDOS) ---
                    if n20 > 0:
                        # Fórmulas Hiley Normales
                        a = depth / 10.0 + 0.25
                        e = 20.0 / n20
                        n_val = 0.7 - 0.7 / 19.0 * (e - 1.0)
                        c_val = 0.5 - 0.5 / 19.0 * (e - 1.0)
                        term1 = 63.5 * 76.0 * (1.0 + (n_val ** 2) * a)
                        term2 = (e + c_val) * (1.0 + a) * 20.0
                        pres_car = term1 / term2
                        pres_adm = pres_car / factor_f
                    else:
                        # Caso N=0: No hay resistencia
                        n20 = 0.0 # Aseguramos que se muestre 0
                        a = depth / 10.0 + 0.25 # Depende de Z, se mantiene
                        e = 0.0
                        n_val = 0.0
                        c_val = 0.0
                        pres_car = 0.0
                        pres_adm = 0.0
                    
                    final_rows.append({
                        "Ensayo": str(ensayo),
                        "Z(m)": depth, "N(20)": n20, "F": factor_f,
                        "a": a, "e": e, "n": n_val, "c": c_val,
                        "Presión Característica (kg/cm2)": pres_car,
                        "Presión Admisible (kg/cm2)": pres_adm
                    })
                
                st.session_state.resultados = pd.DataFrame(final_rows)
                st.toast("✅ Cálculos actualizados", icon="⚡")
                st.rerun()

    except Exception as e:
        st.error(f"Error crítico: {e}")