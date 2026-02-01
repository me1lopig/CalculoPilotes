
# Programa para el cálculo de valores de Adherencia Límite para Micropilotes
# Según la Guía para el proyecto y la ejecución de Micropilotes
# en Obras de Carreteras
# Dirección General de Carreteras
# Ministerio de Fomento
# Gobierno de España

# Autor: Germán López Pineda
# ICCP Universidad de Granada
# MMC Universitat Jaume I
# Fecha: 2024-11-19         
# Versión: 1.0
# Licencia: GPL-3.0
# Contacto: me1lopig@uco.es


import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Cálculo Geotécnico Micropilotes", layout="wide")

st.title("🧮 Calculadora de valores de Adherencia Límite para  Micropilotes")
st.markdown("Cálculo basado en la **Guía para el proyecto y ejecución de micropilotes en obras de carreteras**.")
st.markdown("Desarrollado por la DT del Departamento de Geotecnia ")
st.markdown("Version febrero de 2026")
# --- CONSTANTES Y REGLAS DE CONTROL ---
PARAM_ARENAS = ["Plim (MPa)", "SPT (N)"]
PARAM_ARCILLAS = ["Plim (MPa)", "qu (MPa)"]
OPCIONES_TODAS = ["Plim (MPa)", "SPT (N)", "qu (MPa)"]

LIMITES = {
    "Arenas y Gravas": {
        "Plim (MPa)": {"min": 0.5, "max": 7.0},
        "SPT (N)": {"min": 10, "max": 100}
    },
    "Arcillas y Limos": {
        "Plim (MPa)": {"min": 0.25, "max": 2.5},
        "qu (MPa)": {"min": 0.05, "max": 0.5}
    }
}

# --- FUNCIONES DE CÁLCULO (AJUSTE POR 3 TRAMOS RECTOS) ---

def calcular_arenas(plim):
    # IRS: Satura en 0.62 @ 2.8 MPa
    x_irs = [0.0, 0.5, 2.8, 10.0]; y_irs = [0.0, 0.14, 0.62, 0.62]
    # IR: Satura en 0.50 @ 3.3 MPa
    x_ir  = [0.0, 0.5, 3.3, 10.0]; y_ir  = [0.0, 0.09, 0.50, 0.50]
    # IU: Satura en 0.40 @ 3.5 MPa
    x_iu  = [0.0, 0.5, 3.5, 10.0]; y_iu  = [0.0, 0.04, 0.40, 0.40]
    return np.interp(plim, x_iu, y_iu), np.interp(plim, x_ir, y_ir), np.interp(plim, x_irs, y_irs)

def calcular_arcillas(plim):
    # Ajuste exacto 3 tramos: 1 (0-0.5), 2 (0.5-Sat), 3 (Meseta)
    x_irs = [0.0, 0.25, 0.5, 1.8, 10.0]; y_irs = [0.0, 0.12, 0.20, 0.40, 0.40]
    x_ir  = [0.0, 0.25, 0.5, 2.1, 10.0]; y_ir  = [0.0, 0.07, 0.12, 0.30, 0.30]
    x_iu  = [0.0, 0.25, 0.5, 2.3, 10.0]; y_iu  = [0.0, 0.05, 0.08, 0.20, 0.20]
    return np.interp(plim, x_iu, y_iu), np.interp(plim, x_ir, y_ir), np.interp(plim, x_irs, y_irs)

# --- TABLA DE ENTRADA CON DESPLEGABLES ---
st.subheader("📝 1. Entrada de Datos")

if 'df_data' not in st.session_state:
    st.session_state.df_data = pd.DataFrame([
        {"Nombre": "UG-01", "Suelo": "Arenas y Gravas", "Ensayo": "SPT (N)", "Valor": 45.0},
        {"Nombre": "UG-02", "Suelo": "Arcillas y Limos", "Ensayo": "qu (MPa)", "Valor": 0.25}
    ])

df_input = st.data_editor(
    st.session_state.df_data,
    num_rows="dynamic",
    column_config={
        "Suelo": st.column_config.SelectboxColumn("Tipo de Suelo", options=list(LIMITES.keys()), required=True),
        "Ensayo": st.column_config.SelectboxColumn("Parámetro/Ensayo", options=OPCIONES_TODAS, required=True),
        "Valor": st.column_config.NumberColumn("Valor Numérico", format="%.2f", min_value=0.0)
    },
    use_container_width=True
)

# --- PROCESAMIENTO Y CONTROLES ---
results = []
valid_inputs_report = []

for i, row in df_input.iterrows():
    if pd.isna(row["Valor"]) or not row["Nombre"]: continue
    
    s_tipo = row["Suelo"]
    e_tipo = row["Ensayo"]
    val = row["Valor"]
    
    # CONTROL 1: Coherencia de Suelo vs Parámetro
    if s_tipo == "Arenas y Gravas" and e_tipo not in PARAM_ARENAS:
        st.error(f"❌ Error en '{row['Nombre']}': Para Arenas usar Plim o SPT.")
        continue
    if s_tipo == "Arcillas y Limos" and e_tipo not in PARAM_ARCILLAS:
        st.error(f"❌ Error en '{row['Nombre']}': Para Arcillas usar Plim o qu.")
        continue

    # CONTROL 2: Límites Normativos
    lims = LIMITES[s_tipo][e_tipo]
    if val < lims["min"] or val > lims["max"]:
        st.error(f"⚠️ Fila '{row['Nombre']}': Valor fuera de rango ({lims['min']} a {lims['max']}).")
        continue

    # Cálculo
    plim = val if "Plim" in e_tipo else (val / 20.0 if "SPT" in e_tipo else val * 5.0)
    iu, ir, irs = calcular_arenas(plim) if s_tipo == "Arenas y Gravas" else calcular_arcillas(plim)
    
    # Almacenar para informe (entrada)
    valid_inputs_report.append({"Nombre": row["Nombre"], "Suelo": s_tipo, "Ensayo": e_tipo, "Valor": val})
    # Almacenar para resultados (salida)
    results.append({
        "Nombre": row["Nombre"], "IU (MPa)": round(iu, 3), 
        "IR (MPa)": round(ir, 3), "IRS (MPa)": round(irs, 3), "Grupo": s_tipo, "_p": plim
    })

df_res = pd.DataFrame(results)
df_in_rep = pd.DataFrame(valid_inputs_report)

# --- GRÁFICOS Y DESCARGA ---

def dibujar_grafico(grupo, df_datos):
    is_arenas = (grupo == "Arenas y Gravas")
    fig, ax1 = plt.subplots(figsize=(8, 6))
    x_max, p_min = (7.0, 0.5) if is_arenas else (2.5, 0.25)
    x_norm = np.linspace(p_min, x_max, 400)
    y_norm = calcular_arenas(x_norm) if is_arenas else calcular_arcillas(x_norm)
    
    ax1.plot(x_norm, y_norm[2], 'k-', label='IRS', linewidth=2)
    ax1.plot(x_norm, y_norm[1], 'k-.', label='IR')
    ax1.plot(x_norm, y_norm[0], 'k--', label='IU')

    if not df_datos.empty:
        colores = plt.cm.tab10(np.linspace(0, 1, len(df_datos)))
        for i, (_, s) in enumerate(df_datos.iterrows()):
            ax1.scatter([s["_p"]]*3, [s["IU (MPa)"], s["IR (MPa)"], s["IRS (MPa)"]], 
                        color=colores[i], label=s["Nombre"], s=60, edgecolors='white', zorder=5)
    
    ax1.set_title(f"Tipo de Suelo: {grupo}")
    ax1.set_xlabel("Plim (MPa)")
    ax1.set_ylabel("rf,lim (MPa)")
    ax1.set_xlim(0, x_max); ax1.set_ylim(0, 0.8 if is_arenas else 0.5)
    ax1.grid(True, alpha=0.3); ax1.legend(fontsize='x-small')
    
    # Eje Dual
    ax2 = ax1.twiny()
    ax2.set_xlim(ax1.get_xlim()); ax2.xaxis.set_ticks_position('bottom'); ax2.xaxis.set_label_position('bottom')
    ax2.spines['bottom'].set_position(('outward', 40))
    ticks = ax1.get_xticks()
    ax2.set_xticks(ticks)
    if is_arenas:
        ax2.set_xticklabels([f"{int(t*20)}" for t in ticks]); ax2.set_xlabel("N (SPT)")
    else:
        ax2.set_xticklabels([f"{t/5:.2f}" for t in ticks]); ax2.set_xlabel("qu (MPa)")
    return fig

# Informe Word
def generar_word(df_in, df_out, f1, f2):
    doc = Document()
    doc.add_heading('Informe de resultados de Adherencia Límite', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for titulo, df in [("1. Datos de Entrada", df_in), ("2. Resultados de Cálculo", df_out)]:
        doc.add_heading(titulo, level=1)
        t = doc.add_table(df.shape[0] + 1, df.shape[1])
        t.style = 'Table Grid'
        for i, col in enumerate(df.columns): t.cell(0, i).text = col
        for i, row in enumerate(df.values):
            for j, val in enumerate(row): t.cell(i+1, j).text = str(val)

    doc.add_heading('3. Gráficos de Cálculo', level=1)
    for n, f in [("Arenas", f1), ("Arcillas", f2)]:
        buf = BytesIO(); f.savefig(buf, format='png', dpi=200, bbox_inches='tight')
        doc.add_picture(buf, width=Inches(5.0))
    
    out = BytesIO(); doc.save(out); out.seek(0)
    return out

# Mostrar resultados finales
if not df_res.empty:
    st.subheader("📋 2. Resultados de cálculo de adherencia límite")
    st.dataframe(df_res.drop(columns=["_p", "Grupo"]), use_container_width=True)
    
    fig1 = dibujar_grafico("Arenas y Gravas", df_res[df_res["Grupo"] == "Arenas y Gravas"])
    fig2 = dibujar_grafico("Arcillas y Limos", df_res[df_res["Grupo"] == "Arcillas y Limos"])
    
    # BOTÓN ÚNICO
    informe = generar_word(df_in_rep, df_res.drop(columns=["_p", "Grupo"]), fig1, fig2)
    st.download_button("📥 Descargar Informe  (.docx)", informe, "Informe_Micropilotes.docx", type="primary")
else:
    fig1 = dibujar_grafico("Arenas y Gravas", pd.DataFrame())
    fig2 = dibujar_grafico("Arcillas y Limos", pd.DataFrame())

st.divider()
c1, c2 = st.columns(2)
with c1: st.pyplot(fig1)
with c2: st.pyplot(fig2)