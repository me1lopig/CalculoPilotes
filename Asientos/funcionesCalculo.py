# =============================================================================
# LIBRERÍA: funcionesCalculo.py
# Propósito: Cálculo de tensiones y asientos para ZAPATA CONTINUA (2D)
# Versión: FINAL GRAN FORMATO (Imágenes Verticales de 14cm)
# =============================================================================

import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
import openpyxl
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime

# --- GESTIÓN DE ARCHIVOS Y DATOS ---

def crea_directorio():
    now = datetime.now()
    directorio = "Informe_Geo_" + str(now.strftime("%Y-%m-%d_%H%M"))
    if not os.path.exists(directorio):
        os.mkdir(directorio)
    return directorio

def datos_carga():
    try:
        libro = openpyxl.load_workbook('datos_rectangular.xlsx') 
    except:
        libro = openpyxl.load_workbook('datos_carga.xlsx')
        
    hoja = libro.active 
    B = hoja.cell(row=2, column=1).value  
    q = hoja.cell(row=2, column=2).value  
    ax = hoja.cell(row=2, column=3).value 
    incrx = hoja.cell(row=2, column=4).value 
    incrz = hoja.cell(row=2, column=5).value 
    return B, q, ax, incrx, incrz

def datos_terreno():
    libro = openpyxl.load_workbook('datos_terreno.xlsx')
    hoja = libro.active
    
    espesor=[0]; cotas=[0]; pe_seco=[0]; pe_saturado=[0]
    E=[0]; poisson=[0]; cohesion=[0]; fi=[0]; tipo_datos=[]

    for row in hoja.iter_rows(min_row=2, values_only=True):
        if row[0] is None: break 
        espesor.append(row[0])
        pe_seco.append(row[2])
        pe_saturado.append(row[3])
        E.append(row[4])
        poisson.append(row[5])
        cohesion.append(row[6])
        fi.append(row[7])

    for i in range(1, len(espesor)):
        cotas.append(cotas[i-1] + espesor[i])
        
    az = cotas[-1] 
    nivel_freatico = hoja.cell(row=2, column=2).value 

    for col in hoja.iter_cols(max_row=1, values_only=True):
        tipo_datos.append(col[0])

    return espesor, cotas, az, nivel_freatico, pe_seco, pe_saturado, E, poisson, cohesion, fi, tipo_datos

# --- MOTORES DE CÁLCULO ---

def tension_zapata_continua(b, q, x, z):
    if z <= 0: z = 0.001 
    theta1 = np.arctan2((x - b), z)
    theta2 = np.arctan2((x + b), z)
    alpha = theta2 - theta1 
    delta = theta1          
    
    term_trig = np.sin(alpha) * np.cos(alpha + 2*delta)
    tensionz = (q / np.pi) * (alpha + term_trig)
    tensionx = (q / np.pi) * (alpha - term_trig)
    tensionxz = (q / np.pi) * (np.sin(alpha) * np.sin(alpha + 2*delta))
    return tensionz, tensionx, tensionxz

def asiento_deformacion_plana(cotas, z, hi, E, poisson, tensionx, tensionz):
    idx = parametro_terreno(cotas, z)
    nu = poisson[idx]
    mod_E = E[idx]
    term_comun = (1 + nu) / mod_E
    epsilon_z = term_comun * ((1 - nu) * tensionz - nu * tensionx)
    return abs(epsilon_z * hi)

def parametro_terreno(cotas, zt):
    for i in range(len(cotas)-1):
        if cotas[i] <= zt <= cotas[i+1]:
            return i + 1
    return len(cotas) - 1 

def tension_geostatica(z, cotas, pe_saturado, pe_seco, nf):
    sigma_v = 0
    z_actual = 0
    dz = 0.1
    pasos = int(z / dz)
    for _ in range(pasos):
        z_centro = z_actual + dz/2
        idx = parametro_terreno(cotas, z_centro)
        gamma = pe_saturado[idx] if z_centro > nf else pe_seco[idx]
        sigma_v += gamma * dz
        z_actual += dz
    return sigma_v

# --- EXPORTACIÓN Y GRÁFICOS ---

def guardar_xlsx_matriz(xcoord, zcoord, data, directorio, nombre):
    wb = openpyxl.Workbook()
    ws = wb.active
    for j, x in enumerate(xcoord):
        ws.cell(row=1, column=j+2, value=x)
    for i, z in enumerate(zcoord):
        ws.cell(row=i+2, column=1, value=z)
        for j, val in enumerate(data[i]):
            ws.cell(row=i+2, column=j+2, value=val)
    wb.save(os.path.join(directorio, f"{nombre}.xlsx"))

def guardar_xlsx_vector(xcoord, data, directorio, nombre):
    wb = openpyxl.Workbook()
    ws = wb.active
    for j, x in enumerate(xcoord):
        ws.cell(row=1, column=j+1, value=x)
        ws.cell(row=2, column=j+1, value=data[j])
    wb.save(os.path.join(directorio, f"{nombre}.xlsx"))

def graficos_tensiones_zapata(xcoord, zcoord, tension, directorio, titulo, tipo, B, cotas, nf):
    plt.rcParams['font.family'] = 'sans-serif'
    # Mantenemos el lienzo fijo y márgenes controlados
    fig, ax = plt.subplots(figsize=(10, 10))
    fig.subplots_adjust(left=0.12, right=0.88, top=0.90, bottom=0.10)
    
    X, Z = np.meshgrid(xcoord, zcoord)
    
    # Estilos
    if tipo == 'continua':
        color_texto = 'white'
        color_linea_unidad = 'white'
        estilo_linea_unidad = '--'
        opacidad_estrato = 0.1
        peso_fuente = 'bold'
    else:
        color_texto = '#424242'
        color_linea_unidad = '#616161'
        estilo_linea_unidad = '-'
        opacidad_estrato = 0.3
        peso_fuente = 'normal'

    # Estratos
    colores_estratos = ['#FFFFFF', '#F9FBE7', '#F3E5F5', '#E3F2FD', '#E8F5E9']
    for i in range(len(cotas)-1):
        c_sup, c_inf = cotas[i], cotas[i+1]
        if c_inf > max(zcoord): c_inf = max(zcoord)
        ax.axhspan(c_sup, c_inf, color=colores_estratos[i % len(colores_estratos)], alpha=opacidad_estrato, zorder=0)
        ax.axhline(c_inf, color=color_linea_unidad, linestyle=estilo_linea_unidad, linewidth=1.0, zorder=20)
        ax.text(max(xcoord)*0.98, (c_sup + c_inf)/2, f"U.G. {i+1}", 
                color=color_texto, fontsize=9, ha='right', fontweight=peso_fuente, style='italic', zorder=20)

    # Nivel Freático
    if 0 < nf < max(zcoord):
        ax.axhline(nf, color='#0277BD', linestyle='-.', linewidth=2.2, zorder=25)
        for side in [min(xcoord)*0.92, max(xcoord)*0.92]:
            ax.plot(side, nf, marker='v', color='#0277BD', markersize=12, markeredgecolor='black', zorder=26)

    # Zapata
    canto_zapata = 0.6 
    b = B / 2
    coords = [(-b, -canto_zapata), (b, -canto_zapata), (b, 0), (-b, 0)]
    ax.add_patch(Polygon(coords, facecolor='#EEEEEE', edgecolor='black', hatch='///', linewidth=1.5, zorder=30))

    # Mapas
    if tipo == 'isolinea':
        CS = ax.contour(X, Z, tension, levels=15, colors='#0D47A1', linewidths=1.3, zorder=10)
        ax.clabel(CS, inline=True, fontsize=9, fmt='%1.1f', colors='black')
    else:
        cf = ax.contourf(X, Z, tension, levels=60, cmap='jet', alpha=0.9, extend='both', zorder=5)
        cbar = fig.colorbar(cf, ax=ax, shrink=0.7)
        cbar.set_label('Incremento Tensión [kN/m²]', fontweight='bold')

    ax.invert_yaxis()
    ax.set_ylim(max(zcoord), -canto_zapata * 2.5) 
    ax.set_xlim(min(xcoord), max(xcoord))
    ax.set_aspect('equal', adjustable='box')
    
    ax.set_title(f"{titulo} ({tipo.upper()})", fontsize=13, fontweight='bold', pad=15)
    ax.set_xlabel("Distancia X [m]", fontsize=11, fontweight='bold')
    ax.set_ylabel("Profundidad Z [m]", fontsize=11, fontweight='bold')
    ax.axhline(0, color='black', linewidth=2.0, zorder=31)

    plt.savefig(os.path.join(directorio, f"{titulo.replace(' ', '_')}_{tipo}.png"), dpi=300)
    plt.close()

def grafico_asientos(xcoord, asiento, directorio, titulo, cotas, nf, B):
    plt.rcParams['font.family'] = 'sans-serif'
    fig, ax = plt.subplots(figsize=(12, 5))
    fig.subplots_adjust(left=0.1, right=0.95, top=0.88, bottom=0.15)
    
    asiento_mm = np.array(asiento) * 1000 
    max_asiento = np.max(asiento_mm) if np.max(asiento_mm) > 0 else 1
    b = B / 2.0

    ax.axvline(-b, color='#555555', linestyle='--', linewidth=1, alpha=0.8)
    ax.axvline(b, color='#555555', linestyle='--', linewidth=1, alpha=0.8)
    
    y_cota = -max_asiento * 0.2 
    ax.annotate('', xy=(-b, y_cota), xytext=(b, y_cota),
                arrowprops=dict(arrowstyle='<->', color='black', lw=1.2))
    ax.text(0, y_cota, f" B = {B} m ", ha='center', va='center', 
            fontsize=10, backgroundcolor='white')

    ax.axhline(0, color='black', linewidth=1.5)
    ax.plot(xcoord, asiento_mm, color='#0D47A1', linewidth=2.0, label='Asiento', zorder=5)
    ax.fill_between(xcoord, 0, asiento_mm, color='#E3F2FD', alpha=0.4)

    puntos_interes = [0, -b, b]
    for p_x in puntos_interes:
        idx = (np.abs(xcoord - p_x)).argmin()
        val_y = asiento_mm[idx]
        ax.plot(xcoord[idx], val_y, 'o', color='white', markeredgecolor='#0D47A1', markersize=6, zorder=6)
        ax.text(xcoord[idx], val_y + (max_asiento*0.15), f"{val_y:.1f}mm", 
                ha='center', va='top', fontsize=8, color='#0D47A1', fontweight='bold')

    ax.invert_yaxis() 
    ax.set_ylim(max_asiento * 2.0, -max_asiento * 0.5) 
    ax.set_xlim(-1.5 * B, 1.5 * B) 

    ax.set_title(titulo, fontsize=12, fontweight='bold', pad=15)
    ax.set_ylabel("Asiento [mm]", fontsize=11, fontweight='bold')
    ax.set_xlabel("Distancia al eje [m]", fontsize=11, fontweight='bold')
    
    ax.grid(True, which='major', linestyle=':', color='gray', alpha=0.3)
    ax.minorticks_on()

    plt.savefig(os.path.join(directorio, f"{titulo}.png"), dpi=300)
    plt.close()

def guardar_reporte_docx(B, q, ax, az, directorio, espesor, pe_seco, pe_saturado, E, poisson, nf):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # TITULO
    titulo = doc.add_heading('Informe de Cálculo Geotécnico: Zapata Continua', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("_" * 70)

    # 1. DATOS
    doc.add_heading('1. Datos de Diseño', level=1)
    
    table_geo = doc.add_table(rows=1, cols=2)
    table_geo.style = 'Light Shading Accent 1'
    table_geo.alignment = WD_TABLE_ALIGNMENT.CENTER
    datos = [("Ancho (B)", f"{B} m"), ("Carga (q)", f"{q} kN/m²"), 
             ("Profundidad (z)", f"{az} m"), ("Nivel Freático", f"{nf} m")]
    table_geo.rows[0].cells[0].text = 'Parámetro'
    table_geo.rows[0].cells[1].text = 'Valor'
    for d, v in datos:
        row = table_geo.add_row().cells
        row[0].text = d; row[1].text = v
    doc.add_paragraph() 

    # 2. ESTRATIGRAFÍA
    doc.add_heading('2. Estratigrafía', level=1)
    headers = ['Capa', 'Espesor', 'Gamma Sat', 'E (kPa)', 'Nu']
    table_terr = doc.add_table(rows=1, cols=len(headers))
    table_terr.style = 'Light List Accent 1'
    table_terr.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table_terr.rows[0].cells[i].text = h
    for i in range(1, len(espesor)):
        row = table_terr.add_row().cells
        row[0].text = f"U.G. {i}"
        row[1].text = str(espesor[i]); row[2].text = str(pe_saturado[i])
        row[3].text = str(E[i]); row[4].text = str(poisson[i])
    doc.add_page_break()

    # 3. RESULTADOS GRÁFICOS
    doc.add_heading('3. Resultados Gráficos', level=1)
    
    # A. ASIENTOS
    doc.add_heading('3.1. Asientos', level=2)
    path_asientos = os.path.join(directorio, 'Perfil de Asientos.png')
    if os.path.exists(path_asientos):
        doc.add_picture(path_asientos, width=Cm(16))
        c = doc.add_paragraph('Fig 1. Perfil de asientos (Técnico)')
        c.style = 'Caption'; c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # B. TENSIONES (Verticales Grandes 14cm)
    def agregar_pareja_vertical(titulo_seccion, archivo_iso, archivo_cont, fig_num):
        doc.add_heading(titulo_seccion, level=2)
        
        # Imagen 1: Isolíneas
        path_iso = os.path.join(directorio, archivo_iso)
        if os.path.exists(path_iso):
            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p1.add_run()
            # Ancho grande para que se vea bien (14 cm)
            run1.add_picture(path_iso, width=Cm(14))
            
            c1 = doc.add_paragraph('(a) Isolíneas')
            c1.style = 'Caption'; c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph() # Espacio entre imágenes

        # Imagen 2: Continua
        path_cont = os.path.join(directorio, archivo_cont)
        if os.path.exists(path_cont):
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run()
            # Mismo ancho grande
            run2.add_picture(path_cont, width=Cm(14))
            
            c2 = doc.add_paragraph('(b) Mapa Continuo')
            c2.style = 'Caption'; c2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Título General de la Figura
        c_main = doc.add_paragraph(f'Fig {fig_num}. Comparativa de {titulo_seccion}')
        c_main.style = 'Caption'; c_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_main.runs[0].font.bold = True
        
        doc.add_page_break()

    agregar_pareja_vertical('3.2. Tensión Vertical', 'Tensión_Vertical_Sigma_Z_isolinea.png', 'Tensión_Vertical_Sigma_Z_continua.png', 2)
    agregar_pareja_vertical('3.3. Tensión Horizontal', 'Tensión_Horizontal_Sigma_X_isolinea.png', 'Tensión_Horizontal_Sigma_X_continua.png', 3)
    agregar_pareja_vertical('3.4. Tensión Cortante', 'Tensión_Cortante_Tau_XZ_isolinea.png', 'Tensión_Cortante_Tau_XZ_continua.png', 4)

    nombre_informe = os.path.join(directorio, "Informe_Geotecnico_Final.docx")
    doc.save(nombre_informe)
    print(f"\n[INFO] Informe generado exitosamente: {nombre_informe}")