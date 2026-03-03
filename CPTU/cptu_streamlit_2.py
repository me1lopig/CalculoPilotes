import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import re
import io
import zipfile
from docx import Document
from docx.shared import Inches

# --- CONFIGURACIÓN INICIAL ---
st.set_page_config(page_title="Visor CPTU Profesional", layout="wide")

st.title("📊 Análisis Geotécnico Avanzado de Ensayo CPTU")
st.markdown("Generador de informes técnicos y parámetros de diseño **Robertson (2010/2012)**.")
st.divider()

# --- FUNCIONES GEOTÉCNICAS ---
@st.cache_data 
def calcular_geotecnia(df):
    pa = 0.1
    gamma_w = 9.81
    df_c = df[(df['Qc'] > 0.01) & (df['Rf'] > 0.01)].copy()
    df_c['Qc_pa'] = df_c['Qc'] / pa
    df_c['Ic'] = np.sqrt((3.47 - np.log10(df_c['Qc_pa']))**2 + (np.log10(df_c['Rf']) + 1.22)**2)
    condiciones = [(df_c['Ic'] < 1.31), (df_c['Ic'] >= 1.31) & (df_c['Ic'] < 2.05), (df_c['Ic'] >= 2.05) & (df_c['Ic'] < 2.60), (df_c['Ic'] >= 2.60) & (df_c['Ic'] < 2.95), (df_c['Ic'] >= 2.95) & (df_c['Ic'] < 3.60), (df_c['Ic'] >= 3.60)]
    zonas_sbt = ['7. Grava/Arena', '6. Arena a arena limosa', '5. Arena limosa', '4. Limo arcilloso', '3. Arcilla limosa', '2. Orgánico/Turba']
    zonas_num = [7, 6, 5, 4, 3, 2]
    df_c['SBT_Name'] = np.select(condiciones, zonas_sbt, default='Desconocido')
    df_c['SBT_Zone'] = np.select(condiciones, zonas_num, default=0)
    df_c['Gamma_kN3'] = (0.27 * np.log10(df_c['Rf']) + 0.36 * np.log10(df_c['Qc_pa']) + 1.236) * gamma_w
    df_c['Gamma_kN3'] = np.clip(df_c['Gamma_kN3'], 12, 22)
    df_c['dz'] = df_c['Depth_m'].diff().fillna(0)
    df_c['sigma_v0_kPa'] = (df_c['Gamma_kN3'] * df_c['dz']).cumsum()
    Nkt = 14 
    df_c['Su_kPa'] = np.where(df_c['Ic'] > 2.6, ((df_c['Qc'] * 1000) - df_c['sigma_v0_kPa']) / Nkt, np.nan)
    df_c['Phi_deg'] = np.where(df_c['Ic'] <= 2.6, 17.6 + 11.0 * np.log10(df_c['Qc_pa']), np.nan)
    df_c['N60'] = df_c['Qc_pa'] / (10 ** (1.1268 - 0.2817 * df_c['Ic']))
    return df_c

SBT_COLORS = {7: '#d8b365', 6: '#f5f5dc', 5: '#c7eae5', 4: '#80cdc1', 3: '#dfc27d', 2: '#8c510a'}
LEYENDA_SUELOS = [plt.Rectangle((0,0),1,1, color=SBT_COLORS[k], label=v) for k, v in zip([7,6,5,4,3,2], ['7. Grava/Arena', '6. Arena limosa', '5. Arena limosa fina', '4. Limo arcilloso', '3. Arcilla limosa', '2. Orgánico/Turba'])]

def plot_stratigraphy_col(ax, df, preforo):
    if preforo > 0:
        ax.axhspan(0, preforo, color='gray', alpha=0.3, hatch='//')
    for zone in df['SBT_Zone'].unique():
        if zone == 0: continue
        ax.fill_betweenx(df['Depth_m'], 0, 1, where=df['SBT_Zone']==zone, color=SBT_COLORS.get(zone, '#fff'), step='mid')
    ax.set_xlim(0, 1); ax.set_xticks([]); ax.set_ylabel('Profundidad (m)', fontsize=11, fontweight='bold')
    ax.invert_yaxis(); ax.set_title('Estratigrafía', fontsize=10, fontweight='bold')

# --- INTERFAZ ---
uploaded_file = st.file_uploader("📂 Sube el archivo CPTU (.CSV)", type=["csv", "CSV"])

if uploaded_file is not None:
    # 1. METADATOS
    content = uploaded_file.read().decode('utf-8').splitlines()
    header_data = {}
    comentario_preforo = 0.0 
    for line in content[:20]:
        if ';' in line:
            key, val = line.split(';', 1)
            clean_key = key.strip().rstrip(':'); clean_val = val.strip().strip(';')
            if clean_key and clean_val:
                header_data[clean_key] = clean_val
                if clean_key.lower() == 'comments':
                    match = re.search(r'(\d+(?:[.,]\d+)?)', clean_val)
                    if match: comentario_preforo = float(match.group(1).replace(',', '.'))
    
    with st.expander("📋 Ver Datos de la Campaña (Cabecera)", expanded=False):
        items = list(header_data.items()); mitad = len(items)//2 + len(items)%2
        c1, c2 = st.columns(2)
        c1.table(pd.DataFrame(items[:mitad], columns=["Parámetro", "Valor"]).set_index("Parámetro"))
        c2.table(pd.DataFrame(items[mitad:], columns=["Parámetro", "Valor"]).set_index("Parámetro"))
                
    uploaded_file.seek(0)
    df = pd.read_csv(uploaded_file, sep=';', decimal=',', skiprows=23)
    df['Depth_m'] = df['Depth'] / 100.0 
    df_calc = calcular_geotecnia(df)
    
    datos_validos = df[df['Qc'] > 0.05]
    cota_analitica = datos_validos['Depth_m'].min() if not datos_validos.empty else 0.0
    st.sidebar.header("⚙️ Configuración")
    cota_preforo = st.sidebar.number_input("Preforo (m)", 0.0, float(df['Depth_m'].max()), float(max(cota_analitica, comentario_preforo)))
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📉 Básicos", "🏗️ Diseño", "🕵️ Calidad", "📑 Capas", "📋 Detalles"])
    leyenda_preforo = [plt.Rectangle((0,0),1,1, color='gray', alpha=0.3, hatch='//', label='PREFORO')] + LEYENDA_SUELOS if cota_preforo > 0 else LEYENDA_SUELOS

    # --- RENDERIZADO DE GRÁFICAS CORREGIDO ---
    def generar_figura_estandar(tipo='basicos'):
        fig, axs = plt.subplots(1, 6, figsize=(16, 8), sharey=True, gridspec_kw={'width_ratios': [1, 2, 2, 2, 2, 1.5]})
        plot_stratigraphy_col(axs[0], df_calc, cota_preforo)
        
        if tipo == 'basicos':
            axs[1].plot(df_calc['Qc'], df_calc['Depth_m'], '#1f77b4'); axs[1].set_title('Qc (MPa)')
            axs[2].plot(df_calc['Fs'], df_calc['Depth_m'], '#ff7f0e'); axs[2].set_title('Fs (kPa)')
            axs[3].plot(df_calc['U2'], df_calc['Depth_m'], '#d62728'); axs[3].set_title('U2 (kPa)')
            axs[4].plot(df_calc['Ic'], df_calc['Depth_m'], '#2ca02c'); axs[4].axvline(2.6, color='red', ls='--'); axs[4].set_title('SBT Index (Ic)')
            for ax in axs[1:5]: ax.grid(True, ls='--', alpha=0.5)
        elif tipo == 'diseno':
            axs[1].plot(df_calc['Su_kPa'], df_calc['Depth_m'], 'brown'); axs[1].set_title('Su (kPa)')
            axs[2].plot(df_calc['Phi_deg'], df_calc['Depth_m'], 'orange'); axs[2].set_title('Phi (°)')
            axs[3].plot(df_calc['Gamma_kN3'], df_calc['Depth_m'], 'purple'); axs[3].set_title('Gamma (kN/m³)')
            axs[4].plot(df_calc['N60'], df_calc['Depth_m'], 'black'); axs[4].set_title('SPT N60 Eq.')
            for ax in axs[1:5]: ax.grid(True, ls='--', alpha=0.5)
        elif tipo == 'calidad':
            axs[1].plot(df_calc['Tilt'], df_calc['Depth_m'], 'red'); axs[1].axvline(15, color='k', ls='--'); axs[1].set_title('Inclinación (°)')
            axs[2].plot(df_calc['Speed'], df_calc['Depth_m'], 'teal'); axs[2].axvline(2.0, color='g', alpha=0.3, lw=4); axs[2].set_title('Velocidad (cm/s)')
            # Ocultamos los paneles que sobran para mantener la simetría
            axs[3].axis('off'); axs[4].axis('off')
            axs[1].grid(True, ls='--', alpha=0.5); axs[2].grid(True, ls='--', alpha=0.5)

        if cota_preforo > 0:
            for i in range(1, 5):
                if not axs[i].axison: continue # No sombrear si el eje está apagado
                axs[i].axhspan(0, cota_preforo, color='gray', alpha=0.3, hatch='//')

        axs[5].axis('off'); axs[5].legend(handles=leyenda_preforo, loc='center', title="Robertson (SBT)", fontsize=9)
        plt.tight_layout()
        return fig

    with tab1: st.pyplot(generar_figura_estandar('basicos'))
    with tab2: st.pyplot(generar_figura_estandar('diseno'))
    with tab3: st.pyplot(generar_figura_estandar('calidad'))
    
    # Procesar Capas
    df_valido = df_calc[df_calc['Depth_m'] >= cota_preforo].copy()
    df_valido['Depth_Interval'] = np.floor(df_valido['Depth_m'])
    res_1m = df_valido.groupby('Depth_Interval').agg(SBT_Predominante=('SBT_Name', lambda x: x.mode()[0]), Qc_Medio=('Qc', 'mean'), U2_Medio=('U2', 'mean'), Su_Medio=('Su_kPa', 'mean'), Phi_Medio=('Phi_deg', 'mean'), Gamma_Medio=('Gamma_kN3', 'mean'), N60_Medio=('N60', 'mean'), Z_min=('Depth_m', 'min'), Z_max=('Depth_m', 'max')).reset_index()
    res_1m = res_1m[res_1m['Qc_Medio'] > 0.05].copy()
    res_1m['Grupo'] = (res_1m['SBT_Predominante'] != res_1m['SBT_Predominante'].shift()).cumsum()
    capas = res_1m.groupby(['Grupo', 'SBT_Predominante']).agg(Desde_m=('Z_min', 'min'), Hasta_m=('Z_max', 'max'), Qc_MPa=('Qc_Medio', 'mean'), U2_kPa=('U2_Medio', 'mean'), Su_kPa=('Su_Medio', 'mean'), Phi_deg=('Phi_Medio', 'mean'), Gamma=('Gamma_Medio', 'mean'), SPT_N60=('N60_Medio', 'mean')).reset_index()
    capas['Espesor_m'] = capas['Hasta_m'] - capas['Desde_m']
    capas_final = capas[['Desde_m', 'Hasta_m', 'Espesor_m', 'SBT_Predominante', 'Qc_MPa', 'U2_kPa', 'Gamma', 'Su_kPa', 'Phi_deg', 'SPT_N60']]
    
    with tab4: st.dataframe(capas_final, hide_index=True)
    with tab5: st.dataframe(df_valido[['Depth_m', 'SBT_Name', 'Ic', 'Gamma_kN3', 'Su_kPa', 'Phi_deg', 'N60']], hide_index=True)

    # --- EXPORTACIÓN ---
    st.sidebar.divider()
    if st.sidebar.button("🚀 GENERAR DOCUMENTACIÓN", use_container_width=True):
        progress_bar = st.sidebar.progress(0)
        status_text = st.sidebar.empty()
        
        status_text.text("Generando Word...")
        doc = Document()
        doc.add_heading(f'Informe CPTU: {header_data.get("Location", "Ensayo")}', 0)
        for i, t in enumerate(['basicos', 'diseno', 'calidad']):
            fig = generar_figura_estandar(t)
            img_stream = io.BytesIO()
            fig.savefig(img_stream, format='png', dpi=150)
            doc.add_heading(f'Gráfico de {t.capitalize()}', level=1)
            doc.add_picture(img_stream, width=Inches(6))
            progress_bar.progress(20 + i*20)

        status_text.text("Generando Excel...")
        excel_stream = io.BytesIO()
        with pd.ExcelWriter(excel_stream, engine='xlsxwriter') as writer:
            pd.DataFrame(list(header_data.items())).to_excel(writer, sheet_name='Cabecera', index=False)
            capas_final.to_excel(writer, sheet_name='Estratigrafía', index=False)
            df_valido.to_excel(writer, sheet_name='Detalle', index=False)
        progress_bar.progress(80)

        zip_stream = io.BytesIO()
        with zipfile.ZipFile(zip_stream, 'w') as zf:
            word_stream = io.BytesIO()
            doc.save(word_stream)
            zf.writestr("Informe_CPTU.docx", word_stream.getvalue())
            zf.writestr("Datos_CPTU.xlsx", excel_stream.getvalue())
        
        progress_bar.progress(100)
        status_text.success("¡Hecho!")
        st.sidebar.download_button("📥 DESCARGAR (.ZIP)", zip_stream.getvalue(), f"Resultados_{header_data.get('Location', 'Ensayo')}.zip", "application/zip", use_container_width=True)