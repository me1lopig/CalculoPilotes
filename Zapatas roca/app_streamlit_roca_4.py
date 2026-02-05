import streamlit as st
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. CONFIGURACIÓN DE PÁGINA Y ESTILOS
st.set_page_config(page_title="Cimentaciones superficiales en Roca", layout="wide")

st.markdown("""
    <style>
    .main { background-color: #fcfcfc; }
    .titulo-seccion { background-color: #1b5e20; color: white; padding: 12px; font-weight: bold; border-radius: 4px; text-align: center; margin-bottom: 20px; }
    .titulo-norma { background-color: #1b5e20; color: white; padding: 8px 15px; font-weight: bold; border-radius: 4px 4px 0 0; font-size: 14px; margin-top: 10px; border: 1px solid #1b5e20; }
    .tabla-profesional { width: 100%; border-collapse: collapse; font-size: 13px; margin-bottom: 25px; border: 1px solid #d1d1d1; background-color: white; }
    .tabla-profesional th { background-color: #f5f5f5; color: #333; padding: 10px; border: 1px solid #d1d1d1; font-weight: bold; }
    .tabla-profesional td { padding: 10px; border: 1px solid #d1d1d1; text-align: center; color: #444; }
    .text-left { text-align: left !important; padding-left: 15px !important; }
    .grupo-roca { background-color: #fafafa; font-weight: 500; }
    .requisitos { background-color: #f1f8e9; padding: 15px; border-left: 5px solid #2e7d32; border-radius: 4px; font-size: 14px; line-height: 1.6; }
    .nota-pie-tabla { font-size: 11px; color: #666; margin-top: -15px; margin-bottom: 20px; padding-left: 5px; }
    </style>
    """, unsafe_allow_html=True)

st.title("🏗️ Cálculo de Presión Vertical Admisible en Roca")

# --- FUNCIÓN GENERADORA DE INFORME WORD ---
def generar_word(qu_inp, s_val, a_val, estado_j, b_range_info, checks, df_data, fig_plot):
    doc = Document()
    
    # Configuración de fuente base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # --- ESTILO UNIFICADO PARA TABLAS ---
    estilo_tablas = 'Light List Accent 1'

    # 1. ENCABEZADO
    titulo = doc.add_heading('Informe Técnico: Cimentaciones en Roca', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Fecha de emisión: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").bold = True
    p_meta.add_run("Metodología: Cálculo Analítico Simplificado (CTE)")

    # 2. PARÁMETROS DE DISEÑO
    doc.add_heading('1. Parámetros de Diseño', level=1)
    
    p_params = doc.add_paragraph()
    p_params.add_run(f"• Resistencia Compresión Simple (qu): ").bold = True
    p_params.add_run(f"{qu_inp} MPa\n")
    p_params.add_run(f"• Espaciamiento de juntas (s): ").bold = True
    p_params.add_run(f"{s_val} mm\n")
    p_params.add_run(f"• Apertura de juntas (a): ").bold = True
    p_params.add_run(f"{a_val} mm\n")
    p_params.add_run(f"• Estado de juntas: ").bold = True
    p_params.add_run(f"{estado_j}\n")
    p_params.add_run(f"• Rango de Anchos (B): ").bold = True
    p_params.add_run(f"{b_range_info[0]:.2f} m - {b_range_info[1]:.2f} m (Inc: {b_range_info[2]} m)")

    # 3. VERIFICACIÓN DE HIPÓTESIS (TABLA 1)
    doc.add_heading('2. Verificación de Hipótesis', level=1)
    
    table_checks = doc.add_table(rows=1, cols=2)
    table_checks.style = estilo_tablas
    
    hdr_cells = table_checks.rows[0].cells
    hdr_cells[0].text = 'Parámetro de Control'
    hdr_cells[1].text = 'Estado'
    
    for check_name, status in checks.items():
        row_cells = table_checks.add_row().cells
        row_cells[0].text = check_name
        
        p = row_cells[1].paragraphs[0]
        run = p.add_run(status)
        if "NO CUMPLE" in status:
            run.font.color.rgb = RGBColor(200, 0, 0)
            run.bold = True
        else:
            run.font.color.rgb = RGBColor(0, 100, 0)

    doc.add_paragraph()

    # 4. RESULTADOS (TABLA 2)
    doc.add_heading('3. Tabla de Resultados', level=1)
    
    t_res = doc.add_table(df_data.shape[0] + 1, df_data.shape[1])
    t_res.style = estilo_tablas
    
    for j in range(df_data.shape[-1]):
        t_res.cell(0, j).text = df_data.columns[j]
    
    for i in range(df_data.shape[0]):
        for j in range(df_data.shape[-1]):
            val = df_data.values[i, j]
            if isinstance(val, float):
                t_res.cell(i + 1, j).text = f"{val:.2f}"
            else:
                t_res.cell(i + 1, j).text = str(val)

    # 5. GRÁFICA (MEJORADA)
    doc.add_heading('4. Gráfica de Presión Admisible', level=1)
    
    try:
        # MEJORA: Aumentar resolución (scale=3) y tamaño de la imagen en píxeles
        img_bytes = fig_plot.to_image(format="png", width=1000, height=600, scale=3)
        # MEJORA: Ajustar el ancho de la imagen en Word
        doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph(f"[No se pudo generar la imagen del gráfico. Error: {e}]")
        doc.add_paragraph("Nota: Asegúrese de tener instalado el paquete 'kaleido'.")

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# --- SECCIÓN A: INDICACIONES Y FORMULACIÓN ---
st.markdown('<div class="titulo-seccion">INDICACIONES TÉCNICAS Y FORMULACIÓN</div>', unsafe_allow_html=True)

col_req, col_form = st.columns([0.6, 0.4])

with col_req:
    st.markdown("""
    <div class="requisitos">
        <strong>Requisitos del Cálculo Analítico Simplificado:</strong><br>
        • Para roca sana o poco meteorizada (<strong>qu ≥ 2.5 MPa, RQD > 25 y GM < IV</strong>).<br>
        • Superficie de la roca esencialmente horizontal y sin problemas de estabilidad lateral.<br>
        • Carga con componente tangencial inferior al 10% de la carga normal.<br>
        • En rocas sedimentarias los estratos deben ser horizontales o subhorizontales.
    </div>
    """, unsafe_allow_html=True)

with col_form:
    st.markdown("**Formulación Matemática:**")
    st.latex(r"K_{sp} = \frac{3 + \frac{s}{1000 \cdot B}}{10 \sqrt{1 + 300 \frac{a}{s}}}")
    st.latex(r"q_d = q_u \cdot K_{sp}")
    st.caption("Donde: $s$ y $a$ en mm, $B$ en m.")

st.divider()

# --- SECCIÓN B: NORMAS  ---
st.subheader("📚 Normas y Códigos de Uso Habitual")
col_izq, col_der = st.columns(2)

with col_izq:
    # DIN 1054
    st.markdown('<div class="titulo-norma">DIN 1054</div>', unsafe_allow_html=True)
    st.markdown("""<table class="tabla-profesional">
        <tr><th rowspan="2">Estado del macizo</th><th colspan="2">Presión Admisible (MPa)</th></tr>
        <tr><th>Roca sana / poco alterada</th><th>Roca quebradiza / alterada</th></tr>
        <tr><td class="text-left grupo-roca">Homogéneo</td><td>4.00</td><td>1.50</td></tr>
        <tr><td class="text-left grupo-roca">Estratificado o diaclasado</td><td>2.00</td><td>1.00</td></tr>
    </table>""", unsafe_allow_html=True)

    # CTE 2006
    st.markdown('<div class="titulo-norma">CTE 2006 (España)</div>', unsafe_allow_html=True)
    st.markdown("""
    <table class="tabla-profesional">
        <tr><th>Tipo de roca</th><th>q<sub>adm</sub> (MPa)</th></tr>
        <tr><td class="text-left grupo-roca">Rocas ígneas y metamórficas sanas (1)</td><td>10.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas metamórficas foliadas sanas (1) (2)</td><td>3.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas sedimentarias sanas (1) (2)</td><td>1.00 - 4.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas arcillosas sanas (2) (4)</td><td>0.50 - 1.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas diaclasadas (s > 0.30m)</td><td>1.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas muy diaclasadas o meteorizadas</td><td>(ver nota 3)</td></tr>
    </table>
    <div class="nota-pie-tabla">
        (1) Los valores indicados serán aplicables para estratificación y/o foliación subhorizontal. Los macizos rocosos con discontinuidades inclinadas, especialmente en las cercanías de taludes, deben ser objeto de análisis especial.<br>
        (2) Se admiten pequeñas discontinuidades con espaciamiento superior a 1m.<br>
        (3) Estos casos deben ser investigados "in situ".<br>
        (4) Rocas arcillosas sanas.
    </div>
    """, unsafe_allow_html=True)

with col_der:
    # CP 2004
    st.markdown('<div class="titulo-norma">CP 2004 / 1972</div>', unsafe_allow_html=True)
    st.markdown("""<table class="tabla-profesional">
        <tr><th>Tipo de roca</th><th>q<sub>adm</sub> (MPa)</th></tr>
        <tr><td class="text-left grupo-roca">Rocas ígneas (granitos y gneises), sanas</td><td>10.00</td></tr>
        <tr><td class="text-left grupo-roca">Calizas y areniscas duras</td><td>4.00</td></tr>
        <tr><td class="text-left grupo-roca">Esquistos y pizarras</td><td>3.00</td></tr>
        <tr><td class="text-left grupo-roca">Argilitas/limolitas duras, areniscas blandas</td><td>2.00</td></tr>
        <tr><td class="text-left grupo-roca">Arenas cementadas</td><td>1.00</td></tr>
    </table>""", unsafe_allow_html=True)

st.divider()

# --- SECCIÓN C: SIDEBAR Y LÓGICA ---
with st.sidebar:
    st.header("⚙️ Parámetros del Macizo")
    qu_input = st.number_input("Resistencia qu [MPa]", value=15.0, step=0.5)
    
    if qu_input < 2.5:
        st.warning(f"⚠️ Resistencia qu ({qu_input} MPa) inferior al límite. Se usará 2.5 MPa.")
        qu_calc = 2.5
    else:
        qu_calc = qu_input

    s = st.number_input("Espaciamiento s [mm]", value=301, step=10)
    a = st.number_input("Apertura a [mm]", value=3.0, step=0.1)
    
    op_limpia = "Limpias"
    op_rellena = "Rellenas con suelo"
    estado_junta = st.selectbox("Estado de juntas", [op_limpia, op_rellena])
    
    st.divider()
    st.header("📏 Rango de Anchos (B)")
    b_min = st.number_input("Ancho Mínimo B (m)", value=1.0, min_value=0.10, step=0.10)
    b_max = st.number_input("Ancho Máximo B (m)", value=3.00, min_value=b_min, step=0.50)
    b_step = st.selectbox("Incremento de B (m)", [0.25, 0.50, 1.00], index=1)

# Comprobaciones
c_s = s > 300
c_a = a < 5 if estado_junta == op_limpia else a < 25
rel_as = a/s
c_rel = 0 < rel_as < 0.02
c_qu = qu_input >= 2.5

checks_dict = {
    "Resistencia qu ≥ 2.5 MPa": "CUMPLE" if c_qu else "NO CUMPLE",
    "Espaciamiento s > 300mm": "CUMPLE" if c_s else "NO CUMPLE",
    f"Apertura a < {'5' if estado_junta == op_limpia else '25'}mm": "CUMPLE" if c_a else "NO CUMPLE",
    "Relación a/s < 0.02": "CUMPLE" if c_rel else "NO CUMPLE"
}

st.subheader("✅ Comprobaciones de Seguridad")
v1, v2, v3, v4 = st.columns(4)
v1.metric("qu ≥ 2.5 MPa", checks_dict["Resistencia qu ≥ 2.5 MPa"])
v2.metric("s > 300mm", checks_dict["Espaciamiento s > 300mm"])
v3.metric("Apertura (a)", checks_dict[f"Apertura a < {'5' if estado_junta == op_limpia else '25'}mm"])
v4.metric("Relación a/s", checks_dict["Relación a/s < 0.02"])

st.divider()

# --- SECCIÓN D: CÁLCULOS Y RESULTADOS ---
def calc_ksp(s_val, B_val, a_val):
    return (3 + (s_val / (B_val * 1000))) / (10 * np.sqrt(1 + 300 * (a_val / s_val)))

anchos_b = np.arange(b_min, b_max + 0.001, b_step)
filas = []

for b in anchos_b:
    k = calc_ksp(s, b, a)
    qd = qu_calc * k
    valido = "SÍ" if 0.05 < (s / (b * 1000)) < 2 else "NO"
    filas.append({
        "B (m)": b,
        "Válido 0,05<s/B<2": valido,
        "Ksp": k,
        "qd (MPa)": qd,
        "qd (kg/cm²)": qd * 1000/98.1
    })

df_res = pd.DataFrame(filas)

# Generación Gráfica (MEJORADA)
b_smooth = np.linspace(b_min, b_max, 100)
qd_smooth = [qu_calc * calc_ksp(s, b, a) for b in b_smooth]

fig = go.Figure(go.Scatter(x=b_smooth, y=qd_smooth, mode='lines', line=dict(color='#1b5e20', width=4)))
fig.update_layout(
    title="Curva de Presión Admisible",
    xaxis_title="Ancho de Cimentación B (m)",
    yaxis_title="Presión Admisible qd (MPa)",
    plot_bgcolor='white',
    height=500, # MEJORA: Altura aumentada
    # MEJORA: Márgenes aumentados para evitar cortes
    margin=dict(l=50, r=50, t=60, b=50)
)
fig.update_xaxes(showgrid=True, gridwidth=1, gridcolor='#eee')
fig.update_yaxes(showgrid=True, gridwidth=1, gridcolor='#eee')

col_res_t, col_res_g = st.columns([0.45, 0.55], gap="large")

with col_res_t:
    st.subheader("📋 Resultados Cálculo Analítico")
    st.dataframe(df_res, hide_index=True, use_container_width=True,
                 column_config={"B (m)": st.column_config.NumberColumn(format="%.2f"),
                                "Ksp": st.column_config.NumberColumn(format="%.3f"),
                                "qd (MPa)": st.column_config.NumberColumn(format="%.2f"),
                                "qd (kg/cm²)": st.column_config.NumberColumn(format="%.2f")})

with col_res_g:
    st.subheader("📈 Gráfico de Diseño")
    st.plotly_chart(fig, use_container_width=True)

# --- BOTÓN DE DESCARGA (SIDEBAR) ---
with st.sidebar:
    st.divider()
    st.header("📄 Informe Técnico")
    
    if st.button("Generar informe Word"):
        with st.spinner("Generando documento..."):
            try:
                doc_buffer = generar_word(
                    qu_inp=qu_calc,
                    s_val=s,
                    a_val=a,
                    estado_j=estado_junta,
                    b_range_info=[b_min, b_max, b_step],
                    checks=checks_dict,
                    df_data=df_res,
                    fig_plot=fig
                )
                
                file_name = f"Calculo_Roca_{datetime.now().strftime('%Y%m%d')}.docx"
                
                st.success("✅ Informe generado correctamente")
                st.download_button(
                    label="📥 Descargar .docx",
                    data=doc_buffer.getvalue(),
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Error: {e}")
                st.warning("Verifique que 'kaleido' esté instalado para exportar el gráfico.")