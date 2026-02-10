# Aplicacion para el cálculo de la carga admisible en roca
# Basada en el método analítico simplificado (CTE-DB-SE-C 2019)
# Desarrollado por: Germán López Pineda
# Fecha: febrero 2026

import streamlit as st
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. CONFIGURACIÓN DE PÁGINA
st.set_page_config(page_title="Cimentaciones superficiales en Roca", layout="wide")

# --- FUNCIÓN PARA RESETEAR INFORME ---
def reset_informe():
    if 'informe_buffer' in st.session_state:
        st.session_state.informe_buffer = None

# Estilos CSS
st.markdown("""
    <style>
    .main { background-color: #fcfcfc; }
    .titulo-seccion { background-color: #1b5e20; color: white; padding: 12px; font-weight: bold; border-radius: 4px; text-align: center; margin-bottom: 20px; }
    .titulo-norma { background-color: #1b5e20; color: white; padding: 8px 15px; font-weight: bold; border-radius: 4px 4px 0 0; font-size: 14px; margin-top: 10px; border: 1px solid #1b5e20; }
    .tabla-profesional { width: 100%; border-collapse: collapse; font-size: 13px; margin-bottom: 25px; border: 1px solid #d1d1d1; background-color: white; }
    .tabla-profesional th { background-color: #f5f5f5; color: #333; padding: 10px; border: 1px solid #d1d1d1; font-weight: bold; }
    .tabla-profesional td { padding: 10px; border: 1px solid #d1d1d1; text-align: center; color: #444; }
    .text-left { text-align: left !important; padding-left: 15px !important; }
    .grupo-roca { background-color: #fafafa; font-weight: 500; }
    .requisitos { background-color: #f1f8e9; padding: 15px; border-left: 5px solid #2e7d32; border-radius: 4px; font-size: 14px; line-height: 1.6; }
    .nota-pie-tabla { font-size: 11px; color: #666; margin-top: -15px; margin-bottom: 20px; padding-left: 5px; }
    </style>
    """, unsafe_allow_html=True)

st.title("🏗️ Cálculo de Presión Vertical Admisible en Roca CTE-DB-SE-C")

# --- CONSTANTES ---
OP_LIMPIA = "Limpias"
OP_RELLENA = "Rellenas con suelo"

# --- FUNCIÓN GENERADORA DE INFORME WORD ---
def generar_informe_word(qu_inp, s_val, a_val, estado_j, b_range_info, checks, df_data, fig_plot):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ENCABEZADO PRINCIPAL
    titulo = doc.add_heading('Informe resultados: Cimentaciones en Roca', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Fecha de emisión: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").bold = True
    p_meta.add_run("Metodología: Cálculo Analítico Simplificado (CTE-DB-SE-C 2019)")

    # 1. LIMITACIONES (NUEVO APARTADO BASADO EN LA IMAGEN)
    doc.add_heading('1. Limitaciones y Requisitos del Método', level=1)
    doc.add_paragraph("Este cálculo se basa en el método analítico simplificado, el cual es válido únicamente si se cumplen las siguientes condiciones del terreno y la carga:")
    
    # Lista de viñetas
    doc.add_paragraph("Para roca sana o poco meteorizada (qu ≥ 2.5 MPa, RQD > 25 y GM < IV).", style='List Bullet')
    doc.add_paragraph("Superficie de la roca esencialmente horizontal y sin problemas de estabilidad lateral.", style='List Bullet')
    doc.add_paragraph("Carga con componente tangencial inferior al 10% de la carga normal.", style='List Bullet')
    doc.add_paragraph("En rocas sedimentarias los estratos deben ser horizontales o subhorizontales.", style='List Bullet')

    # 2. PARÁMETROS (Re-enumerado)
    doc.add_heading('2. Parámetros de Diseño', level=1)
    p_params = doc.add_paragraph()
    p_params.add_run(f"• Resistencia (qu): {qu_inp} MPa\n")
    p_params.add_run(f"• Espaciamiento (s): {s_val} mm\n")
    p_params.add_run(f"• Apertura (a): {a_val} mm\n")
    p_params.add_run(f"• Estado juntas: {estado_j}\n")
    p_params.add_run(f"• Anchos (B): {b_range_info[0]:.2f}m - {b_range_info[1]:.2f}m")

    # 3. VERIFICACIÓN (Re-enumerado)
    doc.add_heading('3. Verificación de Hipótesis del Modelo', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parámetro de Control'
    hdr_cells[1].text = 'Estado'
    
    for check_name, status in checks.items():
        row_cells = table.add_row().cells
        row_cells[0].text = check_name
        p = row_cells[1].paragraphs[0]
        run = p.add_run(status)
        if "NO CUMPLE" in status:
            run.font.color.rgb = RGBColor(200, 0, 0)
            run.bold = True
        else:
            run.font.color.rgb = RGBColor(0, 100, 0)

    doc.add_paragraph()

    # 4. RESULTADOS (Re-enumerado)
    doc.add_heading('4. Tabla de Resultados', level=1)
    t_res = doc.add_table(df_data.shape[0] + 1, df_data.shape[1])
    t_res.style = 'Light List Accent 1'
    
    for j in range(df_data.shape[-1]):
        t_res.cell(0, j).text = df_data.columns[j]
    
    for i in range(df_data.shape[0]):
        for j in range(df_data.shape[-1]):
            val = df_data.values[i, j]
            if isinstance(val, float):
                t_res.cell(i + 1, j).text = f"{val:.2f}"
            else:
                t_res.cell(i + 1, j).text = str(val)

    # 5. GRÁFICA (Re-enumerado)
    doc.add_heading('5. Gráfica de Resultados', level=1)
    try:
        img_bytes = fig_plot.to_image(format="png", width=1000, height=600, scale=3)
        doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph(f"[No se pudo generar el gráfico: {e}]")

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# --- SECCIÓN A: INFO ---
st.markdown('<div class="titulo-seccion">INDICACIONES TÉCNICAS Y FORMULACIÓN</div>', unsafe_allow_html=True)
col_req, col_form = st.columns([0.6, 0.4])

with col_req:
    st.markdown("""
    <div class="requisitos">
        <strong>Requisitos del Cálculo Analítico Simplificado:</strong><br>
        • Para roca sana o poco meteorizada (<strong>qu ≥ 2.5 MPa, RQD > 25 y GM < IV</strong>).<br>
        • Superficie de la roca esencialmente horizontal y sin problemas de estabilidad lateral.<br>
        • Carga con componente tangencial inferior al 10% de la carga normal.<br>
        • En rocas sedimentarias los estratos deben ser horizontales o subhorizontales.
    </div>
    """, unsafe_allow_html=True)

with col_form:
    st.markdown("**Formulación usada:**")
    st.latex(r"K_{sp} = \frac{3 + \frac{s}{1000 \cdot B}}{10 \sqrt{1 + 300 \frac{a}{s}}}")
    st.latex(r"q_d = q_u \cdot K_{sp}")
    st.caption("Donde: $s$ y $a$ en mm, $B$ en m.")

st.divider()

# --- SECCIÓN B: NORMAS ---
st.subheader("📚 Normas y Códigos de referencia")
col_izq, col_der = st.columns(2)

with col_izq:
    st.markdown('<div class="titulo-norma">DIN 1054</div>', unsafe_allow_html=True)
    st.markdown("""<table class="tabla-profesional">
        <tr><th rowspan="2">Estado del macizo</th><th colspan="2">Presión Admisible (MPa)</th></tr>
        <tr><th>Roca sana / poco alterada</th><th>Roca quebradiza / alterada</th></tr>
        <tr><td class="text-left grupo-roca">Homogéneo</td><td>4.00</td><td>1.50</td></tr>
        <tr><td class="text-left grupo-roca">Estratificado o diaclasado</td><td>2.00</td><td>1.00</td></tr>
    </table>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown('<div class="titulo-norma">CP 2004 / 1972</div>', unsafe_allow_html=True)
    st.markdown("""<table class="tabla-profesional">
        <tr><th>Tipo de roca</th><th>q<sub>adm</sub> (MPa)</th></tr>
        <tr><td class="text-left grupo-roca">Rocas ígneas (granitos y gneises), sanas</td><td>10.00</td></tr>
        <tr><td class="text-left grupo-roca">Calizas y areniscas duras</td><td>4.00</td></tr>
        <tr><td class="text-left grupo-roca">Esquistos y pizarras</td><td>3.00</td></tr>
        <tr><td class="text-left grupo-roca">Argilitas/limolitas duras, areniscas blandas</td><td>2.00</td></tr>
        <tr><td class="text-left grupo-roca">Arenas cementadas</td><td>1.00</td></tr>
        <tr><td class="text-left grupo-roca">Argilitas y limolitas blandas</td><td>0.60 - 1.00</td></tr>
        <tr><td class="text-left grupo-roca">Calizas blandas y porosas</td><td>0.60</td></tr>
    </table>""", unsafe_allow_html=True)

with col_der:
    st.markdown('<div class="titulo-norma">CTE 2019 (España)</div>', unsafe_allow_html=True)
    st.markdown("""
    <table class="tabla-profesional">
        <tr><th>Tipo de roca</th><th>q<sub>adm</sub> (MPa)</th></tr>
        <tr><td class="text-left grupo-roca">Rocas ígneas y metamórficas sanas (1):<br><span style="font-size:11px; color:#555">Granito, diorita, basalto, gneis</span></td><td>10.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas metamórficas foliadas sanas (1) (2):<br><span style="font-size:11px; color:#555">Esquistos, pizarras</span></td><td>3.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas sedimentarias sanas (1) (2):<br><span style="font-size:11px; color:#555">Pizarras cementadas, limolitas, areniscas, calizas sin karstificar, conglomerados cementados</span></td><td>1.00 - 4.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas arcillosas sanas (2) (4)</td><td>0.50 - 1.00</td></tr>
        <tr><td class="text-left grupo-roca">Rocas diaclasadas de cualquier tipo con espaciamiento > 0.30m (excepto rocas arcillosas)</td><td>1.00</td></tr>
        <tr><td class="text-left grupo-roca">Calizas, areniscas y rocas pizarrosas con pequeño espaciamiento de estratificación</td><td>(ver 3)</td></tr>
        <tr><td class="text-left grupo-roca">Rocas muy diaclasadas o meteorizadas</td><td>(ver 3)</td></tr>
    </table>
    <div class="nota-pie-tabla">
        (1) Aplicable para estratificación subhorizontal.<br>
        (2) Discontinuidades peq. con s > 1m.<br>
        (3) Investigar "in situ". (4) Rocas arcillosas sanas.
    </div>
    """, unsafe_allow_html=True)

st.divider()

# --- SECCIÓN C: SIDEBAR (CON RESETEO AUTOMÁTICO) ---
with st.sidebar:
    st.header("⚙️ Parámetros del Macizo")
    
    qu_input = st.number_input("Resistencia qu [MPa]", value=15.0, step=0.5, on_change=reset_informe)
    
    # --- LÓGICA DE VALOR REAL CON ADVERTENCIA ---
    if qu_input < 2.5:
        st.warning(f"⚠️ Resistencia qu ({qu_input} MPa) inferior al rango normativo (mínimo 2.5 MPa). Se utilizará el valor real para el cálculo, aunque esté fuera del ámbito de aplicación estricto del método.")
        qu_calc = qu_input 
    else:
        qu_calc = qu_input
    # ---------------------------------------------------

    s = st.number_input("Espaciamiento s [mm]", value=301, step=10, min_value=1, on_change=reset_informe)
    a = st.number_input("Apertura a [mm]", value=3.0, step=0.1, on_change=reset_informe)
    
    estado_junta = st.selectbox("Estado de juntas", [OP_LIMPIA, OP_RELLENA], on_change=reset_informe)
    
    st.divider()
    st.header("📏 Rango de Anchos (B)")
    b_min = st.number_input("Ancho Mínimo B (m)", value=1.0, min_value=0.10, step=0.10, on_change=reset_informe)
    b_max = st.number_input("Ancho Máximo B (m)", value=3.00, min_value=b_min, step=0.50, on_change=reset_informe)
    b_step = st.selectbox("Incremento de B (m)", [0.25, 0.50, 1.00], index=1, on_change=reset_informe)

# --- LÓGICA DE COMPROBACIÓN ---
c_s = s > 300
c_qu = qu_input >= 2.5
c_rel = 0 < (a/s) < 0.02

if estado_junta == OP_LIMPIA:
    limite_a = 5
    texto_apertura = "Apertura a < 5 mm (Junta Limpia)"
else:
    limite_a = 25
    texto_apertura = "Apertura a < 25 mm (Junta Rellena)"

c_a = a < limite_a

checks_dict = {
    "Resistencia qu ≥ 2.5 MPa": "CUMPLE" if c_qu else "NO CUMPLE",
    "Espaciamiento s > 300mm": "CUMPLE" if c_s else "NO CUMPLE",
    texto_apertura: "CUMPLE" if c_a else "NO CUMPLE",
    "Relación a/s < 0.02": "CUMPLE" if c_rel else "NO CUMPLE"
}

# --- VISUALIZACIÓN: TABLA SIN ÍNDICE ---
st.subheader("✅ Verificación de las Hipótesis del Modelo")

df_checks = pd.DataFrame(list(checks_dict.items()), columns=["Parámetro de Control", "Estado"])
df_checks.set_index("Parámetro de Control", inplace=True)

def estilo_estado(val):
    color = '#2e7d32' if "CUMPLE" in val and "NO" not in val else '#c62828'
    return f'color: {color}; font-weight: bold'

st.table(df_checks.style.applymap(estilo_estado, subset=['Estado']))

st.divider()

# --- SECCIÓN D: CÁLCULOS Y GRÁFICO ---
def calc_ksp(s_val, B_val, a_val):
    return (3 + (s_val / (B_val * 1000))) / (10 * np.sqrt(1 + 300 * (a_val / s_val)))

anchos_b = np.arange(b_min, b_max + 0.001, b_step)
filas = []

for b in anchos_b:
    k = calc_ksp(s, b, a)
    qd = qu_calc * k
    valido = "SÍ" if 0.05 < (s / (b * 1000)) < 2 else "NO"
    filas.append({
        "B (m)": b,
        "Válido 0,05<s/B<2": valido,
        "Ksp": k,
        "qd (MPa)": qd,
        "qd (kg/cm²)": qd * 1000/98.1
    })

df_res = pd.DataFrame(filas)

# Gráfico
b_smooth = np.linspace(b_min, b_max, 100)
qd_smooth = [qu_calc * calc_ksp(s, b, a) for b in b_smooth]

fig = go.Figure(go.Scatter(x=b_smooth, y=qd_smooth, mode='lines', line=dict(color='#1b5e20', width=4)))
fig.update_layout(
    title="Curva de Presión Admisible",
    xaxis_title="Ancho B (m)", yaxis_title="qd (MPa)",
    plot_bgcolor='white', height=500, margin=dict(l=50, r=50, t=60, b=50)
)
fig.update_xaxes(showgrid=True, gridwidth=1, gridcolor='#eee')
fig.update_yaxes(showgrid=True, gridwidth=1, gridcolor='#eee')

col_t, col_g = st.columns([0.45, 0.55], gap="large")
with col_t:
    st.subheader("📋 Resultados Tabulados")
    st.dataframe(df_res, hide_index=True, use_container_width=True,
                 column_config={"B (m)": st.column_config.NumberColumn(format="%.2f"),
                                "Ksp": st.column_config.NumberColumn(format="%.3f"),
                                "qd (MPa)": st.column_config.NumberColumn(format="%.2f"),
                                "qd (kg/cm²)": st.column_config.NumberColumn(format="%.2f")})
with col_g:
    st.subheader("📈 Gráfico de Diseño")
    st.plotly_chart(fig, use_container_width=True)

# --- GESTIÓN DE ESTADO Y DESCARGA ---
if 'informe_buffer' not in st.session_state:
    st.session_state.informe_buffer = None

with st.sidebar:
    st.divider()
    st.header("📄 Informe Resultados")
    
    # 1. Generar
    if st.button("Generar informe Word"):
        with st.spinner("Generando informe..."):
            try:
                buffer = generar_informe_word(qu_calc, s, a, estado_junta, [b_min, b_max, b_step],
                                            checks_dict, df_res, fig)
                st.session_state.informe_buffer = buffer
                st.session_state.informe_nombre = f"Informe_CTE_{datetime.now().strftime('%H%M')}.docx"
                st.success("¡informe generado!")
            except Exception as e:
                st.error(f"Error generando informe: {e}")

    # 2. Descargar
    if st.session_state.informe_buffer is not None:
        st.download_button(
            label="📥 Descargar Documento",
            data=st.session_state.informe_buffer.getvalue(),
            file_name=st.session_state.informe_nombre,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )