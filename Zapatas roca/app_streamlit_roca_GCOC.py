import streamlit as st
import numpy as np
import pandas as pd
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# 1. CONFIGURACIÓN DE PÁGINA
st.set_page_config(page_title="Carga Admisible en Roca - GCOC 2009", layout="centered")

# --- FUNCIÓN PARA RESETEAR INFORME ---
def reset_informe():
    if 'informe_buffer' in st.session_state:
        st.session_state.informe_buffer = None

# Estilos CSS
st.markdown("""
    <style>
    .main { background-color: #fcfcfc; }
    .titulo-seccion { background-color: #0d47a1; color: white; padding: 12px; font-weight: bold; border-radius: 4px; text-align: center; margin-bottom: 20px; }
    .titulo-norma { background-color: #0d47a1; color: white; padding: 8px 15px; font-weight: bold; border-radius: 4px 4px 0 0; font-size: 14px; margin-top: 10px; border: 1px solid #0d47a1; }
    .tabla-profesional { width: 100%; border-collapse: collapse; font-size: 12px; margin-bottom: 25px; border: 1px solid #d1d1d1; background-color: white; }
    .tabla-profesional th { background-color: #365F91; color: white; padding: 8px; border: 1px solid #d1d1d1; font-weight: bold; vertical-align: middle;}
    .tabla-profesional td { padding: 8px; border: 1px solid #d1d1d1; text-align: center; color: #444; vertical-align: middle; }
    .text-left { text-align: left !important; padding-left: 10px !important; }
    .requisitos { background-color: #e3f2fd; padding: 15px; border-left: 5px solid #1565c0; border-radius: 4px; font-size: 14px; line-height: 1.6; }
    </style>
    """, unsafe_allow_html=True)

st.title("🛣️ Cálculo de Presión Vertical Admisible en Roca (GCOC 2009)")

# --- FUNCIONES AUXILIARES PARA WORD ---
def set_cell_background(cell, color_hex):
    """Establece el color de fondo de una celda en Word."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def estilizar_tabla(table):
    """Aplica el estilo azul corporativo a la tabla."""
    table.style = 'Table Grid'
    
    # Estilo del Encabezado (Fila 0)
    for cell in table.rows[0].cells:
        set_cell_background(cell, "365F91") # Azul oscuro profesional
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255) # Texto blanco
                run.bold = True
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- FUNCIÓN GENERADORA DE INFORME WORD ---
def generar_informe_word(inputs, resultados, checks):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # 0. ENCABEZADO
    titulo = doc.add_heading('Informe de resultados', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Fecha de emisión: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").bold = True
    p_meta.add_run("Metodología: Guía de Cimentaciones en Obras de Carretera (Cap. 4.5.3)")

    # 1. LIMITACIONES DEL MÉTODO
    doc.add_heading('1. Condiciones y Limitaciones del Método', level=1)
    p_intro = doc.add_paragraph("Según el apartado 4.5.3 de la GCOC, la validez de este cálculo está supeditada al cumplimiento de las siguientes condiciones de contorno:")
    p_intro.italic = True
    
    limitaciones = [
        "La pendiente del terreno no supera el 10%.",
        "La inclinación de las acciones no supera el 10% (tg δ < 0,10).",
        "No existe flujo de agua con gradiente importante (I ≤ 0,2).",
        "El área de apoyo es menor a 100 m² (para áreas mayores verificar asientos).",
        "La presión de servicio resultante no supera los 5 MPa."
    ]
    
    for item in limitaciones:
        doc.add_paragraph(item, style='List Bullet')

    # 2. PARÁMETROS DE ENTRADA
    doc.add_heading('2. Parámetros del Macizo Rocoso', level=1)
    p = doc.add_paragraph()
    p.add_run(f"• Resistencia Compresión Simple (qu): {inputs['qu']} MPa\n")
    p.add_run(f"• RQD: {inputs['rqd']} %\n")
    p.add_run(f"• Espaciamiento discontinuidades (s): {inputs['s']:.2f} m\n")
    p.add_run(f"• Tipo de Roca: {inputs['txt_a1']}\n")
    p.add_run(f"• Grado Meteorización: {inputs['txt_a2']}")

    # 3. VERIFICACIÓN (TABLA)
    doc.add_heading('3. Verificación de las Hipótesis del Modelo', level=1)
    t_checks = doc.add_table(rows=1, cols=2)
    
    # Encabezados
    hdr_cells = t_checks.rows[0].cells
    hdr_cells[0].text = 'Criterio'
    hdr_cells[1].text = 'Estado'
    
    # Datos
    for check_name, status in checks.items():
        row_cells = t_checks.add_row().cells
        row_cells[0].text = check_name
        p_row = row_cells[1].paragraphs[0]
        run = p_row.add_run(status)
        p_row.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if "NO CUMPLE" in status or "SUELO" in status:
            run.font.color.rgb = RGBColor(200, 0, 0) # Rojo
            run.bold = True
        else:
            run.font.color.rgb = RGBColor(0, 128, 0) # Verde oscuro
            run.bold = True

    # Aplicar estilo único
    estilizar_tabla(t_checks)
    doc.add_paragraph()

    # 4. FICHA DE CÁLCULO (RESULTADOS)
    doc.add_heading('4. Resultados del Cálculo', level=1)
    
    t_res = doc.add_table(rows=1, cols=3) # Creamos solo header, luego rows
    
    # Encabezados
    row0 = t_res.rows[0].cells
    row0[0].text = "Parámetro"
    row0[1].text = "Valor"
    row0[2].text = "Descripción"
    
    # Datos
    data_rows = [
        ("Coef. α1", f"{inputs['a1']}", "Factor por Tipo de Roca"),
        ("Coef. α2", f"{inputs['a2']}", "Factor por Meteorización"),
        ("Coef. α3", f"{resultados['a3']:.3f}", "Factor por Espaciamiento"),
        ("Presión Base (p0)", "1.0 MPa", "Presión de referencia"),
        ("PRESIÓN ADMISIBLE", f"{resultados['q_adm']:.2f} MPa", "Valor de diseño")
    ]
    
    for i, (param, val, desc) in enumerate(data_rows):
        cells = t_res.add_row().cells
        cells[0].text = param
        cells[1].text = val
        cells[2].text = desc
        
        # Alinear columnas
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Resaltar la última fila (Resultado)
        if i == len(data_rows) - 1:
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # Aplicar estilo único
    estilizar_tabla(t_res)

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# --- SECCIÓN A: INFO (LIMITACIONES) ---
st.markdown('<div class="titulo-seccion">MÉTODO GCOC (4.5.3)</div>', unsafe_allow_html=True)

st.warning("⚠️ **Limitaciones del procedimiento de referencia:**")
st.markdown("""
<div style="font-size: 13px; line-height: 1.5; margin-bottom: 20px;">
    <ul>
        <li>La pendiente del terreno no debe superar el <strong>10%</strong>.</li>
        <li>La inclinación de las acciones no debe superar el <strong>10% (tg δ < 0,10)</strong>. Si es mayor, se debe aplicar el factor reductor <em>i = (1,1 - tg δ)³</em>.</li>
        <li>No debe existir un flujo de agua con gradiente importante <strong>(I ≤ 0,2)</strong>.</li>
        <li>El área de apoyo debe ser menor a <strong>100 m²</strong>.</li>
        <li>La presión de servicio no superará los <strong>5 MPa</strong>.</li>
    </ul>
</div>
""", unsafe_allow_html=True)

col_req, col_form = st.columns([0.4, 0.6])

with col_req:
    st.markdown("""
    <div class="requisitos">
        <strong>Ámbito (Ap. 4.5.3):</strong><br>
        • qu ≥ 1 MPa.<br>
        • RQD ≥ 10%.<br>
        • Meteorización I, II, III.
    </div>
    """, unsafe_allow_html=True)

with col_form:
    st.markdown("**Formulación corregida (4.5.3.3):**")
    st.latex(r"p_{v,adm} = \alpha_1 \cdot \alpha_2 \cdot \alpha_3 \cdot q_u")
    st.latex(r"\alpha_3 = \min \left( \sqrt{\frac{s}{1}}, \sqrt{\frac{RQD}{100}} \right)")

st.divider()

# --- SECCIÓN B: TABLAS COMPLETAS ---
st.subheader("📚 Tablas de Coeficientes")

# TABLA 4.3 COMPLETA (Visualización web estilizada para coincidir con el informe)
st.markdown('<div class="titulo-norma">Tabla 4.3: Valores de α1 según el tipo de roca</div>', unsafe_allow_html=True)
st.markdown("""
<table class="tabla-profesional">
  <tr>
    <th style="width:10%">Grupo</th>
    <th style="width:30%">Nombre Genérico</th>
    <th style="width:50%">Ejemplos</th>
    <th style="width:10%">α1</th>
  </tr>
  <tr>
    <td>1</td>
    <td class="text-left">Rocas carbonatadas con estructura bien desarrollada</td>
    <td class="text-left">• Calizas, dolomías y mármoles puros<br>• Calcarenitas de baja porosidad</td>
    <td>1,0</td>
  </tr>
  <tr>
    <td>2</td>
    <td class="text-left">Rocas ígneas y rocas metamórficas (*)</td>
    <td class="text-left">• Granitos, cuarcitas<br>• Andesitas, riolitas<br>• Pizarras, esquistos y gneises (esquistosidad subhorizontal)</td>
    <td>0,8</td>
  </tr>
  <tr>
    <td>3</td>
    <td class="text-left">Rocas sedimentarias (**) y algunas metamórficas</td>
    <td class="text-left">• Calizas margosas, argilitas, limolitas, areniscas y conglomerados<br>• Pizarras y esquistos (esquistosidad verticalizada)<br>• Yesos</td>
    <td>0,6</td>
  </tr>
  <tr>
    <td>4</td>
    <td class="text-left">Rocas poco soldadas</td>
    <td class="text-left">• Areniscas, limolitas y conglomerados poco cementados<br>• Margas</td>
    <td>0,4</td>
  </tr>
</table>
<div style="font-size: 11px; color: #666; margin-top: -15px; margin-bottom: 20px; padding-left: 10px;">
    (*) A excepción de las indicadas en los grupos 1 y 3.<br>
    (**) A excepción de las indicadas en los grupos 1 y 4.
</div>
""", unsafe_allow_html=True)

# TABLA METEORIZACIÓN
st.markdown('<div class="titulo-norma">Tabla 4.4: Valores de α2 según meteorización</div>', unsafe_allow_html=True)
st.markdown("""
<table class="tabla-profesional">
    <tr><th>Grado</th><th>Descripción</th><th>α2</th></tr>
    <tr><td>I</td><td class="text-left">Roca sana o fresca</td><td>1.0</td></tr>
    <tr><td>II</td><td class="text-left">Ligeramente meteorizada</td><td>0.7</td></tr>
    <tr><td>III</td><td class="text-left">Moderadamente meteorizada</td><td>0.5</td></tr>
</table>
""", unsafe_allow_html=True)

st.divider()

# --- SECCIÓN C: SIDEBAR (INPUTS) ---
with st.sidebar:
    st.header("⚙️ Parámetros del Macizo")
    qu_input = st.number_input("Resistencia qu [MPa]", value=10.0, step=0.5, on_change=reset_informe)
    rqd_input = st.number_input("RQD [%]", value=50.0, min_value=0.0, max_value=100.0, step=1.0, on_change=reset_informe)
    s_input = st.number_input("Espaciamiento s [m]", value=0.30, step=0.05, min_value=0.01, format="%.2f", on_change=reset_informe)

    st.divider()
    st.subheader("⛰️ Factores de Calidad (α)")
    
    opciones_a1 = {
        "G1: Carbonatadas bien desarrolladas (Calizas, Dolomías)": 1.0,
        "G2: Ígneas/Metamórficas (Esquistosidad Horizontal)": 0.8,
        "G3: Sedimentarias/Metamórficas (Esquistosidad Vertical)": 0.6,
        "G4: Rocas poco soldadas (Margas, Areniscas flojas)": 0.4
    }
    txt_a1 = st.selectbox("Tipo de Roca (α1)", list(opciones_a1.keys()), on_change=reset_informe)
    val_a1 = opciones_a1[txt_a1]

    opciones_a2 = {
        "Grado I: Roca sana": 1.0,
        "Grado II: Ligeramente meteorizada": 0.7,
        "Grado III: Moderadamente meteorizada": 0.5,
        "Grado IV o superior": 0.0
    }
    txt_a2 = st.selectbox("Meteorización (α2)", list(opciones_a2.keys()), on_change=reset_informe)
    val_a2 = opciones_a2[txt_a2]

# --- CÁLCULO ---
c_qu = qu_input >= 1.0
c_rqd = rqd_input >= 10.0
c_met = val_a2 > 0.0

checks_dict = {
    "Resistencia qu ≥ 1 MPa": "CUMPLE" if c_qu else "NO CUMPLE",
    "RQD ≥ 10%": "CUMPLE" if c_rqd else "NO CUMPLE",
    "Meteorización < Grado IV": "CUMPLE" if c_met else "NO CUMPLE"
}

val_a3 = min(np.sqrt(s_input), np.sqrt(rqd_input / 100.0))

if c_qu and c_rqd and c_met:
    q_calc = val_a1 * val_a2 * val_a3 * qu_input
    q_final = min(q_calc, 5.0)
    nota_limite = "Limitado a 5 MPa (Máx Norma)" if q_calc > 5.0 else None
else:
    q_final = 0.0
    nota_limite = "Fuera de Norma"

# --- VISUALIZACIÓN ---
st.subheader("✅ Verificación de Hipótesis")
df_checks = pd.DataFrame(list(checks_dict.items()), columns=["Criterio", "Estado"]).set_index("Criterio")
st.table(df_checks.style.applymap(lambda v: f"color: {'#c62828' if 'NO' in v else '#2e7d32'}; font-weight: bold"))

st.divider()

st.subheader("📋 Ficha de Cálculo")
col_res1, col_res2 = st.columns([0.6, 0.4])

with col_res1:
    df_resumen = pd.DataFrame([
        {"Parámetro": "α1 (Tipo)", "Valor": val_a1, "Descripción": txt_a1.split(":")[0]},
        {"Parámetro": "α2 (Meteor.)", "Valor": val_a2, "Descripción": txt_a2.split(":")[0]},
        {"Parámetro": "α3 (Discont.)", "Valor": f"{val_a3:.3f}", "Descripción": "min(√s, √RQD/100)"},
        {"Parámetro": "Resistencia (qu)", "Valor": f"{qu_input} MPa", "Descripción": "Dato de entrada"},
    ])
    st.table(df_resumen.set_index("Parámetro"))

with col_res2:
    st.markdown("### Presión Admisible")
    st.markdown(f"""
    <div style="background-color: #e8f5e9; padding: 20px; border-radius: 10px; border: 2px solid #2e7d32; text-align: center;">
        <span style="font-size: 16px; color: #2e7d32; font-weight: bold;">Valor de Diseño (qd)</span><br>
        <span style="font-size: 42px; color: #1b5e20; font-weight: bold;">{q_final:.2f} MPa</span><br>
        <span style="font-size: 14px; color: #555;">{(q_final * 10.197):.2f} kg/cm²</span>
    </div>
    """, unsafe_allow_html=True)
    if nota_limite:
        st.warning(f"⚠️ {nota_limite}")

# --- DESCARGA ---
with st.sidebar:
    st.divider()
    if st.button("Generar informe Word"):
        inputs_repo = {"qu": qu_input, "rqd": rqd_input, "s": s_input, "a1": val_a1, "txt_a1": txt_a1, "a2": val_a2, "txt_a2": txt_a2}
        res_repo = {"a3": val_a3, "q_adm": q_final}
        st.session_state.informe_buffer = generar_informe_word(inputs_repo, res_repo, checks_dict)
        st.session_state.informe_nombre = f"Informe_GCOC_{datetime.now().strftime('%H%M')}.docx"

    if st.session_state.get('informe_buffer'):
        st.download_button("📥 Descargar Documento", st.session_state.informe_buffer.getvalue(), st.session_state.informe_nombre, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")