import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ==========================================
# CONSTANTES
# ==========================================
GAMMA_AGUA = 9.81  # kN/m³

# ==========================================
# FUNCIONES: TENSIONES DE HOLL BAJO EL CENTRO
# ==========================================
# La solución de Holl da tensiones bajo la ESQUINA de un rectángulo BxL.
# Para obtener tensiones en el CENTRO de la zapata se divide en 4 cuadrantes
# iguales (B/2 x L/2), se calcula en la esquina de cada uno y se suman los 4
# valores (principio de superposición — UNE-EN 1997-1, 6.6.2(15)).

def holl_esquina(p, B, L, z):
    """
    Tensiones bajo la ESQUINA de una carga rectangular BxL a profundidad z.
    Solución de Holl (ver imagen Ec. Holl).
    Devuelve (σz, σx, σy) en kPa.
    """
    if z <= 1e-6:
        return p, p / 2.0, p / 2.0

    R1 = np.sqrt(L**2 + z**2)
    R2 = np.sqrt(B**2 + z**2)
    R3 = np.sqrt(L**2 + B**2 + z**2)

    arctan_term = np.arctan((B * L) / (z * R3))

    sigma_z = (p / (2 * np.pi)) * (arctan_term + B * L * (1/R1**2 + 1/R2**2) * (z / R3))
    sigma_x = (p / (2 * np.pi)) * (arctan_term - (B * L * z) / (R1**2 * R3))
    sigma_y = (p / (2 * np.pi)) * (arctan_term - (B * L * z) / (R2**2 * R3))

    return sigma_z, sigma_x, sigma_y


def holl_centro(p, B, L, z):
    """
    Tensiones bajo el CENTRO de la cimentación a profundidad z.
    Se divide la zapata en 4 cuadrantes de (B/2 x L/2) y se superponen.
    """
    sz, sx, sy = holl_esquina(p, B / 2.0, L / 2.0, z)
    return 4 * sz, 4 * sx, 4 * sy


# ==========================================
# FUNCIÓN: ASIENTO EC. 68
# ==========================================
def calcular_asiento_ec68(p, B, L, df_terreno, z_max):
    """
    Calcula el asiento elástico usando la Ec. 68 de la Guía EC7:
        s = Σ [ h/E · (Δσz - ν·(Δσx + Δσy)) ]_i

    Las tensiones se evalúan en el punto medio de cada estrato.
    Devuelve (asiento_total_m, lista_de_resultados_por_capa).
    """
    asiento_total = 0.0
    resultados = []
    z_actual = 0.0

    for _, row in df_terreno.iterrows():
        if z_actual >= z_max:
            break

        h_i   = float(row["Espesor (m)"])
        E_i   = float(row["E (kPa)"])
        nu_i  = float(row["nu"])
        nombre = str(row["Descripción"])

        z_techo = z_actual
        z_base  = min(z_actual + h_i, z_max)   # recorta en z_max
        h_efec  = z_base - z_techo             # espesor efectivo del tramo
        z_mid   = (z_techo + z_base) / 2.0    # punto medio del estrato

        # Tensiones de Holl en el centro de la zapata, a z_mid
        dsz, dsx, dsy = holl_centro(p, B, L, z_mid)

        # Deformación unitaria vertical (Ec. 68)
        dep_z = (dsz - nu_i * (dsx + dsy)) / E_i

        # Asiento aportado por este estrato
        delta_s = dep_z * h_efec
        asiento_total += delta_s

        resultados.append({
            "Capa":              nombre,
            "z Techo [m]":       round(z_techo, 3),
            "z Base [m]":        round(z_base,  3),
            "h_ef [m]":          round(h_efec,  3),
            "z_mid [m]":         round(z_mid,   3),
            "Δσz [kPa]":         round(dsz,     3),
            "Δσx [kPa]":         round(dsx,     3),
            "Δσy [kPa]":         round(dsy,     3),
            "Δεz [-]":           round(dep_z,   6),
            "Δs [mm]":           round(delta_s * 1000, 3),
        })

        z_actual = z_base

    return asiento_total, resultados


# ==========================================
# FUNCIÓN: TENSIÓN EFECTIVA VERTICAL σ'v0
# ==========================================
def calcular_sigma_v0(z, df_terreno, NF):
    """σ'v0 (kPa) a profundidad z considerando el nivel freático NF."""
    sigma_v0 = 0.0
    z_actual = 0.0

    for _, row in df_terreno.iterrows():
        espesor  = float(row["Espesor (m)"])
        gamma    = float(row["Peso Esp. (kN/m³)"])
        gamma_sat= float(row["Peso Esp. Sat (kN/m³)"])

        z_techo = z_actual
        z_base  = z_actual + espesor

        if z <= z_techo:
            break

        z_eval = min(z, z_base)

        # Tramo seco
        z_seco_base = min(z_eval, NF)
        if z_seco_base > z_techo:
            sigma_v0 += gamma * (z_seco_base - z_techo)

        # Tramo sumergido
        z_sat_techo = max(z_techo, NF)
        if z_eval > z_sat_techo:
            sigma_v0 += (gamma_sat - GAMMA_AGUA) * (z_eval - z_sat_techo)

        z_actual = z_base

    return sigma_v0


# ==========================================
# FUNCIÓN: PROFUNDIDAD DE INFLUENCIA EC7
# ==========================================
def calcular_z_influencia_ec7(p, B, L, df_terreno, NF):
    """z_i donde Δσz ≤ 0.20·σ'v0 (criterio EC7)."""
    espesor_total = float(pd.to_numeric(df_terreno["Espesor (m)"]).sum())
    z = 0.05
    while z <= espesor_total:
        dsz, _, _ = holl_centro(p, B, L, z)
        sv0 = calcular_sigma_v0(z, df_terreno, NF)
        if sv0 > 0 and dsz <= 0.20 * sv0:
            return z
        z += 0.05
    return espesor_total


# ==========================================
# INFORME WORD
# ==========================================
def _fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    return buf

def _cab(celda, texto):
    celda.text = texto
    run = celda.paragraphs[0].runs[0]
    run.font.bold  = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size  = Pt(8)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1A3A5C')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def _tabla_word(doc, df, titulo):
    doc.add_heading(titulo, level=2)
    cols = list(df.columns)
    tbl  = doc.add_table(rows=1 + len(df), cols=len(cols))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(cols):
        _cab(tbl.cell(0, j), c)
    for i, (_, row) in enumerate(df.iterrows(), 1):
        for j, val in enumerate(row):
            celda = tbl.cell(i, j)
            celda.text = str(val)
            celda.paragraphs[0].runs[0].font.size = Pt(7)
            if i % 2 == 0:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tc = celda._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'E8F0FE')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    doc.add_paragraph()

def generar_informe_word(B, L, p, NF, z_max, z_i, df_terreno, df_resultados,
                         asiento_total, fig_bulbo_bytes, fig_bar_bytes):
    """Genera el informe Word y devuelve bytes."""
    fecha = datetime.now().strftime("%d/%m/%Y %H:%M")
    doc   = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # Portada
    t = doc.add_heading('INFORME DE CÁLCULO DE ASIENTOS', level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph('Formulación Ec. 68 — Guía de Cimentaciones EC7')
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.italic = True
    f = doc.add_paragraph(f'Fecha: {fecha}')
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Datos de entrada
    doc.add_heading('1. Datos de Entrada', level=1)
    doc.add_heading('1.1 Parámetros Geométricos y de Carga', level=2)
    params = [
        ('Ancho (B)',               f'{B:.2f} m'),
        ('Longitud (L)',             f'{L:.2f} m'),
        ('Esbeltez (n = L/B)',       f'{L/B:.3f}'),
        ('Presión neta (p)',         f'{p:.1f} kPa'),
        ('Nivel freático (NF)',       f'{NF:.1f} m'),
        ('Prof. de corte (z_max)',   f'{z_max:.1f} m'),
        ('Prof. influencia EC7 (z_i)', f'{z_i:.2f} m'),
    ]
    tgeo = doc.add_table(rows=len(params), cols=2)
    tgeo.style = 'Table Grid'
    for i, (par, val) in enumerate(params):
        tgeo.cell(i, 0).text = par
        tgeo.cell(i, 1).text = val
        tgeo.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        for j in range(2):
            tgeo.cell(i, j).paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()
    _tabla_word(doc, df_terreno, '1.2 Estratigrafía del Terreno')
    doc.add_page_break()

    # 2. Resultados
    doc.add_heading('2. Resultados del Cálculo', level=1)
    _tabla_word(doc, df_resultados, '2.1 Asientos por Estrato (Ec. 68)')
    p_tot = doc.add_paragraph()
    run   = p_tot.add_run(f'   ASIENTO TOTAL ESTIMADO: {asiento_total * 1000:.3f} mm')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(26, 58, 92)
    doc.add_paragraph()

    doc.add_heading('2.2 Asiento por Estrato — Gráfico', level=2)
    doc.add_picture(fig_bar_bytes, width=Cm(13))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 3. Bulbo
    doc.add_heading('3. Bulbo de Presiones y Zona de Influencia EC7', level=1)
    doc.add_picture(fig_bulbo_bytes, width=Cm(12))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    nota = doc.add_paragraph(
        f'B={B} m · L={L} m · p={p} kPa · NF={NF} m — {fecha}')
    nota.runs[0].font.size = Pt(8)
    nota.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==========================================
# GESTIÓN DE ESTADO (SESSION STATE)
# ==========================================
def reset_calculo():
    st.session_state.calculo_realizado = False

if 'calculo_realizado' not in st.session_state:
    st.session_state.calculo_realizado = False

if 'df_terreno' not in st.session_state:
    st.session_state.df_terreno = pd.DataFrame({
        "Descripción":           ["Relleno",  "Arcilla", "Grava"],
        "Espesor (m)":           [1.5,         3.0,       5.0],
        "E (kPa)":               [10000.0,     5000.0,    40000.0],
        "nu":                    [0.30,         0.45,      0.25],
        "Peso Esp. (kN/m³)":     [18.0,         19.0,      21.0],
        "Peso Esp. Sat (kN/m³)": [20.0,         20.0,      22.0],
    })

# ==========================================
# CONFIGURACIÓN DE PÁGINA
# ==========================================
st.set_page_config(
    page_title="Cálculo Asientos EC7 — Ec. 68",
    layout="wide",
    page_icon="🏗️"
)

st.sidebar.title("Navegación")
modo_vista = st.sidebar.radio(
    "Selecciona vista:",
    ("🧮 Panel de Cálculo", "📋 Detalle por Estrato",
     "📉 Bulbo de Presiones", "📖 Fundamento Teórico")
)

st.sidebar.markdown("---")
st.sidebar.header("📥 Datos de Entrada")

B  = st.sidebar.number_input("Ancho zapata (B) [m]",         min_value=0.1, value=2.0,   step=0.1,  on_change=reset_calculo)
L  = st.sidebar.number_input("Longitud zapata (L) [m]",      min_value=0.1, value=3.0,   step=0.1,  on_change=reset_calculo)
p  = st.sidebar.number_input("Presión neta (p) [kPa]",       min_value=1.0, value=150.0, step=10.0, on_change=reset_calculo)
NF = st.sidebar.number_input("Nivel Freático [m desde cim.]",min_value=0.0, value=100.0, step=0.5,  on_change=reset_calculo)

if L < B:
    B, L = L, B
    st.sidebar.warning("⚠️ L < B: se han intercambiado.")
st.sidebar.info(f"**Esbeltez n = L/B:** {L/B:.2f}")

st.sidebar.markdown("---")

# --- Profundidad de influencia ---
espesor_total = float(pd.to_numeric(st.session_state.df_terreno["Espesor (m)"]).sum())
espesor_total = max(espesor_total, 0.1)

z_i_ec7 = calcular_z_influencia_ec7(p, B, L, st.session_state.df_terreno, NF)

st.sidebar.subheader("📐 Profundidad de Cálculo")
if z_i_ec7 >= espesor_total - 0.05:
    st.sidebar.warning(f"⚠️ Bulbo supera la estratigrafía definida (> {espesor_total:.1f} m).")
else:
    st.sidebar.success(f"💡 **Recomendación EC7:** z_i = **{z_i_ec7:.2f} m**")

z_max_user = st.sidebar.number_input(
    "Profundidad de corte (z_max) [m]",
    min_value=0.1,
    max_value=espesor_total,
    value=float(min(round(z_i_ec7, 1), espesor_total)),
    step=0.1,
    on_change=reset_calculo
)

# --- BOTÓN CALCULAR ---
st.sidebar.markdown("---")
if st.sidebar.button("🚀 Calcular Asiento Total", type="primary", use_container_width=True):
    asiento_total, resultados = calcular_asiento_ec68(
        p, B, L, st.session_state.df_terreno, z_max_user
    )
    st.session_state.asiento_total   = asiento_total
    st.session_state.df_resultados   = pd.DataFrame(resultados)
    st.session_state.calculo_realizado = True

# --- BOTÓN INFORME WORD ---
st.sidebar.markdown("---")
if st.session_state.calculo_realizado:

    # Generar figuras en memoria para el informe
    df_res = st.session_state.df_resultados

    # Figura barras
    fig_bar_rep, ax_br = plt.subplots(figsize=(6, 3))
    bars = ax_br.barh(df_res["Capa"], df_res["Δs [mm]"], color='#1a3a5c', edgecolor='white')
    ax_br.bar_label(bars, fmt='%.3f mm', padding=3, fontsize=8)
    ax_br.set_xlabel("Δs (mm)")
    ax_br.set_title("Asiento por estrato (Ec. 68)")
    ax_br.invert_yaxis()
    ax_br.grid(axis='x', linestyle=':', alpha=0.5)
    ax_br.spines[['top', 'right']].set_visible(False)
    plt.tight_layout()
    fig_bar_bytes = _fig_to_bytes(fig_bar_rep)
    plt.close(fig_bar_rep)

    # Figura bulbo
    z_vals = np.linspace(0.05, espesor_total, 200)
    sz_v, sx_v, sy_v, sv0_v = [], [], [], []
    for z in z_vals:
        sz, sx, sy = holl_centro(p, B, L, z)
        sz_v.append(sz); sx_v.append(sx); sy_v.append(sy)
        sv0_v.append(calcular_sigma_v0(z, st.session_state.df_terreno, NF) * 0.20)

    fig_bul_rep, ax_bul = plt.subplots(figsize=(6, 8))
    ax_bul.plot(sz_v,  z_vals, label=r'Vertical $\Delta\sigma_z$',           color='red',    lw=2)
    ax_bul.plot(sx_v,  z_vals, label=r'Horiz. Trans. $\Delta\sigma_x$',      color='blue',   ls='--')
    ax_bul.plot(sy_v,  z_vals, label=r'Horiz. Long. $\Delta\sigma_y$',       color='purple', ls='-.')
    ax_bul.plot(sv0_v, z_vals, label=r"$0.20\,\sigma'_{v0}$ (criterio EC7)", color='green',  lw=2)
    if z_i_ec7 < espesor_total:
        ax_bul.axhline(y=z_i_ec7, color='orange', ls=':', lw=1.5, label=f'z_i = {z_i_ec7:.2f} m')
    if NF < espesor_total:
        ax_bul.axhline(y=NF, color='deepskyblue', ls='-.', lw=1.2, label=f'NF = {NF:.1f} m')
    ax_bul.set_ylim(espesor_total, 0)
    ax_bul.set_xlim(left=0)
    ax_bul.set_xlabel("Tensión (kPa)")
    ax_bul.set_ylabel("Profundidad z (m)")
    ax_bul.set_title("Bulbo de presiones — centro de la zapata")
    ax_bul.legend(fontsize=8)
    ax_bul.grid(True, linestyle=':', alpha=0.4)
    ax_bul.spines[['top', 'right']].set_visible(False)
    plt.tight_layout()
    fig_bulbo_bytes = _fig_to_bytes(fig_bul_rep)
    plt.close(fig_bul_rep)

    word_bytes = generar_informe_word(
        B, L, p, NF, z_max_user, z_i_ec7,
        st.session_state.df_terreno,
        df_res,
        st.session_state.asiento_total,
        fig_bulbo_bytes,
        fig_bar_bytes,
    )
    st.sidebar.download_button(
        label="📝 Descargar Informe Word",
        data=word_bytes,
        file_name=f"informe_asientos_ec68_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
else:
    st.sidebar.button("📝 Descargar Informe Word", disabled=True,
                      use_container_width=True,
                      help="Primero calcula el asiento total.")

# ==========================================
# ÁREA PRINCIPAL
# ==========================================
st.title("🏗️ Cálculo de Asientos — Formulación Ec. 68 (EC7)")
st.markdown(
    r"Integración directa de deformaciones unitarias: "
    r"$s = \sum \left[\dfrac{h}{E}\left(\Delta\sigma_z - \nu(\Delta\sigma_x + \Delta\sigma_y)\right)\right]_i$"
)
st.markdown("---")

# --- VISTA 1: PANEL DE CÁLCULO ---
if modo_vista == "🧮 Panel de Cálculo":
    st.header("1. Estratigrafía del Terreno")
    st.info("Introduce aquí las capas de terreno. Los **pesos específicos** se usan para el criterio de profundidad de influencia EC7.")

    df_edit = st.data_editor(
        st.session_state.df_terreno,
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "E (kPa)":               st.column_config.NumberColumn("E (kPa)",           min_value=100.0,  step=500.0),
            "nu":                    st.column_config.NumberColumn("ν (Poisson)",        min_value=0.0, max_value=0.5, step=0.01, format="%.2f"),
            "Peso Esp. (kN/m³)":     st.column_config.NumberColumn("γ (kN/m³)",         min_value=10.0, max_value=25.0, step=0.5),
            "Peso Esp. Sat (kN/m³)": st.column_config.NumberColumn("γsat (kN/m³)",      min_value=10.0, max_value=25.0, step=0.5),
        }
    )
    if not df_edit.equals(st.session_state.df_terreno):
        st.session_state.df_terreno = df_edit
        st.session_state.calculo_realizado = False
        st.rerun()

    st.markdown("---")
    st.header("2. Resultados")

    if st.session_state.calculo_realizado:
        df_res = st.session_state.df_resultados

        st.dataframe(
            df_res,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Capa":        st.column_config.TextColumn("Capa",        width="medium"),
                "z Techo [m]": st.column_config.NumberColumn("z Techo [m]", width="small", format="%.2f"),
                "z Base [m]":  st.column_config.NumberColumn("z Base [m]",  width="small", format="%.2f"),
                "h_ef [m]":    st.column_config.NumberColumn("h_ef [m]",    width="small", format="%.2f"),
                "z_mid [m]":   st.column_config.NumberColumn("z_mid [m]",   width="small", format="%.3f"),
                "Δσz [kPa]":   st.column_config.NumberColumn("Δσz [kPa]",  width="small", format="%.2f"),
                "Δσx [kPa]":   st.column_config.NumberColumn("Δσx [kPa]",  width="small", format="%.2f"),
                "Δσy [kPa]":   st.column_config.NumberColumn("Δσy [kPa]",  width="small", format="%.2f"),
                "Δεz [-]":     st.column_config.NumberColumn("Δεz [-]",     width="small", format="%.6f"),
                "Δs [mm]":     st.column_config.NumberColumn("Δs [mm]",     width="medium",format="%.3f"),
            }
        )

        st.markdown("---")
        st.success("✅ Cálculo completado.")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("🏗️ Asiento Total Estimado",
                      f"{st.session_state.asiento_total * 1000:.3f} mm")
        with col2:
            st.metric("📐 Profundidad de Influencia z_i (EC7)",
                      f"{z_i_ec7:.2f} m")

        st.markdown("---")
        st.bar_chart(df_res.set_index("Capa")["Δs [mm]"])

    else:
        st.info("👈 Pulsa **Calcular Asiento Total** en el panel izquierdo.")

# --- VISTA 2: DETALLE POR ESTRATO ---
elif modo_vista == "📋 Detalle por Estrato":
    st.header("📋 Cálculos Detallados por Estrato")

    if not st.session_state.calculo_realizado:
        st.warning("⚠️ Calcula primero el asiento total.")
    else:
        df_res = st.session_state.df_resultados

        cfg4 = {"format": "%.4f", "width": "small"}
        cfg3 = {"format": "%.3f", "width": "small"}
        cfg2 = {"format": "%.2f", "width": "small"}

        # Tabla tensiones
        st.markdown("##### ⚡ Tensiones de Holl en z_mid (centro de la zapata)")
        df_tens = df_res[["Capa", "z_mid [m]", "Δσz [kPa]", "Δσx [kPa]", "Δσy [kPa]"]].copy()
        st.dataframe(df_tens, use_container_width=True, hide_index=True,
                     column_config={
                         "Capa":       st.column_config.TextColumn("Capa",    width="medium"),
                         "z_mid [m]":  st.column_config.NumberColumn("z_mid [m]",  **cfg3),
                         "Δσz [kPa]":  st.column_config.NumberColumn("Δσz [kPa]",  **cfg3),
                         "Δσx [kPa]":  st.column_config.NumberColumn("Δσx [kPa]",  **cfg3),
                         "Δσy [kPa]":  st.column_config.NumberColumn("Δσy [kPa]",  **cfg3),
                     })

        # Tabla deformaciones
        st.markdown("##### 📐 Deformación unitaria y asiento (Ec. 68)")
        df_def = df_res[["Capa", "h_ef [m]", "Δεz [-]", "Δs [mm]"]].copy()
        st.dataframe(df_def, use_container_width=True, hide_index=True,
                     column_config={
                         "Capa":    st.column_config.TextColumn("Capa",    width="medium"),
                         "h_ef [m]":st.column_config.NumberColumn("h_ef [m]", **cfg3),
                         "Δεz [-]": st.column_config.NumberColumn("Δεz [-]",  format="%.6f", width="medium"),
                         "Δs [mm]": st.column_config.NumberColumn("Δs [mm]",  **cfg3),
                     })

        st.info(f"**Asiento total = Σ Δs** = **{st.session_state.asiento_total * 1000:.3f} mm**")

# --- VISTA 3: BULBO DE PRESIONES ---
elif modo_vista == "📉 Bulbo de Presiones":
    st.header("Bulbo de Presiones y Zona de Influencia EC7")
    st.write(
        r"Tensiones bajo el **centro** de la zapata ($\times 4$ superposición de cuadrantes $B/2 \times L/2$) "
        r"y criterio EC7: $\Delta\sigma_z \leq 0.20 \cdot \sigma'_{v0}$."
    )

    col1, col2 = st.columns([1, 3])
    with col1:
        z_max_gr = st.slider("Profundidad máxima a graficar [m]:",
                             min_value=1.0, max_value=espesor_total,
                             value=min(espesor_total, 15.0), step=0.5)
        st.markdown(f"**p:** {p} kPa · **B:** {B} m · **L:** {L} m")
        if NF < 100.0:
            st.markdown(f"**NF:** {NF:.1f} m")
        st.metric("📐 z_i (EC7)", f"{z_i_ec7:.2f} m")

    with col2:
        z_vals = np.linspace(0.05, z_max_gr, 200)
        sz_v, sx_v, sy_v, sv0_v, umb20_v = [], [], [], [], []
        for z in z_vals:
            sz, sx, sy = holl_centro(p, B, L, z)
            sz_v.append(sz); sx_v.append(sx); sy_v.append(sy)
            sv0 = calcular_sigma_v0(z, st.session_state.df_terreno, NF)
            sv0_v.append(sv0)
            umb20_v.append(0.20 * sv0)

        fig, ax = plt.subplots(figsize=(9, 7))
        ax.plot(sz_v,   z_vals, label=r'Vertical $\Delta\sigma_z$',           color='red',         lw=2)
        ax.plot(sx_v,   z_vals, label=r'Horiz. Trans. $\Delta\sigma_x$',      color='blue',        ls='--')
        ax.plot(sy_v,   z_vals, label=r'Horiz. Long. $\Delta\sigma_y$',       color='purple',      ls='-.')
        ax.plot(sv0_v,  z_vals, label=r"$\sigma'_{v0}$ (tensión efectiva)",   color='saddlebrown', ls=':', lw=1.5)
        ax.plot(umb20_v,z_vals, label=r"$0.20\,\sigma'_{v0}$ (criterio EC7)", color='green',       lw=2)

        if z_i_ec7 <= z_max_gr:
            ax.axhline(y=z_i_ec7, color='orange', ls='--', lw=1.5,
                       label=f'z_i EC7 = {z_i_ec7:.2f} m')
            ax.annotate(f' z_i = {z_i_ec7:.2f} m',
                        xy=(0, z_i_ec7), xytext=(p * 0.04, z_i_ec7 - z_max_gr * 0.03),
                        color='darkorange', fontsize=10, fontweight='bold')
        if NF < z_max_gr and NF < 100.0:
            ax.axhline(y=NF, color='deepskyblue', ls='-.', lw=1.2,
                       label=f'NF = {NF:.1f} m')

        ax.set_ylim(z_max_gr, 0)
        ax.set_xlim(left=0)
        ax.set_xlabel("Tensión (kPa)", fontsize=11)
        ax.set_ylabel("Profundidad z (m)", fontsize=11)
        ax.set_title("Bulbo de presiones — Centro de la cimentación", fontsize=13)
        ax.legend(loc='lower right', fontsize=9)
        ax.grid(True, linestyle=':', alpha=0.5)
        ax.spines[['top', 'right']].set_visible(False)
        st.pyplot(fig)
        plt.close(fig)

# --- VISTA 4: TEORÍA ---
elif modo_vista == "📖 Fundamento Teórico":
    st.header("Fundamento Teórico — Metodología EC7 (Ec. 68)")

    st.subheader("1. Fórmula de Asiento (Ec. 68)")
    st.markdown(
        "El asiento elástico se calcula integrando la deformación unitaria vertical "
        "en cada estrato de la estratigrafía:"
    )
    st.latex(r"s = \sum_{i=1}^{n} \left[ \Delta\varepsilon_z \cdot h \right]_i"
             r"= \sum_{i=1}^{n} \left[ \frac{h}{E} \left( \Delta\sigma_z - \nu(\Delta\sigma_x + \Delta\sigma_y) \right) \right]_i")
    st.markdown(r"""
Donde para cada estrato $i$:
- $h_i$ → espesor del estrato (m)
- $E_i$ → módulo de elasticidad (kPa)
- $\nu_i$ → coeficiente de Poisson
- $\Delta\sigma_z, \Delta\sigma_x, \Delta\sigma_y$ → incrementos de tensión en el punto medio del estrato

Las tensiones se evalúan en **$z_{mid} = (z_{techo} + z_{base})/2$** de cada estrato.
""")

    st.subheader("2. Tensiones de Holl — Bajo la Esquina")
    st.markdown(
        r"La solución de Holl da los incrementos de tensión bajo la **esquina** "
        r"de una carga rectangular $B \times L$ a profundidad $z$:"
    )
    st.latex(r"\sigma_z = \frac{p}{2\pi}\left[\arctan\frac{BL}{zR_3} + BL\left(\frac{1}{R_1^2}+\frac{1}{R_2^2}\right)\frac{z}{R_3}\right]")
    st.latex(r"\sigma_x = \frac{p}{2\pi}\left[\arctan\frac{BL}{zR_3} - \frac{BLz}{R_1^2 R_3}\right]")
    st.latex(r"\sigma_y = \frac{p}{2\pi}\left[\arctan\frac{BL}{zR_3} - \frac{BLz}{R_2^2 R_3}\right]")
    st.latex(r"R_1 = \sqrt{L^2+z^2} \quad R_2 = \sqrt{B^2+z^2} \quad R_3 = \sqrt{L^2+B^2+z^2}")

    st.subheader("3. Tensiones en el Centro — Superposición de Cuadrantes")
    st.markdown(r"""
Para obtener las tensiones bajo el **centro** de la zapata se aplica el principio
de superposición (UNE-EN 1997-1, 6.6.2(15)):

> Se divide el área en **4 cuadrantes iguales** de $B/2 \times L/2$, se calcula la
> tensión en la esquina de cada uno y se suman los 4 valores.

$$\Delta\sigma_{z,centro}(z) = 4 \cdot \sigma_z\!\left(\frac{B}{2}, \frac{L}{2}, z\right)$$

Lo mismo aplica para $\Delta\sigma_x$ y $\Delta\sigma_y$.
""")

    st.subheader("4. Profundidad de Influencia (Criterio EC7)")
    st.markdown(r"La profundidad de influencia $z_i$ se determina donde:")
    st.latex(r"\Delta\sigma_z(z_i) \leq 0.20 \cdot \sigma'_{v0}(z_i)")
    st.markdown(r"""
Con $\sigma'_{v0}$ = tensión efectiva vertical geoestática, considerando el nivel freático:
$$\sigma'_{v0}(z) = \sum_i \gamma'_i \cdot \Delta z_i \qquad
  \gamma'_i = \begin{cases} \gamma_i & \text{sobre NF} \\ \gamma_{sat,i}-\gamma_w & \text{bajo NF} \end{cases}$$
""")
