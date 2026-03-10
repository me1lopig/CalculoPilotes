# Programa para el cálculo de valores de Adherencia Límite para Micropilotes
# Según la Guía para el proyecto y la ejecución de Micropilotes
# en Obras de Carreteras
# Dirección General de Carreteras
# Ministerio de Fomento
# Gobierno de España

import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Cálculo Geotécnico Micropilotes", layout="wide")

st.title("🧮 Calculadora de valores de Adherencia Límite para Micropilotes")
st.markdown("Cálculo basado en la **Guía para el proyecto y ejecución de micropilotes en obras de carreteras**.")
st.markdown("Desarrollado por la DT del Departamento de Geotecnia ")
st.markdown("Version febrero de 2026")

# --- CONSTANTES Y REGLAS DE CONTROL ---
PARAM_ARENAS = ["Plim (MPa)", "SPT (N)"]
PARAM_ARCILLAS = ["Plim (MPa)", "qu (MPa)"]
OPCIONES_TODAS = ["Plim (MPa)", "SPT (N)", "qu (MPa)"]

LIMITES = {
    "Arenas y Gravas": {
        "Plim (MPa)": {"min": 0.5, "max": 7.0},
        "SPT (N)": {"min": 10, "max": 100}
    },
    "Arcillas y Limos": {
        "Plim (MPa)": {"min": 0.25, "max": 2.5},
        "qu (MPa)": {"min": 0.05, "max": 0.5}
    }
}

# --- DATOS DE SECCIONES Y MATERIALES (INCRUSTADOS) ---

DATA_DIAMETROS = """Dp de Perforación (mm);D del Micropilote (mm);de de la Armadura Tubular (mm) (*)
120;114,3;60,3 - 73,0
140;133,0;60,3 - 73,0 - 88,9
160;152,4;73,0 - 88,9 - 101,6
185;177,8;88,9 - 101,6 - 114,3 - 127,0
200;193,7;101,6 - 114,3 - 127,0 - 139,0
225;219,1;114,3 - 127,0 - 139,0 - 168,3"""

DATA_ACEROS = """EN 10025;API 5CT;Límite elástico min (N/mm2);Resistencia a tracción (N/mm2);Alargamiento min (%)
S235;-;235;340-470;26
S275;-;275;410-560;22
S355;-;355;490-630;22
S550;-;550;600-750;16
-;J55;379;517;24
-;K55;379;655;19
-;N80;552-758;689;18"""

DATA_CS550 = """Diametro Exterior (mm);Espesor (mm);Modulo Resistente (cm3);Momento de Inercia (cm4);Seccion Acero (cm2);Peso (kg/m);Carga Adm. Compresion (Tn)
43;3;3.52;7.58;3.77;3;19.8
48;4;5.62;13.48;5.53;5;29.1
60.3;3.2;7.78;23.46;5.74;5;30.2
60.3;6;12.66;38.16;10.23;8;53.8
73;6;19.56;71.40;12.62;11;66.4
73;6.5;20.75;75.74;13.57;12;71.4
73;7;21.88;79.88;14.51;12;76.3
73;8;23.98;87.54;16.33;13;85.8
73;8.5;24.95;91.08;17.22;14;90.5
73;9;25.87;94.43;18.09;15;95.1
73;10;27.57;100.62;19.78;16;104.0
88.9;3;16.81;74.73;8.09;7;42.5
88.9;6;30.34;134.87;15.62;13;82.1
88.9;6.5;32.31;143.62;16.82;14;88.4
88.9;7;34.20;152.04;18.00;14;94.6
88.9;7.5;36.02;160.12;19.17;16;100.8
88.9;8;37.77;167.88;20.32;16;106.8
88.9;8.5;39.44;175.33;21.46;17;112.8
88.9;9;41.05;182.47;22.58;19;118.7
88.9;9.5;42.59;189.32;23.69;19;124.5
88.9;10;44.07;195.88;24.77;20;130.2
88.9;12;49.36;219.41;28.98;23;152.3
101.6;6;40.66;206.57;18.01;15;94.7
101.6;7;46.04;233.87;20.79;17;109.3
101.6;7.5;48.59;246.84;22.16;18;116.5
101.6;8;51.06;259.37;23.51;19;123.6
101.6;9;55.74;283.14;26.17;21;137.6
101.6;10;60.09;305.26;28.76;23;151.2
101.6;12.5;69.66;353.87;34.97;28;183.8
114.3;7;59.64;340.86;23.58;19;124.0
114.3;8;66.37;379.30;26.70;21;140.4
114.3;9;72.70;415.46;29.76;24;156.4
114.3;10;78.64;449.43;32.75;26;172.2
127;8;83.71;531.53;29.89;25;157.1
127;9;91.93;583.78;33.35;28;175.3
127;10;99.72;633.23;36.74;29;193.1
127;12;114.04;724.13;43.33;35;227.8"""

DATA_CSTM80 = """Diametro Exterior (mm);Espesor (mm);Modulo Resistente (cm3);Momento de Inercia (cm4);Seccion Acero (cm2);Peso (kg/m);Carga Adm. Compresion (Tn)
139.7;7;92.17;643.81;29.17;23;153.3
139.7;7.5;97.68;682.32;31.13;25;163.7
139.7;8;103.07;719.92;33.08;26;173.9
139.7;9;113.45;792.43;36.94;31;194.2
139.7;10;123.33;861.46;40.73;33;214.1
139.7;11;132.73;927.11;44.45;36;233.7
139.7;12;141.66;989.49;48.12;38;253.0
139.7;12.5;145.95;1019.49;49.93;40;262.5
152.4;6;97.14;740.19;27.58;22;145.0
152.4;9;137.24;1045.77;40.52;35;213.0
152.4;10;149.47;1138.95;44.71;36;235.1
168.3;7;137.28;1155.20;35.45;28;186.4
168.3;9;170.24;1432.56;45.02;36;236.7
168.3;10;185.76;1563.19;49.71;40;261.3
168.3;11;200.67;1688.64;54.33;43;285.6
168.3;16;266.55;2242.98;76.52;61;402.2
177.8;9;191.66;1703.85;47.70;38;250.8
177.8;9.5;200.58;1783.19;50.20;40;263.9
177.8;10;209.34;1861.04;52.69;42;277.0
177.8;11;226.36;2012.36;57.61;46;302.9
177.8;12.5;250.69;2228.66;64.88;51;341.1
193.7;6.3;168.22;1629.22;37.07;30;194.9
193.7;12;293.01;2837.76;68.46;54;359.9
203;6;177.55;1802.16;37.11;30;195.1
203;9;254.66;2584.76;54.82;44;288.2
203;12;324.61;3294.82;71.97;57;378.3
219;6;208.00;2277.58;40.13;32;211.0
219;8;269.77;2953.93;53.00;42;278.6
219;9;299.31;3277.46;59.35;47;312.0
219;10;327.99;3591.46;65.63;52;345.0
219;11;355.81;3896.15;71.84;57;377.7
219;20;570.66;6248.76;124.97;99;657.0
244;10;412.97;5038.25;73.48;58;386.3
244;11.5;466.13;5686.80;83.96;66;441.4
244;12.5;500.39;6104.74;90.86;72;477.7
273;6;328.56;4484.81;50.30;40;298.0
273;14;846.82;12617.60;124.85;99;656.3"""

# --- FUNCIONES DE CÁLCULO ---
def calcular_arenas(plim):
    x_irs = [0.0, 0.5, 2.8, 10.0]; y_irs = [0.0, 0.14, 0.62, 0.62]
    x_ir  = [0.0, 0.5, 3.3, 10.0]; y_ir  = [0.0, 0.09, 0.50, 0.50]
    x_iu  = [0.0, 0.5, 3.5, 10.0]; y_iu  = [0.0, 0.04, 0.40, 0.40]
    return np.interp(plim, x_iu, y_iu), np.interp(plim, x_ir, y_ir), np.interp(plim, x_irs, y_irs)

def calcular_arcillas(plim):
    x_irs = [0.0, 0.25, 0.5, 1.8, 10.0]; y_irs = [0.0, 0.12, 0.20, 0.40, 0.40]
    x_ir  = [0.0, 0.25, 0.5, 2.1, 10.0]; y_ir  = [0.0, 0.07, 0.12, 0.30, 0.30]
    x_iu  = [0.0, 0.25, 0.5, 2.3, 10.0]; y_iu  = [0.0, 0.05, 0.08, 0.20, 0.20]
    return np.interp(plim, x_iu, y_iu), np.interp(plim, x_ir, y_ir), np.interp(plim, x_irs, y_irs)

# --- CREACIÓN DE PESTAÑAS ---
tab1, tab2 = st.tabs(["📊 Cálculos de Adherencia", "⚙️ Catálogos y Secciones"])

# ==========================================
# PESTAÑA 1: CÁLCULOS DE ADHERENCIA
# ==========================================
with tab1:
    st.subheader("📝 1. Entrada de Datos")

    if 'df_data' not in st.session_state:
        st.session_state.df_data = pd.DataFrame([
            {"Nombre": "UG-01", "Suelo": "Arenas y Gravas", "Ensayo": "SPT (N)", "Valor": 45.0},
            {"Nombre": "UG-02", "Suelo": "Arcillas y Limos", "Ensayo": "qu (MPa)", "Valor": 0.25}
        ])

    df_input = st.data_editor(
        st.session_state.df_data,
        num_rows="dynamic",
        column_config={
            "Suelo": st.column_config.SelectboxColumn("Tipo de Suelo", options=list(LIMITES.keys()), required=True),
            "Ensayo": st.column_config.SelectboxColumn("Parámetro/Ensayo", options=OPCIONES_TODAS, required=True),
            "Valor": st.column_config.NumberColumn("Valor Numérico", format="%.2f", min_value=0.0)
        },
        use_container_width=True
    )

    # --- PROCESAMIENTO Y CONTROLES ---
    results = []
    valid_inputs_report = []

    for i, row in df_input.iterrows():
        if pd.isna(row["Valor"]) or not row["Nombre"]: continue
        
        s_tipo = row["Suelo"]
        e_tipo = row["Ensayo"]
        val = row["Valor"]
        
        if s_tipo == "Arenas y Gravas" and e_tipo not in PARAM_ARENAS:
            st.error(f"❌ Error en '{row['Nombre']}': Para Arenas usar Plim o SPT.")
            continue
        if s_tipo == "Arcillas y Limos" and e_tipo not in PARAM_ARCILLAS:
            st.error(f"❌ Error en '{row['Nombre']}': Para Arcillas usar Plim o qu.")
            continue

        lims = LIMITES[s_tipo][e_tipo]
        if val < lims["min"] or val > lims["max"]:
            st.error(f"⚠️ Fila '{row['Nombre']}': Valor fuera de rango ({lims['min']} a {lims['max']}).")
            continue

        plim = val if "Plim" in e_tipo else (val / 20.0 if "SPT" in e_tipo else val * 5.0)
        iu, ir, irs = calcular_arenas(plim) if s_tipo == "Arenas y Gravas" else calcular_arcillas(plim)
        
        valid_inputs_report.append({"Nombre": row["Nombre"], "Suelo": s_tipo, "Ensayo": e_tipo, "Valor": val})
        results.append({
            "Nombre": row["Nombre"], "IU (MPa)": round(iu, 3), 
            "IR (MPa)": round(ir, 3), "IRS (MPa)": round(irs, 3), "Grupo": s_tipo, "_p": plim
        })

    df_res = pd.DataFrame(results)
    df_in_rep = pd.DataFrame(valid_inputs_report)

    # --- GRÁFICOS Y DESCARGA ---
    def dibujar_grafico(grupo, df_datos):
        is_arenas = (grupo == "Arenas y Gravas")
        fig, ax1 = plt.subplots(figsize=(8, 6))
        x_max, p_min = (7.0, 0.5) if is_arenas else (2.5, 0.25)
        x_norm = np.linspace(p_min, x_max, 400)
        y_norm = calcular_arenas(x_norm) if is_arenas else calcular_arcillas(x_norm)
        
        ax1.plot(x_norm, y_norm[2], 'k-', label='IRS', linewidth=2)
        ax1.plot(x_norm, y_norm[1], 'k-.', label='IR')
        ax1.plot(x_norm, y_norm[0], 'k--', label='IU')

        if not df_datos.empty:
            colores = plt.cm.tab10(np.linspace(0, 1, len(df_datos)))
            for i, (_, s) in enumerate(df_datos.iterrows()):
                ax1.scatter([s["_p"]]*3, [s["IU (MPa)"], s["IR (MPa)"], s["IRS (MPa)"]], 
                            color=colores[i], label=s["Nombre"], s=60, edgecolors='white', zorder=5)
        
        ax1.set_title(f"Tipo de Suelo: {grupo}")
        ax1.set_xlabel("Plim (MPa)")
        ax1.set_ylabel("rf,lim (MPa)")
        ax1.set_xlim(0, x_max); ax1.set_ylim(0, 0.8 if is_arenas else 0.5)
        ax1.grid(True, alpha=0.3); ax1.legend(fontsize='x-small')
        
        ax2 = ax1.twiny()
        ax2.set_xlim(ax1.get_xlim()); ax2.xaxis.set_ticks_position('bottom'); ax2.xaxis.set_label_position('bottom')
        ax2.spines['bottom'].set_position(('outward', 40))
        ticks = ax1.get_xticks()
        ax2.set_xticks(ticks)
        if is_arenas:
            ax2.set_xticklabels([f"{int(t*20)}" for t in ticks]); ax2.set_xlabel("N (SPT)")
        else:
            ax2.set_xticklabels([f"{t/5:.2f}" for t in ticks]); ax2.set_xlabel("qu (MPa)")
        return fig

    def generar_word(df_in, df_out, f1, f2):
        doc = Document()
        doc.add_heading('Informe de resultados de Adherencia Límite', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for titulo, df in [("1. Datos de Entrada", df_in), ("2. Resultados de Cálculo", df_out)]:
            doc.add_heading(titulo, level=1)
            t = doc.add_table(df.shape[0] + 1, df.shape[1])
            t.style = 'Table Grid'
            for i, col in enumerate(df.columns): t.cell(0, i).text = col
            for i, row in enumerate(df.values):
                for j, val in enumerate(row): t.cell(i+1, j).text = str(val)

        doc.add_heading('3. Gráficos de Cálculo', level=1)
        for n, f in [("Arenas", f1), ("Arcillas", f2)]:
            buf = BytesIO(); f.savefig(buf, format='png', dpi=200, bbox_inches='tight')
            doc.add_picture(buf, width=Inches(5.0))
        
        out = BytesIO(); doc.save(out); out.seek(0)
        return out

    if not df_res.empty:
        st.subheader("📋 2. Resultados de cálculo de adherencia límite")
        st.dataframe(df_res.drop(columns=["_p", "Grupo"]), use_container_width=True)
        
        fig1 = dibujar_grafico("Arenas y Gravas", df_res[df_res["Grupo"] == "Arenas y Gravas"])
        fig2 = dibujar_grafico("Arcillas y Limos", df_res[df_res["Grupo"] == "Arcillas y Limos"])
        
        informe = generar_word(df_in_rep, df_res.drop(columns=["_p", "Grupo"]), fig1, fig2)
        st.download_button("📥 Descargar Informe (.docx)", informe, "Informe_Micropilotes.docx", type="primary")
    else:
        fig1 = dibujar_grafico("Arenas y Gravas", pd.DataFrame())
        fig2 = dibujar_grafico("Arcillas y Limos", pd.DataFrame())

    st.divider()
    c1, c2 = st.columns(2)
    with c1: st.pyplot(fig1)
    with c2: st.pyplot(fig2)

# ==========================================
# PESTAÑA 2: SECCIONES Y MATERIALES
# ==========================================
with tab2:
    st.subheader("🕳️ Diámetros Habituales de Perforación y Armadura")
    st.markdown("Relación entre el diámetro de perforación ($D_P$), el diámetro final del micropilote ($D$) y los diámetros exteriores de la armadura tubular ($d_e$).")
    df_diametros = pd.read_csv(StringIO(DATA_DIAMETROS), sep=";")
    st.dataframe(df_diametros, use_container_width=True, hide_index=True)
    st.caption("(*) El valor mayor de cada una de las filas de diámetros exteriores de armadura tubular $d_e$, será válido únicamente en uniones roscadas.")

    st.divider()

    st.subheader("🛠️ Tipos de Acero más Habituales")
    st.markdown("Consulta las propiedades mecánicas de los aceros estructurales y de entubación (API 5CT).")
    
    # Cargar y mostrar la tabla de aceros
    df_aceros = pd.read_csv(StringIO(DATA_ACEROS), sep=";")
    st.dataframe(df_aceros, use_container_width=True, hide_index=True)
    
    st.divider()
    
    st.subheader("📏 Catálogo de Secciones de Tubería para Micropilotes")
    st.markdown("Consulta las propiedades mecánicas y geométricas de las armaduras tubulares según su calidad de acero.")
    
    # Cargar las tablas de tuberías
    df_cs550 = pd.read_csv(StringIO(DATA_CS550), sep=";")
    df_cstm80 = pd.read_csv(StringIO(DATA_CSTM80), sep=";")
    
    col_a, col_b = st.columns(2)
    
    with col_a:
        st.markdown("### Calidad CS550")
        st.dataframe(df_cs550, use_container_width=True, hide_index=True)
        
    with col_b:
        st.markdown("### Calidad CSTM80")
        st.dataframe(df_cstm80, use_container_width=True, hide_index=True)