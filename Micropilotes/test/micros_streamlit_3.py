import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE LA PÁGINA ---
st.set_page_config(page_title="Cálculo de Adherencia límite (Micropilotes)", layout="wide")

st.title("🧮 Calculadora de Adherencia límite por Fuste")
st.markdown("""
Esta herramienta calcula la **Adherencia límite** ($r_{f,lim}$) utilizando **modelos lineales a trozos** ajustados a las gráficas de la normativa.
""")

# --- FUNCIONES DE CÁLCULO (PIECEWISE LINEAR) ---

def calcular_arenas(plim):
    """Modelo lineal a trozos para Arenas y Gravas"""
    x_irs = [0.0, 0.5, 2.8, 10.0]; y_irs = [0.0, 0.14, 0.62, 0.62]
    x_ir  = [0.0, 0.5, 3.3, 10.0]; y_ir  = [0.0, 0.09, 0.50, 0.50]
    x_iu  = [0.0, 0.5, 3.5, 10.0]; y_iu  = [0.0, 0.04, 0.40, 0.40]
    
    tau_irs = np.interp(plim, x_irs, y_irs)
    tau_ir  = np.interp(plim, x_ir, y_ir)
    tau_iu  = np.interp(plim, x_iu, y_iu)
    return tau_iu, tau_ir, tau_irs

def calcular_arcillas(plim):
    """Modelo lineal a trozos para Arcillas y Limos"""
    x_irs = [0.0, 0.25, 0.5, 1.0, 1.8, 10.0]; y_irs = [0.0, 0.12, 0.20, 0.30, 0.40, 0.40]
    x_ir  = [0.0, 0.25, 0.5, 1.0, 2.1, 10.0]; y_ir  = [0.0, 0.07, 0.12, 0.18, 0.30, 0.30]
    x_iu  = [0.0, 0.25, 0.5, 1.0, 2.3, 10.0]; y_iu  = [0.0, 0.05, 0.08, 0.12, 0.20, 0.20]
    
    tau_irs = np.interp(plim, x_irs, y_irs)
    tau_ir  = np.interp(plim, x_ir, y_ir)
    tau_iu  = np.interp(plim, x_iu, y_iu)
    return tau_iu, tau_ir, tau_irs

# --- FUNCIÓN PARA CREAR EL GRÁFICO (DOBLE EJE) ---

def crear_figura(tipo_suelo, plim_calculo, res_iu, res_ir, res_irs, x_max_plot, p_transicion):
    fig, ax1 = plt.subplots(figsize=(8, 6))
    
    x_solid = np.linspace(p_transicion, x_max_plot, 200)
    x_dotted = np.linspace(0, p_transicion, 50)
    
    calc_func = calcular_arenas if tipo_suelo == "Arenas y Gravas" else calcular_arcillas
    y_sol = calc_func(x_solid)
    y_dot = calc_func(x_dotted)
    
    # Dibujar líneas
    labels = ['IU', 'IR', 'IRS']
    styles = ['--', '-.', '-']
    for i in range(3):
        ax1.plot(x_solid, y_sol[i], 'k', linestyle=styles[i], label=labels[i], linewidth=2)
        ax1.plot(x_dotted, y_dot[i], 'k', linestyle=':', alpha=0.4, linewidth=1)
    
    # Indicador de cálculo
    ax1.axvline(x=plim_calculo, color='red', linestyle=':', alpha=0.6)
    ax1.scatter([plim_calculo]*3, [res_iu, res_ir, res_irs], color='red', zorder=5)
    
    ax1.set_xlabel('Presión límite $P_{lim}$ (MPa)')
    ax1.set_ylabel('Rozamiento unitario límite $r_{f,lim}$ (MPa)')
    ax1.set_xlim(0, x_max_plot)
    ax1.set_ylim(0, 0.8 if tipo_suelo == "Arenas y Gravas" else 0.45)
    ax1.grid(True, which='both', linestyle='--', alpha=0.3)
    ax1.legend()

    # Eje secundario sincronizado
    ax2 = ax1.twiny()
    ax2.set_xlim(ax1.get_xlim())
    ax2.xaxis.set_ticks_position('bottom')
    ax2.xaxis.set_label_position('bottom')
    ax2.spines['bottom'].set_position(('outward', 45))
    
    ticks = ax1.get_xticks()
    ax2.set_xticks(ticks)
    if tipo_suelo == "Arenas y Gravas":
        ax2.set_xticklabels([f"{int(t * 20)}" for t in ticks])
        ax2.set_xlabel("Índice SPT (N)")
    else:
        ax2.set_xticklabels([f"{t / 5:.2f}" for t in ticks])
        ax2.set_xlabel("Compresión Simple $q_u$ (MPa)")
    
    fig.tight_layout()
    return fig

# --- FUNCIÓN GENERAR INFORME DOCX ---

def generar_informe(tipo_suelo, input_label, res_iu, res_ir, res_irs, fig):
    doc = Document()
    titulo = doc.add_heading('Informe de Cálculo: Adherencia Límite en Micropilotes', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Datos de Entrada', level=1)
    doc.add_paragraph(f'Tipo de Suelo: {tipo_suelo}')
    doc.add_paragraph(f'Valor introducido: {input_label}')

    doc.add_heading('2. Resultados de Adherencia Límite (rf,lim)', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Procedimiento de Inyección'
    hdr_cells[1].text = 'rf,lim (MPa)'

    data = [
        ('IRS (Inyección Repetitiva Selectiva)', f'{res_irs:.3f}'),
        ('IR (Inyección Repetitiva)', f'{res_ir:.3f}'),
        ('IU (Inyección Única)', f'{res_iu:.3f}')
    ]
    for item, valor in data:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = valor

    doc.add_heading('3. Gráfico de Cálculo', level=1)
    img_buf = BytesIO()
    fig.savefig(img_buf, format='png', dpi=200, bbox_inches='tight')
    doc.add_picture(img_buf, width=Inches(5.5))
    
    doc_buf = BytesIO()
    doc.save(doc_buf)
    doc_buf.seek(0)
    return doc_buf

# --- BARRA LATERAL (ENTRADAS RESTAURADAS) ---

st.sidebar.header("⚙️ Datos de Entrada")
tipo_suelo = st.sidebar.selectbox("Selecciona el Tipo de Suelo:", ("Arenas y Gravas", "Arcillas y Limos"))

plim_calculo = 0.0
input_label = ""

if tipo_suelo == "Arenas y Gravas":
    st.sidebar.info("Rango: $P_{lim} \geq 0.5$ MPa ($N \geq 10$)")
    tipo_dato = st.sidebar.radio("Entrada:", ["Presión Límite (Plim)", "Índice SPT (N)"])
    if tipo_dato == "Presión Límite (Plim)":
        val = st.sidebar.number_input("Plim (MPa)", 0.5, 7.0, 2.0, 0.1)
        plim_calculo = val
        input_label = f"Plim = {val} MPa"
    else:
        val = st.sidebar.number_input("SPT (N)", 10, 100, 30, 1)
        plim_calculo = val / 20.0
        input_label = f"N = {val} (Plim ≈ {plim_calculo:.2f} MPa)"
    x_max_plot, p_transicion = 7.0, 0.5
else:
    st.sidebar.info("Rango: $P_{lim} \geq 0.25$ MPa ($q_u \geq 0.05$ MPa)")
    tipo_dato = st.sidebar.radio("Entrada:", ["Presión Límite (Plim)", "Compresión Simple (qu)"])
    if tipo_dato == "Presión Límite (Plim)":
        val = st.sidebar.number_input("Plim (MPa)", 0.25, 2.5, 1.0, 0.05)
        plim_calculo = val
        input_label = f"Plim = {val} MPa"
    else:
        val = st.sidebar.number_input("qu (MPa)", 0.05, 0.5, 0.2, 0.01)
        plim_calculo = val * 5.0
        input_label = f"qu = {val} MPa (Plim ≈ {plim_calculo:.2f} MPa)"
    x_max_plot, p_transicion = 2.5, 0.25

# --- CÁLCULOS Y VISUALIZACIÓN ---

res_iu, res_ir, res_irs = calcular_arenas(plim_calculo) if tipo_suelo == "Arenas y Gravas" else calcular_arcillas(plim_calculo)
figura_final = crear_figura(tipo_suelo, plim_calculo, res_iu, res_ir, res_irs, x_max_plot, p_transicion)

col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("📋 Resultados")
    st.metric("IRS", f"{res_irs:.3f} MPa")
    st.metric("IR", f"{res_ir:.3f} MPa")
    st.metric("IU", f"{res_iu:.3f} MPa")
    
    st.markdown("---")
    informe_word = generar_informe(tipo_suelo, input_label, res_iu, res_ir, res_irs, figura_final)
    st.download_button(
        label="📄 Descargar Informe Word",
        data=informe_word,
        file_name=f"Informe_Micropilote_{tipo_suelo.replace(' ','_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

with col2:
    st.subheader("📈 Gráfico de Cálculo")
    st.pyplot(figura_final)