import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Configuración de la página
st.set_page_config(page_title="Clasificación Potencial Expansivo", layout="wide")

# ==========================================
# UTILIDADES WORD Y ESTILOS
# ==========================================
def set_cell_bg(cell, color_hex):
    """Establece el color de fondo de una celda en Word."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def safe_format(val):
    if val is None or pd.isna(val) or val == "" or val == "---":
        return "--"
    if isinstance(val, (int, float)):
        return f"{val:.2f}"
    return str(val).replace("\\>", ">")

def insertar_tabla_referencia(doc, df, titulo):
    doc.add_heading(titulo, level=2)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = 'Table Grid'
    for j, col_name in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col_name)
        set_cell_bg(cell, "D9E2F3")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            val = df.iloc[i, j]
            table.cell(i + 1, j).text = safe_format(val)
            if table.cell(i + 1, j).paragraphs[0].runs:
                table.cell(i + 1, j).paragraphs[0].runs[0].font.size = Pt(9)

NIVELES = ["BAJA", "MEDIA", "ALTA", "MUY ALTA"]

# ==========================================
# TABLAS DE REFERENCIA
# ==========================================
df_ref_chen = pd.DataFrame({
    "Grado Expansión": ["Baja", "Media", "Alta", "Muy alta"],
    "% #200 (Finos)": ["< 30", "30 - 60", "60 - 95", "> 95"],
    "LL (%)": ["< 30", "30 - 40", "40 - 60", "> 60"],
    "IP (%)": ["< 15", "15 - 25", "25 - 35", "> 35"],
    "Presión (kg/cm²)": ["< 0.5", "1.5 - 2.5", "2.5 - 10", "> 10"]
})

# Tabla exacta de la imagen proporcionada (Ortiz)
df_ref_ortiz = pd.DataFrame({
    "Grado Expansión": ["Baja", "Media", "Alta", "Muy alta"],
    "Retracción (%)": ["> 15", "12 - 16", "8 - 12", "< 10"],
    "IP (%)": ["< 18", "15 - 28", "25 - 40", "> 35"],
    "LL(%)": ["< 30", "30 - 40", "40 - 60", "> 60"],
    "% #200": ["< 30", "30 - 60", "60 - 95", "> 95"],
    "Coloides (<0.001mm)": ["< 15", "13 - 23", "20 - 30", "> 30"]
})

# ==========================================
# LÓGICA DE CLASIFICACIÓN
# ==========================================
def safe_float(val):
    if val is None or pd.isna(val) or val == "": return None
    try:
        if isinstance(val, str): val = val.replace(",", ".").strip()
        return float(val)
    except ValueError: return None

def clasificar_parametro_chen(valor, tipo):
    val = safe_float(valor)
    if val is None: return None
    
    if tipo == "LL":
        if val > 60: return "MUY ALTA"
        elif val >= 40: return "ALTA"
        elif val >= 30: return "MEDIA"
        return "BAJA"
    if tipo == "FINOS":
        if val > 95: return "MUY ALTA"
        elif val >= 60: return "ALTA"
        elif val >= 30: return "MEDIA"
        return "BAJA"
    if tipo == "IP":
        if val > 35: return "MUY ALTA"
        elif val >= 25: return "ALTA"
        elif val >= 15: return "MEDIA"
        return "BAJA"
    return None

def clasificar_parametro_ortiz(valor, tipo):
    """Clasifica según la tabla de Rodríguez Ortiz (Imagen)."""
    val = safe_float(valor)
    if val is None: return None

    # Límite de Retracción (Inverso: menor valor = mayor riesgo)
    if tipo == "RETRACCION":
        if val < 10: return "MUY ALTA"
        if val <= 12: return "ALTA"   # Priorizamos Alta en solape
        if val <= 16: return "MEDIA"
        return "BAJA"

    # IP (Solapes: Media 15-28, Alta 25-40, Muy Alta >35)
    if tipo == "IP":
        if val > 35: return "MUY ALTA"
        if val >= 25: return "ALTA"
        if val >= 15: return "MEDIA"
        return "BAJA"

    # LL
    if tipo == "LL":
        if val > 60: return "MUY ALTA"
        if val >= 40: return "ALTA"
        if val >= 30: return "MEDIA"
        return "BAJA"

    # #200
    if tipo == "FINOS":
        if val > 95: return "MUY ALTA"
        if val >= 60: return "ALTA"
        if val >= 30: return "MEDIA"
        return "BAJA"

    # Coloides (< 0.001 mm)
    if tipo == "COLOIDES":
        if val > 30: return "MUY ALTA"
        if val >= 20: return "ALTA"
        if val >= 13: return "MEDIA"
        return "BAJA"

    return None

def obtener_est_chen(clasif):
    mapping = {
        "MUY ALTA": ("> 10.00", "> 10.00 kg/cm²"), 
        "ALTA": ("3.00 - 10.00", "2.50 - 10.00 kg/cm²"), 
        "MEDIA": ("1.00 - 5.00", "1.50 - 2.50 kg/cm²"), 
        "BAJA": ("< 1.00", "< 0.50 kg/cm²"), 
        "---": ("", "")
    }
    return mapping.get(clasif, ("", ""))

# ==========================================
# GENERADOR INFORME WORD
# ==========================================
def generar_informe_word(df_chen, df_ortiz, ref_chen, ref_ortiz):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    t = doc.add_heading('INFORME RESULTADOS DE EXPANSIVIDAD', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- CHEN ---
    doc.add_heading('1. Método Chen (1988)', level=1)
    insertar_tabla_referencia(doc, ref_chen, "1.1 Criterios de Clasificación")

    doc.add_heading("1.2 Resultados Chen", level=2)
    t_c = doc.add_table(rows=1, cols=8)
    t_c.style = "Table Grid"
    h_c = ["ID", "Finos", "LL", "IP", "Niv. Finos", "Niv. LL", "Niv. IP", "CLASIFICACIÓN"]
    for j, h in enumerate(h_c):
        cell = t_c.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "D9E2F3")
        cell.paragraphs[0].runs[0].bold = True

    for _, row in df_chen.iterrows():
        cells = t_c.add_row().cells
        vals = [row["ID"], row['Finos #200'], row['LL'], row['IP'], 
                row.get("Nivel_Finos"), row.get("Nivel_LL"), row.get("Nivel_IP"), row['Clasificación']]
        for i, v in enumerate(vals):
            cells[i].text = safe_format(v)
        if row['Clasificación'] in ["ALTA", "MUY ALTA"]:
            set_cell_bg(cells[7], "FFCCCC")

    # --- ORTIZ ---
    doc.add_heading('2. Método Rodríguez Ortiz (1975)', level=1)
    insertar_tabla_referencia(doc, ref_ortiz, "2.1 Criterios de Clasificación")

    doc.add_heading("2.2 Resultados R. Ortiz", level=2)
    t_o = doc.add_table(rows=1, cols=7)
    t_o.style = "Table Grid"
    h_o = ["ID", "Retr.", "IP", "LL", "#200", "Col.", "CLASIFICACIÓN"]
    for j, h in enumerate(h_o):
        cell = t_o.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "D9E2F3")
        cell.paragraphs[0].runs[0].bold = True

    for _, row in df_ortiz.iterrows():
        cells = t_o.add_row().cells
        vals = [row["ID"], row['Retr.'], row['IP'], row['LL'], row['#200'], row['Col.'], row['Clasificación']]
        for i, v in enumerate(vals):
            cells[i].text = safe_format(v)
        if row['Clasificación'] in ["ALTA", "MUY ALTA"]:
            set_cell_bg(cells[6], "FFCCCC")

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# ==========================================
# INTERFAZ (SIDEBAR CON DICCIONARIO)
# ==========================================
st.sidebar.title("🛠️ Panel de Control")
st.sidebar.subheader("Métodos")
st.sidebar.info("Se calculan ambas clasificaciones de forma independiente.")

# --- DICCIONARIO EN LA BARRA LATERAL ---
st.sidebar.divider()
st.sidebar.subheader("📖 Diccionario")
with st.sidebar.expander("Ver términos", expanded=True):
    st.markdown("""
    * **ID**: Identificador de muestra.
    * **LL**: Límite Líquido (%).
    * **LP**: Límite Plástico (%).
    * **Retr**: Límite de Retracción (%).
    * **#200**: % Pasa tamiz #200 (Finos).
    * **Col**: % Coloides (< 0.001 mm).
    """)

download_btn_container = st.sidebar.empty()

# ==========================================
# ÁREA PRINCIPAL
# ==========================================
st.title("🏗️ Clasificación del Potencial Expansivo")

st.subheader("1. Entrada de Datos")

if 'data_app' not in st.session_state:
    st.session_state.data_app = pd.DataFrame({
        "ID": ["Sondeo A", "Sondeo B"], "LL": [65.00, 35.00], "LP": [25.00, 20.00],
        "Retracción": [None, None], "% Pasa #200": [None, None], "Coloides": [None, None]
    })

df_in = st.data_editor(st.session_state.data_app, num_rows="dynamic", use_container_width=True)

if not df_in.empty:
    res_chen, res_ortiz = [], []
    
    for _, row in df_in.iterrows():
        # Extracción de datos
        id_m = row["ID"]
        ll = safe_float(row["LL"])
        lp = safe_float(row["LP"])
        retr = safe_float(row["Retracción"])
        f200 = safe_float(row["% Pasa #200"])
        col = safe_float(row["Coloides"])
        ip = ll - lp if (ll is not None and lp is not None) else None

        # --- CHEN ---
        vc = {
            "LL": clasificar_parametro_chen(ll, "LL"),
            "Finos": clasificar_parametro_chen(f200, "FINOS"),
            "IP": clasificar_parametro_chen(ip, "IP")
        }
        vc_ok = {k: v for k, v in vc.items() if v}
        cl_chen = max(vc_ok.values(), key=lambda x: NIVELES.index(x)) if vc_ok else "---"
        crit_chen = [k for k, v in vc_ok.items() if v == cl_chen]
        _, pres_chen = obtener_est_chen(cl_chen)
        
        res_chen.append({
            "ID": id_m, "Finos #200": f200, "LL": ll, "IP": ip,
            "Nivel_Finos": vc.get("Finos"), "Nivel_LL": vc.get("LL"), "Nivel_IP": vc.get("IP"),
            "Clasificación": cl_chen, "Presión": pres_chen, "Críticos": crit_chen
        })

        # --- ORTIZ ---
        vo = {
            "Retr.": clasificar_parametro_ortiz(retr, "RETRACCION"),
            "IP": clasificar_parametro_ortiz(ip, "IP"),
            "LL": clasificar_parametro_ortiz(ll, "LL"),
            "#200": clasificar_parametro_ortiz(f200, "FINOS"),
            "Col.": clasificar_parametro_ortiz(col, "COLOIDES")
        }
        vo_ok = {k: v for k, v in vo.items() if v}
        cl_ortiz = max(vo_ok.values(), key=lambda x: NIVELES.index(x)) if vo_ok else "---"
        crit_ortiz = [k for k, v in vo_ok.items() if v == cl_ortiz]
        
        res_ortiz.append({
            "ID": id_m, "Retr.": retr, "IP": ip, "LL": ll, "#200": f200, "Col.": col,
            "Clasificación": cl_ortiz, "Críticos": crit_ortiz
        })

    # --- RESULTADOS ---
    def highlight_bad(row):
        styles = ['' for _ in row.index]
        if row['Clasificación'] in ["ALTA", "MUY ALTA"]:
            styles[row.index.get_loc('Clasificación')] = 'background-color: #ffcccc'
        return styles

    st.subheader("2. Resultados: Método Chen")
    st.dataframe(pd.DataFrame(res_chen).style.apply(highlight_bad, axis=1).format("{:.2f}", subset=["Finos #200", "LL", "IP"], na_rep="-"), use_container_width=True, hide_index=True)
    with st.expander("Ver Referencia Chen"): st.dataframe(df_ref_chen, hide_index=True)

    st.subheader("3. Resultados: Método R. Ortiz")
    st.dataframe(pd.DataFrame(res_ortiz).style.apply(highlight_bad, axis=1).format("{:.2f}", subset=["Retr.", "IP", "LL", "#200", "Col."], na_rep="-"), use_container_width=True, hide_index=True)
    with st.expander("Ver Referencia Ortiz (Imagen Actualizada)"): st.dataframe(df_ref_ortiz, hide_index=True)

    try:
        w_f = generar_informe_word(pd.DataFrame(res_chen), pd.DataFrame(res_ortiz), df_ref_chen, df_ref_ortiz)
        download_btn_container.download_button("📄 Descargar Word", data=w_f, file_name="Informe_Expansividad.docx", type="primary", use_container_width=True)
    except Exception as e: st.error(f"Error Word: {e}")