import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Configuración de la página
st.set_page_config(page_title="Geotech - Diagnóstico Crítico", layout="wide")

# ==========================================
# UTILIDADES WORD Y ESTILOS
# ==========================================
def set_cell_bg(cell, color_hex):
    """Establece el color de fondo de una celda en Word."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def insertar_tabla_referencia(doc, df, titulo):
    """Inserta una tabla normativa en el Word."""
    doc.add_heading(titulo, level=2)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = 'Table Grid'
    for j, col_name in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = col_name
        set_cell_bg(cell, "D9E2F3")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iloc[i, j])
            # Aplicar fuente solo si hay texto/run
            if table.cell(i + 1, j).paragraphs[0].runs:
                table.cell(i + 1, j).paragraphs[0].runs[0].font.size = Pt(9)

NIVELES = ["BAJA", "MEDIA", "ALTA", "MUY ALTA"]

# ==========================================
# TABLAS DE REFERENCIA
# ==========================================
df_ref_chen = pd.DataFrame({
    "Grado": ["Bajo", "Medio", "Alto", "Muy alto"],
    "% #200": ["< 30", "30 - 60", "60 - 95", "> 95"],
    "LL (%)": ["< 30", "30 - 40", "40 - 60", "> 60"],
    "Exp. Prob %": ["< 1", "1 - 5", "3 - 10", "> 10"],
    "Presión (kg/cm²)": ["< 0.5", "1.5 - 2.5", "2.5 - 10", "> 10"]
})

df_ref_ortiz = pd.DataFrame({
    "Expansividad": ["Baja", "Media", "Alta", "Muy alta"],
    "Retracción": ["> 15", "12 - 16", "8 - 12", "< 10"],
    "Ip": ["< 18", "15 - 28", "25 - 40", "> 35"],
    "WL (LL)": ["< 30", "30 - 40", "40 - 60", "> 60"],
    "Presión (kg/cm²)": ["< 0.3", "0.3 - 1.2", "1.2 - 3.0", "> 3"],
    "Hinch. Sup (cm)": ["0 - 1", "1 - 3", "3 - 7", "> 7"]
})

# ==========================================
# LÓGICA DE CLASIFICACIÓN
# ==========================================
def clasificar_parametro(valor, tipo):
    if valor is None or pd.isna(valor): return None
    val = float(valor)
    if tipo in ["LL_CHEN", "LL_ORTIZ"]:
        if val > 60: return "MUY ALTA"
        elif val >= 40: return "ALTA"
        elif val >= 30: return "MEDIA"
        return "BAJA"
    if tipo in ["FINOS_CHEN", "FINOS_ORTIZ"]:
        if val > 95: return "MUY ALTA"
        elif val >= 60: return "ALTA"
        elif val >= 30: return "MEDIA"
        return "BAJA"
    if tipo == "IP":
        if val > 35: return "MUY ALTA"
        elif val >= 25: return "ALTA"
        elif val >= 15: return "MEDIA"
        return "BAJA"
    if tipo == "RETRACCION":
        if val < 10: return "MUY ALTA"
        elif val <= 12: return "ALTA"
        elif val <= 16: return "MEDIA"
        return "BAJA"
    if tipo == "COLOIDES":
        if val > 30: return "MUY ALTA"
        elif val >= 20: return "ALTA"
        elif val >= 13: return "MEDIA"
        return "BAJA"
    return None

def obtener_est_chen(clasif):
    mapping = {"MUY ALTA": ("> 10.00", "> 10.00 kg/cm²"), "ALTA": ("3.00 - 10.00", "2.50 - 10.00 kg/cm²"), "MEDIA": ("1.00 - 5.00", "1.50 - 2.50 kg/cm²"), "BAJA": ("< 1.00", "< 0.50 kg/cm²"), "---": ("", "")}
    return mapping.get(clasif, ("", ""))

def obtener_est_ortiz(clasif):
    mapping = {"MUY ALTA": ("> 3.00 kg/cm²", "> 7.00 cm"), "ALTA": ("1.20 - 3.00 kg/cm²", "3.00 - 7.00 cm"), "MEDIA": ("0.30 - 1.20 kg/cm²", "1.00 - 3.00 cm"), "BAJA": ("< 0.30 kg/cm²", "0.00 - 1.00 cm"), "---": ("", "")}
    return mapping.get(clasif, ("", ""))

# ==========================================
# GENERADOR REPORTE WORD (4 TABLAS OBLIGATORIAS)
# ==========================================
def generar_reporte_word(df_chen, df_ortiz):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    t = doc.add_heading('INFORME TÉCNICO DE EXPANSIVIDAD', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. SECCIÓN DE NORMATIVA
    doc.add_heading('1. Criterios de Clasificación (Normativa)', level=1)
    insertar_tabla_referencia(doc, df_ref_chen, "1.1 Criterios de Chen (1988)")
    insertar_tabla_referencia(doc, df_ref_ortiz, "1.2 Criterios de R. Ortiz (1975)")
    doc.add_page_break()

    # 2. SECCIÓN DE RESULTADOS
    doc.add_heading('2. Resultados de las Muestras Analizadas', level=1)
    
    # 2.1 Resultados Chen
    doc.add_heading('2.1 Evaluación según Método Chen', level=2)
    t_c = doc.add_table(rows=1, cols=5)
    t_c.style = 'Table Grid'
    h_c = ['ID', 'Finos #200', 'LL (%)', 'CLASIFICACIÓN', 'Presión Hinch.']
    for i, h in enumerate(h_c):
        cell = t_c.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "D9E2F3")
        cell.paragraphs[0].runs[0].bold = True

    for _, row in df_chen.iterrows():
        cells = t_c.add_row().cells
        cells[0].text = str(row['ID'])
        cells[1].text = f"{row['Finos #200']:.2f}" if pd.notnull(row['Finos #200']) else ""
        cells[2].text = f"{row['LL']:.2f}" if pd.notnull(row['LL']) else ""
        cells[3].text = row['Clasificación'] if row['Clasificación'] != "---" else ""
        cells[4].text = row['Presión']
        if row['Clasificación'] in ["ALTA", "MUY ALTA"]:
            set_cell_bg(cells[3], "FFCCCC")
            if pd.notnull(row['LL']) and "LL" in row['Críticos']: set_cell_bg(cells[2], "FFCCCC")
            if pd.notnull(row['Finos #200']) and "Finos" in row['Críticos']: set_cell_bg(cells[1], "FFCCCC")

    # 2.2 Resultados Ortiz
    doc.add_heading('2.2 Evaluación según Método R. Ortiz', level=2)
    t_o = doc.add_table(rows=1, cols=8)
    t_o.style = 'Table Grid'
    h_o = ['ID', 'Retr.', 'IP', 'LL', '#200', 'Col.', 'DIAGNÓSTICO', 'Presión']
    for i, h in enumerate(h_o):
        cell = t_o.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "D9E2F3")
        cell.paragraphs[0].runs[0].bold = True

    for _, row in df_ortiz.iterrows():
        cells = t_o.add_row().cells
        vals = [row['ID'], row['Retr.'], row['IP'], row['LL'], row['#200'], row['Col.'], row['Clasificación'], row['Presión']]
        for i, v in enumerate(vals):
            if isinstance(v, (int, float)) and pd.notnull(v):
                cells[i].text = f"{v:.2f}"
            elif v == "---" or v is None or pd.isna(v):
                cells[i].text = ""
            else:
                cells[i].text = str(v)
            
            # Formateo seguro de fuente (solo si hay texto)
            p = cells[i].paragraphs[0]
            if p.runs:
                p.runs[0].font.size = Pt(9)
        
        if row['Clasificación'] in ["ALTA", "MUY ALTA"]:
            set_cell_bg(cells[6], "FFCCCC")
            map_cols = {"Retracción": 1, "IP": 2, "LL": 3, "Finos": 4, "Coloides": 5}
            for crit in row['Críticos']:
                if crit in map_cols and cells[map_cols[crit]].text != "":
                    set_cell_bg(cells[map_cols[crit]], "FFCCCC")

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# ==========================================
# INTERFAZ (LATERAL RESTAURADO)
# ==========================================
st.sidebar.title("🛠️ Panel de Control")

st.sidebar.subheader("Métodos de Análisis")
st.sidebar.markdown("""
* **Método de Chen (1988):** Evalúa la expansividad basándose en el porcentaje de finos y el Límite Líquido.
* **Método de R. Ortiz (1975):** Proporciona un diagnóstico integral analizando parámetros físicos y mecánicos combinados.
""")

st.sidebar.subheader("Valores de Entrada")
st.sidebar.info("""
Introduce los datos de laboratorio en la tabla central. Las casillas que dejes vacías no serán consideradas para el diagnóstico final, garantizando la flexibilidad del análisis.
""")

download_btn_container = st.sidebar.empty()

# ==========================================
# ÁREA PRINCIPAL
# ==========================================
st.title("🏗️ Diagnóstico de Expansividad Profesional")

if 'data_app' not in st.session_state:
    st.session_state.data_app = pd.DataFrame({
        "ID": ["Sondeo A", "Sondeo B"], "LL": [65.00, 35.00], "LP": [25.00, 20.00],
        "Retracción": [9.00, 18.00], "% Pasa #200": [98.00, 50.00], "Coloides": [35.00, None]
    })

st.subheader("1. Entrada de Datos de Laboratorio")
df_in = st.data_editor(st.session_state.data_app, num_rows="dynamic", use_container_width=True)

if not df_in.empty:
    res_c, res_o = [], []
    for _, row in df_in.iterrows():
        # Lógica Chen
        v_c = {"LL": clasificar_parametro(row["LL"], "LL_CHEN"), "Finos": clasificar_parametro(row["% Pasa #200"], "FINOS_CHEN")}
        v_c_val = {k: v for k, v in v_c.items() if v}
        peor_c = max(v_c_val.values(), key=lambda x: NIVELES.index(x)) if v_c_val else "---"
        crit_c = [k for k, v in v_c_val.items() if v == peor_c]
        exp_c, pre_c = obtener_est_chen(peor_c)
        res_c.append({"ID": row["ID"], "Finos #200": row["% Pasa #200"], "LL": row["LL"], "Clasificación": peor_c, "Presión": pre_c, "Críticos": crit_c})

        # Lógica Ortiz
        ip = row["LL"] - row["LP"] if pd.notnull(row["LL"]) and pd.notnull(row["LP"]) else None
        v_o = {"Retracción": clasificar_parametro(row["Retracción"], "RETRACCION"), "IP": clasificar_parametro(ip, "IP"), "LL": clasificar_parametro(row["LL"], "LL_ORTIZ"), "Finos": clasificar_parametro(row["% Pasa #200"], "FINOS_ORTIZ"), "Coloides": clasificar_parametro(row["Coloides"], "COLOIDES")}
        v_o_val = {k: v for k, v in v_o.items() if v}
        peor_o = max(v_o_val.values(), key=lambda x: NIVELES.index(x)) if v_o_val else "---"
        crit_o = [k for k, v in v_o_val.items() if v == peor_o]
        pre_o, hin_o = obtener_est_ortiz(peor_o)
        res_o.append({"ID": row["ID"], "Retr.": row["Retracción"], "IP": ip, "LL": row["LL"], "#200": row["% Pasa #200"], "Col.": row["Coloides"], "Clasificación": peor_o, "Presión": pre_o, "Críticos": crit_o})

    # Resultados UI con resaltado crítico
    st.subheader("2. Resultados: Método Chen")
    st.dataframe(pd.DataFrame(res_c).style.apply(lambda r: ['background-color: #ffcccc' if (c in r['Críticos'] or c=='Clasificación') and r['Clasificación'] in ["ALTA", "MUY ALTA"] else '' for c in r.index], axis=1).format({"Finos #200": "{:.2f}", "LL": "{:.2f}"}, na_rep="-"), use_container_width=True)
    with st.expander("Ver Tabla Referencia Chen"): st.table(df_ref_chen)

    st.subheader("3. Resultados: Método R. Ortiz")
    st.dataframe(pd.DataFrame(res_o).style.apply(lambda r: ['background-color: #ffcccc' if (c in r['Críticos'] or c=='Clasificación') and r['Clasificación'] in ["ALTA", "MUY ALTA"] else '' for c in r.index], axis=1).format("{:.2f}", subset=["Retr.", "IP", "LL", "#200", "Col."], na_rep="-"), use_container_width=True)
    with st.expander("Ver Tabla Referencia Ortiz"): st.table(df_ref_ortiz)

    # Botón de Descarga
    try:
        w_f = generar_reporte_word(pd.DataFrame(res_c), pd.DataFrame(res_o))
        download_btn_container.download_button("📄 Descargar Informe Word", data=w_f, file_name="Informe_Expansividad_Final.docx", type="primary", use_container_width=True)
    except Exception as e: st.error(f"Error Word: {e}")